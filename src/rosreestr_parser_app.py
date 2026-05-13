# -*- coding: utf-8 -*-
"""
Локальный парсер справочной информации Росреестра.

Что делает:
- открывает видимый браузер;
- даёт пользователю вручную авторизоваться и вручную пройти капчу, если она появится;
- перебирает квартиры по адресу;
- берёт только актуальные объекты с точным совпадением адреса;
- сохраняет результат в Excel по шаблону.

Важно: программа не обходит капчу, не заказывает выписки и не выполняет юридически значимые действия.
"""

from __future__ import annotations

import json
import os
import queue
import random
import re
import sys
import threading
import time
import traceback
import hashlib
import webbrowser
import urllib.request
from copy import copy
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Iterable

import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk


APP_TITLE = "Парсер Росреестра — V56"
APP_VERSION_CODE = 56

REQUIRED_HEADERS = [
    "Номер помещения (обяз.)",
    "Площадь помещения, м2 (обяз.)",
    "Кадастровый номер",
    "Признак совместной собс. (\"да\" для сов. собс.)",
    "Доля (1/1, 1/2, 1/3 и т.д.) (обяз., сумма в помещении должна быть 1/1)",
    "Муницип. собс. (\"да\" для муниц. собс.)",
    "Номер регистрации (обяз.)",
    "Дата регистрации",
    "Фамилия / Наименование организации (обяз.)",
    "Имя (обяз. для физ лиц)",
    "Отчество",
    "ИНН (обяз. для юр лиц)",
    "ОГРН (ОГРНИП)",
    "Снилс",
]

ERROR_HEADERS = [
    "Дата/время",
    "Номер помещения",
    "Поисковый адрес",
    "Статус",
    "Причина",
    "Кадастровый номер",
    "Адрес из Росреестра",
]

BUILDING_LAND_HEADERS = [
    "Дата/время",
    "Вид объекта",
    "Статус объекта",
    "Кадастровый номер",
    "Адрес (местоположение)",
    "Площадь, кв.м",
    "Назначение / ВРИ",
    "Кадастровая стоимость",
    "Источник поиска",
]

SEARCH_RESULTS_HEADERS = [
    "Дата/время",
    "Поисковый запрос",
    "Номер помещения",
    "Кадастровый номер",
    "Текст строки выдачи",
    "Решение",
    "Причина",
]

DISPUTED_HEADERS = [
    "Дата/время",
    "Поисковый запрос",
    "Номер помещения",
    "Кадастровый номер",
    "Текст строки выдачи / адрес",
    "Что сделал парсер",
    "Причина",
    "Выбранный кадастровый номер",
]

REJECTED_HEADERS = DISPUTED_HEADERS

DUPLICATE_HEADERS = [
    "Дата/время",
    "Тип дубля",
    "Ключ",
    "Где найдено",
    "Кадастровый номер",
    "Номер регистрации",
    "Комментарий",
]

SUMMARY_HEADERS = [
    "Показатель",
    "Значение",
]

def snt_log_headers(headers: list[str]) -> list[str]:
    """Для СНТ в служебных листах заменяет ОСС-термины на участки."""
    return [
        str(h).replace("Номер помещения", "Номер участка").replace("Поисковый адрес", "Поисковый запрос")
        for h in headers
    ]


VALIDATION_HEADERS = [
    "Дата/время",
    "Уровень",
    "Лист",
    "Строка",
    "Поле",
    "Сообщение",
]

CADASTRAL_ONLY_HEADERS = [
    "Дата/время",
    "Поисковый запрос",
    "Номер помещения",
    "Кадастровый номер",
    "Текст строки выдачи / адрес",
    "Статус принятия",
    "Причина",
]

FULL_INFO_HEADERS = [
    "Дата/время",
    "Поисковый запрос",
    "Номер помещения",
    "Кадастровый номер",
    "Адрес (местоположение)",
    "Вид объекта",
    "Статус объекта",
    "Раздел",
    "Поле",
    "Значение",
    "Порядок",
]

SNT_LAND_HEADERS = [
    "Дата/время",
    "Номер участка",
    "Кадастровый номер",
    "Статус объекта",
    "Адрес",
    "Площадь",
    "Категория земель",
    "Вид разрешенного использования",
    "Кадастровая стоимость",
    "Дата определения кадастровой стоимости",
    "Дата внесения кадастровой стоимости",
    "Вид права",
    "Номер регистрации",
    "Дата регистрации",
    "Поисковый запрос",
    "Комментарий",
]

ROSKVARTAL_HEADERS = [
    "Номер помещения",
    "Площадь помещения",
    "Тип помещения",
    "ФИО собственника",
    "Название организации",
    "Доля",
    "Тип документа",
    "Номер регистрации",
    "Дата регистрации",
    "Email",
    "Телефон",
    "Уникальное помещение",
]

ROSKVARTAL_INTERNAL_DATA_HEADERS = [
    "Тип помещения",
    "Вид права",
]

OSS_EXPORT_FORMATS = {
    "Burmistr": "burmistr",
    "Roskvartal": "roskvartal",
}

OSS_EXPORT_FORMAT_LABELS = {v: k for k, v in OSS_EXPORT_FORMATS.items()}


PARSER_MODES = {
    "Реестр для ОСС": "oss",
    "Реестр всех сведений": "full",
    "Реестр кадастровых номеров": "cadastral",
    "Реестр земельных участков / СНТ": "snt",
}

PARSER_MODE_LABELS = {v: k for k, v in PARSER_MODES.items()}

DEFAULT_CONFIG = {
    "login_url": "https://lk.rosreestr.ru/login",
    "service_url": "https://lk.rosreestr.ru/eservices/real-estate-objects-online",
    "address_prefix": "г Город ул Улица д 1 кв",
    "search_numbered_units": True,
    "try_address_variants": True,
    "strict_street_match": False,
    "strict_house_match": True,
    "include_unnumbered_house_objects": False,
    "unnumbered_max_pages": 10,
    "auto_output_filename": True,
    "pause_preset": "Обычно",
    "parser_mode": "oss",
    "oss_export_format": "burmistr",
    "match_oss_export_format": "burmistr",
    "collect_cadastral_only": False,
    "start_flat": 1,
    "end_flat": 100,
    "extra_unit_numbers": "",
    "manual_unit_numbers": "",
    "manual_unit_file_path": "",
    "snt_city": "",
    "snt_name": "",
    "snt_cadastral_quarters": "",
    "snt_include_actual": True,
    "snt_include_redeemed": True,
    "snt_only_land": False,
    "snt_write_buildings": True,
    "snt_write_not_found": True,
    "snt_append_existing": True,
    "snt_try_address_variants": True,
    "pause_min_seconds": 3,
    "pause_max_seconds": 7,
    "retry_count": 3,
    "retry_delay_seconds": 10,
    "rate_limit_pause_seconds": 300,
    "rate_limit_retry_count": 12,
    "browser_profile_dir": "profile_rosreestr",
    "browser_channel": "chrome",
    "template_path": "template_burmistr.xlsx",
    "output_path": "output/result.xlsx",
    "slow_mo_ms": 120,
    "save_after_each_flat": True,
    "close_browser_after_finish": False,
    "resume_enabled": True,
    "skip_done_units": True,
    "skip_done_cadastrals": True,
    "match_new_registry_path": "",
    "match_old_registry_path": "",
    "match_output_path": "",
    "match_transfer_shares": True,
    "match_allow_contracts_for_empty_right": True,
    "match_expand_empty_right_rows": True,
    "match_preserve_existing_fio": True,
    "match_allow_address_surname": False,
    "last_dir_new_registry": "",
    "last_dir_old_registry": "",
    "last_dir_match_output": "",
    "last_dir_template": "",
    "last_dir_output": "",
    "last_dir_manual_queue": "",
    "last_dir_browser_profile": "",
    "repository_url": "https://github.com/serbo26-lab/rosreestr-parser",
    "update_manifest_url": "https://raw.githubusercontent.com/serbo26-lab/rosreestr-parser/main/latest.json",
    "update_channel": "stable",
}


@dataclass
class RightInfo:
    right_type: str
    reg_number: str
    reg_date: str


class RateLimitError(RuntimeError):
    pass


class StopRequested(RuntimeError):
    pass


@dataclass
class ObjectInfo:
    flat_no: str
    area: str
    cadastral_number: str
    ownership_form: str
    address: str
    object_type: str
    status: str
    purpose: str
    floor: str
    cadastral_cost: str
    rights: list[RightInfo]
    raw_text: str = ""


@dataclass
class OldRegistryRecord:
    source_file: str
    sheet: str
    row: int
    flat_no: str
    cadastral_number: str
    right_key: str
    reg_number: str
    reg_date: str
    is_contract: bool
    contract_text: str
    fio_original: str
    last_name: str
    first_name: str
    patronymic: str
    share: str
    raw_right: str
    confidence: str
    old_address: str = ""


# ----------------------------- paths / config -----------------------------


def app_dir() -> Path:
    """Папка программы. Для exe PyInstaller учитываем sys._MEIPASS."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parents[1]


def resource_path(name: str) -> Path:
    """Путь к ресурсу рядом с exe или внутри временной папки PyInstaller."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        p = Path(sys._MEIPASS) / name  # type: ignore[attr-defined]
        if p.exists():
            return p
    return app_dir() / name


def load_config(config_path: Path) -> dict[str, Any]:
    cfg = DEFAULT_CONFIG.copy()
    if config_path.exists():
        try:
            with config_path.open("r", encoding="utf-8") as f:
                loaded = json.load(f)
            if isinstance(loaded, dict):
                cfg.update(loaded)
        except Exception:
            pass
    return cfg


def save_config(config_path: Path, cfg: dict[str, Any]) -> None:
    config_path.parent.mkdir(parents=True, exist_ok=True)
    with config_path.open("w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


# ----------------------------- text helpers -----------------------------


def normalize_text(s: Any) -> str:
    s = "" if s is None else str(s)
    s = s.replace("ё", "е").replace("Ё", "Е")
    s = re.sub(r"\s+", " ", s)
    return s.strip().lower()


def compact_text(s: Any) -> str:
    return re.sub(r"[^0-9a-zа-я]+", "", normalize_text(s))


def extract_flat_no(address: str) -> str:
    """Извлекает номер квартиры/помещения: кв 1, квартира №1, пом. 1, помещение 1, №1."""
    t = normalize_text(address)
    pattern_value = r"([0-9]+[а-яa-z0-9/-]*|[а-яa-z][0-9]+[а-яa-z0-9/-]*)"
    m = re.search(rf"(?:кв\.?|квартира|пом\.?|помещение)[\s\.№:-]*{pattern_value}\b", t, re.I)
    if m:
        return m.group(1).upper()
    # Некоторые адреса пишут как "..., №1" без слова "кв/пом".
    m = re.search(rf"(?:^|[\s,])№\s*{pattern_value}\b", t, re.I)
    if m:
        return m.group(1).upper()
    return ""


def address_has_trailing_unit(address: str, expected_no: str, search_house_no: str = "") -> bool:
    """
    Понимает редкий формат адреса без маркера помещения: '..., д 66, 1'.

    Важно: не считаем голый адрес дома '..., ул Улица, 7' квартирой 7.
    Такая ошибка была причиной того, что земельный участок/дом подтягивался вместо кв. 7.
    Без явного маркера кв/пом допускаем совпадение только если в результате есть
    отдельный дом и отдельный хвост помещения, например 'д 7, 7'.
    """
    expected = normalize_text(expected_no).strip().upper()
    if not expected:
        return False

    text = normalize_text(address)
    parts = [p.strip(" .№") for p in re.split(r",+", text) if p.strip()]
    if len(parts) < 2:
        return False

    last = parts[-1].strip(" .№")
    if last.upper() != expected:
        return False

    # Если последний фрагмент сам является "д 7" / "дом 7" — это дом, не помещение.
    if re.search(r"(?:^|\s)(д|дом)\s*\.?\s*№?\s*" + re.escape(last) + r"$", last):
        return False

    # Если в адресе нет явного дома "д/дом", а хвост равен номеру дома из запроса,
    # это почти наверняка голый адрес дома/земельного участка: "... ул Улица, 7".
    # Не принимаем его как квартиру 7.
    if search_house_no and last.upper() == str(search_house_no).strip().upper():
        has_explicit_house = bool(re.search(r"(?:^|[\s,])(?:д\.?|дом)\s*№?\s*" + re.escape(str(search_house_no).strip()), text, re.I))
        occurrences = len(re.findall(r"(?<!\d)" + re.escape(str(search_house_no).strip()) + r"(?!\d)", text))
        if not has_explicit_house and occurrences <= 1:
            return False

    return True


def has_extra_trailing_unit_without_marker(address: str, search_house_no: str = "") -> bool:
    """
    Для поиска по голому адресу дома: отсекает результаты вида 'д 7, 7',
    то есть когда после дома есть отдельный номер без слов кв/пом/№.
    Голый адрес 'ул Улица, 7' при этом не отсекается.
    """
    text = normalize_text(address)
    parts = [p.strip(" .№") for p in re.split(r",+", text) if p.strip()]
    if len(parts) < 2:
        return False
    last = parts[-1].strip(" .№")
    if not re.fullmatch(r"[0-9]+[а-яa-z0-9/-]*", last, flags=re.I):
        return False
    if re.search(r"(?:^|\s)(д|дом)\s*\.?\s*№?\s*" + re.escape(last) + r"$", last):
        return False

    # Если дом указан явно, и после него отдельный номер — это не голый объект дома,
    # а помещение/квартира с номером без слова.
    if search_house_no:
        has_explicit_house = bool(re.search(r"(?:^|[\s,])(?:д\.?|дом)\s*№?\s*" + re.escape(str(search_house_no).strip()), text, re.I))
        if has_explicit_house:
            return True

    return False


def extract_house_no(address: str) -> str:
    """
    Номер дома.
    Сначала ищем явный маркер д/дом, затем fallback для адресов Росреестра вида:
    "ул Улица, 10, кв 66" — здесь дом 10, а не квартира 66.
    """
    t = normalize_text(address)
    m = re.search(r"(?:д|дом)[\s\.№:-]*([0-9]+[а-яa-z/-]?)\b", t, re.I)
    if m:
        return m.group(1).upper()
    # Росреестр иногда пишет дом без 'д': 'ул Улица, 10, кв 66'.
    m = re.search(r"(?:ул\.?|улица)\s+.+?[,\s]+([0-9]+[а-яa-z/-]?)\s*,?\s*(?:кв\.?|квартира|пом\.?|помещение|№)\b", t, re.I)
    if m:
        return m.group(1).upper()
    return ""


def extract_street_name(address: str) -> str:
    """Извлекает название улицы для строгого сравнения: название А != название Б."""
    t = normalize_text(address)
    # Берём текст между маркером улицы и домом/квартирой/помещением.
    m = re.search(r"(?:^|[,\s])(?:ул\.?|улица)\s+(.+?)(?=\s*,?\s*(?:д\.?|дом|кв\.?|квартира|пом\.?|помещение|№)\b|$)", t, re.I)
    if not m:
        return ""
    street = m.group(1).strip(" ,")
    street = re.sub(r"\s+", " ", street)
    # Убираем служебные хвосты, если regex не поймал дом.
    street = re.sub(r"\s+(?:д\.?|дом)\s*.*$", "", street, flags=re.I).strip(" ,")
    return street


def street_matches(search_address: str, result_address: str, *, strict: bool = False) -> bool:
    if not strict:
        return True
    q = extract_street_name(search_address)
    r = extract_street_name(result_address)
    if not q or not r:
        # Если улицу не удалось достать, не считаем это совпадением в строгом режиме.
        return False
    return compact_text(q) == compact_text(r)


def house_address_from_prefix(address_prefix: str) -> str:
    """Возвращает голый адрес дома без хвоста 'кв/пом/№'."""
    prefix = re.sub(r"\s+", " ", str(address_prefix or "")).strip(" ,")
    if not prefix:
        return ""
    # Убираем явный маркер квартиры в конце.
    prefix = re.sub(
        r"(?:,\s*)?(?:кв|квартира|к\.|пом(?:ещение)?|пом\.)\s*\.?\s*$",
        "",
        prefix,
        flags=re.I,
    ).strip(" ,")
    # Убираем хвост вида ', №' если пользователь задал его как маркер.
    prefix = re.sub(r"(?:,\s*)?№\s*$", "", prefix).strip(" ,")
    return prefix


def filename_slug(value: str, max_len: int = 90) -> str:
    """Безопасная часть имени файла Windows: без точек, запятых и запрещённых символов."""
    v = str(value or "").replace("ё", "е").replace("Ё", "Е")
    v = re.sub(r"[<>:\"/\\|?*.,;№]+", " ", v)
    v = re.sub(r"\s+", "_", v).strip(" _.-")
    if not v:
        v = "address"
    return v[:max_len].strip(" _.-") or "address"


def unique_path(path: Path) -> Path:
    """Если файл уже есть, добавляет _2, _3 и т.д."""
    if not path.exists():
        return path
    for i in range(2, 1000):
        candidate = path.with_name(f"{path.stem}_{i}{path.suffix}")
        if not candidate.exists():
            return candidate
    return path


def build_auto_output_path(base_output_path: Path, address_prefix: str, start_flat: int, end_flat: int) -> Path:
    house = house_address_from_prefix(address_prefix) or address_prefix or "address"
    slug = filename_slug(house)
    stamp = datetime.now().strftime("%d%m%Y")
    name = f"result_{slug}_{stamp}.xlsx"
    return unique_path(base_output_path.parent / name)


def build_matched_output_path(new_registry_path: Path, fallback_dir: Path | None = None) -> Path:
    """Итог сопоставления: имя от нового реестра, result_ заменяем на matched_."""
    folder = new_registry_path.parent if new_registry_path and new_registry_path.parent else (fallback_dir or Path("output"))
    stem = new_registry_path.stem if new_registry_path else "registry"
    if stem.lower().startswith("result_"):
        out_stem = "matched_" + stem[len("result_"):]
    elif stem.lower().startswith("matched_"):
        out_stem = stem
    else:
        out_stem = "matched_" + stem
    return unique_path(folder / f"{out_stem}.xlsx")


def build_log_output_path(output_path: Path, kind: str = "") -> Path:
    """Отдельный лог-файл Excel рядом с результатом."""
    stem = output_path.stem
    if stem.lower().startswith("result_"):
        log_stem = "log_result_" + stem[len("result_"):]
    elif stem.lower().startswith("matched_"):
        log_stem = "log_matched_" + stem[len("matched_"):]
    elif kind:
        log_stem = f"log_{kind}_{stem}"
    else:
        log_stem = "log_" + stem
    return output_path.with_name(log_stem + output_path.suffix)


RESULT_SHEETS = {"Данные", "Здания и участки", "Кадастровые номера", "Все сведения", "Земельные участки"}


def workbook_copy(wb):
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    from openpyxl import load_workbook
    return load_workbook(bio)


def worksheet_has_data(ws) -> bool:
    """True, если на листе есть данные кроме строки заголовков."""
    if ws.max_row <= 1:
        return False
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(str(v or "").strip() for v in row):
            return True
    return False


def looks_like_datetime_value(value: Any) -> bool:
    s = str(value or "").strip()
    if not s:
        return False
    return bool(
        re.match(r"^\d{2}\.\d{2}\.\d{4}(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?$", s)
        or re.match(r"^\d{4}-\d{2}-\d{2}", s)
    )


SNT_LAND_WIDTHS = {1: 20, 2: 18, 3: 24, 4: 18, 5: 70, 6: 16, 7: 30, 8: 42, 9: 18, 10: 22, 11: 22, 12: 35, 13: 28, 14: 18, 15: 55, 16: 45}

SNT_BUILDING_HEADERS = [
    "Дата/время",
    "Номер участка/дома",
    "Кадастровый номер",
    "Статус объекта",
    "Вид объекта",
    "Адрес",
    "Площадь",
    "Назначение",
    "Количество этажей",
    "Материал наружных стен",
    "Год завершения строительства",
    "Кадастровая стоимость",
    "Дата определения кадастровой стоимости",
    "Дата внесения кадастровой стоимости",
    "Вид права",
    "Номер регистрации",
    "Дата регистрации",
    "Поисковый запрос",
    "Комментарий",
]

SNT_BUILDING_WIDTHS = {1: 18, 2: 18, 3: 24, 4: 18, 5: 22, 6: 70, 7: 16, 8: 28, 9: 18, 10: 28, 11: 24, 12: 18, 13: 24, 14: 24, 15: 35, 16: 28, 17: 18, 18: 55, 19: 50}

SNT_UPDATE_STATUS_HEADER = "Статус актуализации"
SNT_UPDATE_COMMENT_HEADER = "Комментарий актуализации"

SNT_UPDATE_CHANGE_HEADERS = [
    "Дата/время",
    "Лист",
    "Строка",
    "Номер участка",
    "Кадастровый номер",
    "Поле",
    "Было",
    "Стало",
    "Статус",
    "Комментарий",
]

SNT_UPDATE_AMBIGUOUS_HEADERS = [
    "Дата/время",
    "Лист",
    "Строка",
    "Номер участка",
    "Кадастровый номер",
    "Причина",
    "Варианты",
    "Комментарий",
]


def normalize_snt_land_sheet_headers(ws) -> None:
    """
    Исправляет старую ошибку шаблона СНТ:
    раньше в template_snt.xlsx не было колонки «Дата/время», а строки уже писались с датой в первом столбце.
    """
    first_header = normalize_text(ws.cell(1, 1).value)
    if first_header == normalize_text("Номер участка"):
        # Если данные были записаны без даты в первом столбце, вставляем пустую колонку даты.
        # Если в первом столбце уже дата, достаточно только переписать заголовки.
        first_data_value = ws.cell(2, 1).value if ws.max_row >= 2 else ""
        if first_data_value not in (None, "") and not looks_like_datetime_value(first_data_value):
            ws.insert_cols(1)
    for idx, header in enumerate(SNT_LAND_HEADERS, start=1):
        ws.cell(1, idx).value = header
        ws.column_dimensions[ws.cell(1, idx).column_letter].width = SNT_LAND_WIDTHS.get(idx, 18)
    try:
        ws.freeze_panes = "A2"
    except Exception:
        pass


def remove_sheets_except(wb, keep_names: set[str], fallback_name: str = "Лист") -> None:
    # Оставляем только листы из keep_names, где реально есть данные после заголовка.
    existing_keep = [name for name in wb.sheetnames if name in keep_names and worksheet_has_data(wb[name])]
    if not existing_keep:
        # Нельзя удалить все листы, поэтому оставляем fallback с пояснением.
        ws = wb.create_sheet(fallback_name)
        ws.cell(1, 1).value = "Нет данных"
        existing_keep = [ws.title]
    for name in list(wb.sheetnames):
        if name not in existing_keep:
            del wb[name]


def save_result_and_log_workbooks(wb, output_path: Path, *, result_sheets: set[str] | None = None, log_path: Path | None = None, log_kind: str = "") -> Path:
    """
    Сохраняет два файла:
    - основной файл только с результатами;
    - отдельный log_ файл со служебными листами.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    result_sheets = set(result_sheets or RESULT_SHEETS)
    log_path = log_path or build_log_output_path(output_path, log_kind)
    log_path.parent.mkdir(parents=True, exist_ok=True)

    result_wb = workbook_copy(wb)
    remove_sheets_except(result_wb, result_sheets, fallback_name="Данные")
    result_wb.save(output_path)

    log_wb = workbook_copy(wb)
    log_sheets = {name for name in (set(log_wb.sheetnames) - result_sheets) if worksheet_has_data(log_wb[name])}
    if log_sheets:
        remove_sheets_except(log_wb, log_sheets, fallback_name="Лог")
        log_wb.save(log_path)
    else:
        log_path = None
    return log_path


def important_address_tokens(address: str) -> list[str]:
    """Ключевые слова адреса без служебных маркеров. название улицы сравнивается без служебных маркеров."""
    skip_tokens = {"край", "город", "дом", "квартира", "улица", "кв", "ул", "г", "д", "№", "пом", "помещение"}
    tokens: list[str] = []
    # Точки заменяем пробелом, чтобы сокращение с точкой не склеивалось в одно слово.
    text = normalize_text(address).replace(".", " ")
    for token in re.split(r"[\s,]+", text):
        token = token.strip(" .№")
        if len(token) >= 3 and token not in skip_tokens:
            ct = compact_text(token)
            if ct and ct not in tokens:
                tokens.append(ct)
    return tokens


def has_apartment_marker(address: str) -> bool:
    """Отличаем обычные квартиры от помещений без номера/нежилых при поиске по голому дому."""
    return bool(re.search(r"\b(кв|квартира|к\.)\b", normalize_text(address), re.I))


def has_unit_marker(address: str) -> bool:
    """Есть ли в адресе явный маркер квартиры/помещения."""
    return bool(re.search(r"\b(кв\.?|квартира|пом\.?|помещение)\b|(?:^|[\s,])№\s*[0-9]", normalize_text(address), re.I))


def numbered_candidate_priority(row_text: str, expected_flat: str, search_house_no: str = "") -> int:
    """
    Чем меньше число, тем раньше открываем результат.
    0 — явная квартира/помещение с нужным номером.
    1 — безопасный хвост после дома "д 7, 7".
    9 — прочее.
    """
    if extract_flat_no(row_text) == str(expected_flat).strip().upper():
        return 0
    if address_has_trailing_unit(row_text, expected_flat, search_house_no):
        return 1
    return 9


def house_address_matches(
    search_house_address: str,
    result_address: str,
    *,
    strict_street: bool = False,
    strict_house: bool = True,
    allow_unit_without_marker: bool = False,
) -> bool:
    """Совпадение по дому без требования квартиры/помещения."""
    q_house = extract_house_no(search_house_address)
    r_house = extract_house_no(result_address)
    if strict_house and q_house and r_house and q_house != r_house:
        return False
    if strict_street and not street_matches(search_house_address, result_address, strict=True):
        return False
    r_compact = compact_text(result_address)
    for token in important_address_tokens(search_house_address):
        # В мягком режиме токен улицы достаточно совпадает; короткие служебные части всё равно игнорируются.
        if token and token not in r_compact:
            return False
    return True


def extract_cadastral(text: str) -> str:
    m = re.search(r"\b\d{2}:\d{2}:\d{3,}:\d+\b", text or "")
    return m.group(0) if m else ""


def build_search_addresses(address_prefix: str, flat_no: str, *, enabled: bool = True) -> list[str]:
    """
    Формирует варианты поискового адреса для домов, где Росреестр по-разному пишет помещение:
    - ..., кв 1 / кв №1
    - ..., квартира 1 / квартира №1
    - ..., пом. 1 / пом. №1
    - ..., помещение 1 / помещение №1
    - ..., №1

    Вариант "без слова" (..., 1) убран из обычного диапазона, потому что он
    часто цепляет голый адрес дома/земли. Для этого есть отдельный этап
    "Поиск помещений без номеров" после диапазона.
    """
    prefix = re.sub(r"\s+", " ", str(address_prefix or "")).strip(" ,")
    flat = str(flat_no).strip()
    if not prefix:
        return [flat]

    # Если варианты выключены, оставляем старое поведение: пользователь сам отвечает за точную строку.
    if not enabled:
        return [f"{prefix} {flat}".strip()]

    house_prefix = house_address_from_prefix(prefix) or prefix
    prefix_has_unit_tail = bool(re.search(r"(?:кв\.?|квартира|пом\.?|помещение|№)\s*$", normalize_text(prefix), re.I))

    variants = []
    if prefix_has_unit_tail:
        variants.extend([
            f"{prefix} {flat}",
            f"{prefix} №{flat}",
            f"{prefix} № {flat}",
        ])

    variants.extend([
        f"{house_prefix}, кв {flat}",
        f"{house_prefix}, кв №{flat}",
        f"{house_prefix}, кв № {flat}",
        f"{house_prefix}, квартира {flat}",
        f"{house_prefix}, квартира №{flat}",
        f"{house_prefix}, квартира № {flat}",
        f"{house_prefix}, пом. {flat}",
        f"{house_prefix}, пом. №{flat}",
        f"{house_prefix}, пом. № {flat}",
        f"{house_prefix}, помещение {flat}",
        f"{house_prefix}, помещение №{flat}",
        f"{house_prefix}, помещение № {flat}",
        f"{house_prefix}, №{flat}",
        f"{house_prefix}, № {flat}",
    ])

    # Убираем дубли, не меняя порядок.
    result: list[str] = []
    seen: set[str] = set()
    for v in variants:
        v = re.sub(r"\s+", " ", v).strip(" ,")
        key = normalize_text(v)
        if v and key not in seen:
            seen.add(key)
            result.append(v)
    return result



def normalize_plot_number(value: Any) -> str:
    """Нормализует номер участка: № 860-А -> 860А, участок 860/1 -> 860/1."""
    s = clean_unit_number(value)
    if not s:
        return ""
    s = re.sub(r"(?<=\d)-(?=[A-ZА-ЯЁ])", "", s.upper())
    return s


def parse_cadastral_quarters(value: Any) -> list[str]:
    raw = str(value or "")
    parts = re.split(r"[\s,;]+", raw)
    result: list[str] = []
    seen: set[str] = set()
    for part in parts:
        m = re.search(r"\b\d{2}:\d{2}:\d{3,}\b", part)
        if not m:
            continue
        q = m.group(0)
        if q not in seen:
            seen.add(q)
            result.append(q)
    return result


def cadastral_quarter_ok(cad: str, quarters: list[str], enabled: bool = True) -> bool:
    if not enabled or not quarters:
        return True
    cad_s = str(cad or "").strip()
    return any(cad_s.startswith(q + ":") for q in quarters)


def extract_snt_plot_number(address_or_row: Any) -> str:
    """
    Извлекает номер участка из адресов СНТ:
    № 860, участок №860, земельный участок № 1309, д. 236, товарищество, 428.
    """
    raw = str(address_or_row or "")
    raw = re.sub(r"\b\d{2}:\d{2}:\d{3,}:\d+\b", " ", raw)
    t = normalize_text(raw).replace("ё", "е")

    pattern_value = r"([0-9]+[а-яa-z]?(?:\s*[-/]\s*[0-9а-яa-z]+)?)"
    patterns = [
        rf"(?:земельный\s+участок|участок|уч\.?)\s*№?\s*{pattern_value}\b",
        rf"(?:уч\s*-\s*к\.?|уч-к\.?|учк)\s*№?\s*{pattern_value}\b",
        rf"(?:д\.?|дом)\s*№?\s*{pattern_value}\b",
        rf"(?:^|[\s,])№\s*{pattern_value}\b",
        rf"[,\\n]\s*{pattern_value}\s*$",
    ]
    for pat in patterns:
        m = re.search(pat, t, re.I)
        if m:
            return normalize_plot_number(m.group(1))
    return ""


def build_snt_search_addresses(city: str, name: str, plot_no: str, try_variants: bool = True) -> list[str]:
    """Варианты поиска участка СНТ. Форма организации используется как вариант запроса, но не как фильтр."""
    city = re.sub(r"\s+", " ", str(city or "")).strip()
    name = re.sub(r"\s+", " ", str(name or "")).strip().strip('"')
    plot = str(plot_no or "").strip()
    base = f"{city} {name}".strip() or name or city

    # Быстрый режим: один основной запрос. Нужен, чтобы не зависать на десятках вариантов,
    # если конкретный участок не найден.
    if not try_variants:
        return [re.sub(r"\s+", " ", f"{base} {plot}").strip()]

    forms = [
        f"{base} {plot}",
        f"{base} №{plot}",
        f"{base} № {plot}",
        f"{base} участок {plot}",
        f"{base} участок №{plot}",
        f"{base} участок № {plot}",
        f"{base} земельный участок {plot}",
        f"{base} земельный участок №{plot}",
        f"{base} земельный участок № {plot}",
        f"{base} уч {plot}",
        f"{base} уч. №{plot}",
        f"{base} д {plot}",
        f"{base} д. {plot}",
        f"{base} дом {plot}",
    ]
    for org in ["СНТ", "СТ", "ДНТ", "ДСК"]:
        forms.extend([
            f"{city} {org} {name} {plot}".strip(),
            f"{city} {org} {name} участок {plot}".strip(),
            f"{city} {org} {name} №{plot}".strip(),
        ])
    result: list[str] = []
    seen: set[str] = set()
    for v in forms:
        v = re.sub(r"\s+", " ", v).strip()
        key = normalize_text(v)
        if v and key not in seen:
            seen.add(key)
            result.append(v)
    return result


def snt_candidate_match_reason(row_text: str, cad: str, *, city: str, name: str, plot_no: str, quarters: list[str]) -> tuple[bool, str]:
    row_compact = compact_text(row_text)
    reasons: list[str] = []

    city_key = compact_text(city)
    if city_key and city_key not in row_compact:
        reasons.append("город не совпал")

    name_key = compact_text(name)
    if name_key and name_key not in row_compact:
        reasons.append("название не найдено")

    if quarters and not cadastral_quarter_ok(cad, quarters, enabled=True):
        reasons.append("кадастровый квартал не совпал")

    expected_plot = normalize_plot_number(plot_no)
    found_plot = extract_snt_plot_number(row_text)
    if expected_plot:
        if not found_plot:
            reasons.append("номер участка не найден в адресе")
        elif found_plot != expected_plot:
            reasons.append(f"номер участка не совпал: найден {found_plot}")

    if reasons:
        return False, "; ".join(reasons)
    return True, "город/название/квартал/номер участка совпали"


def is_land_object(object_type: str) -> bool:
    t = normalize_text(object_type)
    return "земель" in t or "участ" in t


def card_field(text: str, label: str, extra: str | None = None) -> str:
    return find_line_value(split_lines(text), label, extra_label_part=extra)


def clean_unit_number(value: Any) -> str:
    """
    Очищает номер помещения: 'кв № 9А' -> '9А', 'пом. 12/1' -> '12/1'.
    Случайные слова без цифр не считаем номером.
    """
    s = str(value or "").strip()
    if not s:
        return ""
    s = s.replace("№", " № ")
    s = re.sub(r"(?i)^\s*(?:кв\.?|квартира|пом\.?|помещение|офис|каб\.?|комната|№)\s*", "", s)
    s = re.sub(r"\s+", "", s)
    s = s.strip(" ,.;:")
    if not s:
        return ""
    # Номер должен содержать цифру. Это защищает от случайных слов из файлов.
    if not re.search(r"\d", s):
        return ""
    # Разрешаем типичные форматы: 1, 3А, 12/1, 2-Н, 15Б.
    if not re.fullmatch(r"[0-9A-Za-zА-Яа-яЁё/_\\-]+", s):
        return ""
    return s.upper()


def parse_extra_unit_numbers(value: Any) -> list[str]:
    """
    Разбирает номера помещений из свободного текста.

    Поддерживает формы:
    1 2 3 4 5 15 29
    1, 2, 3, 4, 5, 15, 29
    1; 2; 3; 4; 5; 15; 29
    1\\n2\\n3
    кв 1, кв 2, квартира 3А, пом. 4
    """
    raw = str(value or "")
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    # Вставляем пробел после маркеров, если пользователь написал слитно: кв1, пом3А, №5.
    raw = re.sub(r"(?i)\b(кв|квартира|пом|помещение|офис|каб|комната)\.?\s*№?\s*([0-9])", r"\1 \2", raw)
    raw = re.sub(r"№\s*([0-9])", r" \1", raw)

    # Разделители: запятая, точка с запятой, новая строка, табуляция, обычный пробел.
    parts = re.split(r"[\s,;]+", raw)

    result: list[str] = []
    seen: set[str] = set()
    for part in parts:
        cleaned = clean_unit_number(part)
        if not cleaned:
            continue
        key = cleaned.upper()
        if key not in seen:
            seen.add(key)
            result.append(cleaned)
    return result


def read_manual_units_from_file(path_value: Any) -> list[str]:
    """
    Читает ручную очередь из файла:
    - txt/csv — весь текст;
    - xlsx/xlsm — все видимые ячейки всех листов;
    - docx — таблицы и абзацы.
    Структура не важна: строка, столбик, через пробел/запятую/точку с запятой.
    """
    path_str = str(path_value or "").strip().strip('"')
    if not path_str:
        return []
    path = Path(path_str)
    if not path.exists():
        return []

    ext = path.suffix.lower()
    chunks: list[str] = []

    if ext in [".txt", ".csv"]:
        try:
            chunks.append(path.read_text(encoding="utf-8"))
        except UnicodeDecodeError:
            chunks.append(path.read_text(encoding="cp1251", errors="ignore"))
    elif ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True, read_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for value in row:
                    if value is not None:
                        chunks.append(str(value))
        try:
            wb.close()
        except Exception:
            pass
    elif ext == ".docx":
        try:
            from docx import Document
        except Exception as e:
            raise RuntimeError("Для чтения .docx нужен пакет python-docx. Запустите 00_SETUP.bat после обновления.") from e
        doc = Document(str(path))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        chunks.append(cell.text)
        for p in doc.paragraphs:
            if p.text.strip():
                chunks.append(p.text)
    else:
        raise RuntimeError(f"Формат файла ручной очереди не поддержан: {ext}. Поддерживаются .txt/.csv/.xlsx/.xlsm/.docx")

    return parse_extra_unit_numbers("\n".join(chunks))


def merge_unit_lists(*lists: list[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for items in lists:
        for item in items:
            cleaned = clean_unit_number(item)
            if not cleaned:
                continue
            key = cleaned.upper()
            if key not in seen:
                seen.add(key)
                result.append(cleaned)
    return result


def unit_sort_key(unit: str) -> tuple[int, int, str]:
    """
    Сортировка дополнительных номеров:
    9А идёт после 9, 15Б после 15, Н1 — в конец.
    """
    u = clean_unit_number(unit)
    m = re.match(r"^(\d+)(.*)$", u, flags=re.I)
    if m:
        num = int(m.group(1))
        suffix = m.group(2) or ""
        return (0, num, suffix)
    return (1, 10**9, u)


def build_unit_queue(start_flat: int, end_flat: int, extra_numbers: Any = "") -> list[str]:
    """
    Строит очередь обработки:
    диапазон 1..10 + дополнительные 9А, 9Б => 1,2,...,9,9А,9Б,10.
    Дубли убираются без изменения смысла номеров.
    """
    queue: list[str] = []
    seen: set[str] = set()

    def add(value: Any) -> None:
        cleaned = clean_unit_number(value)
        if not cleaned:
            return
        key = cleaned.upper()
        if key not in seen:
            seen.add(key)
            queue.append(cleaned)

    extras = parse_extra_unit_numbers(extra_numbers)
    by_base: dict[int, list[str]] = {}
    tail: list[str] = []
    for extra in extras:
        m = re.match(r"^(\d+)(.*)$", extra, flags=re.I)
        if m:
            base = int(m.group(1))
            if start_flat <= base <= end_flat:
                by_base.setdefault(base, []).append(extra)
            else:
                tail.append(extra)
        else:
            tail.append(extra)

    if end_flat >= start_flat:
        for num in range(start_flat, end_flat + 1):
            add(str(num))
            for extra in sorted(by_base.get(num, []), key=unit_sort_key):
                # Не добавляем "9" второй раз, если его указали в дополнительных.
                if clean_unit_number(extra) != str(num):
                    add(extra)

    for extra in sorted(tail, key=unit_sort_key):
        add(extra)

    return queue


def build_full_unit_queue(
    start_flat: int,
    end_flat: int,
    extra_numbers: Any = "",
    manual_numbers: Any = "",
    manual_file_path: Any = "",
    *,
    include_range: bool = True,
) -> list[str]:
    """
    Итоговая очередь:
    - если include_range=True: диапазон + доп. номера;
    - затем ручная очередь из поля и файла, без дублей.
    Если диапазон выключен, можно использовать только ручную очередь.
    """
    queue = build_unit_queue(start_flat, end_flat, extra_numbers) if include_range else []
    manual_from_text = parse_extra_unit_numbers(manual_numbers)
    manual_from_file = read_manual_units_from_file(manual_file_path)
    return merge_unit_lists(queue, manual_from_text, manual_from_file)


def address_matches(
    search_address: str,
    result_address: str,
    expected_flat: str | None = None,
    *,
    strict_street: bool = False,
    strict_house: bool = True,
    allow_unit_without_marker: bool = False,
) -> bool:
    """
    Точное практическое совпадение: номер помещения должен совпасть.
    Дом и улица проверяются по настройкам строгости.
    """
    q_flat = extract_flat_no(search_address)
    r_flat = extract_flat_no(result_address)
    expected_flat_norm = str(expected_flat or q_flat or "").strip().upper()
    q_house = extract_house_no(search_address)
    r_house = extract_house_no(result_address)

    # Если ищем конкретную квартиру/помещение, результат без номера помещения или с другим номером не подходит.
    # Допускаем только безопасный редкий формат Росреестра без маркера: "..., д 66, 1".
    # Голый адрес дома/участка вроде "..., ул Улица, 7" больше не считается квартирой 7.
    if expected_flat_norm:
        if not r_flat:
            if not allow_unit_without_marker:
                return False
            if not address_has_trailing_unit(result_address, expected_flat_norm, q_house):
                return False
        elif r_flat != expected_flat_norm:
            return False
    elif q_flat:
        if not r_flat or q_flat != r_flat:
            return False

    if strict_house and q_house and r_house and q_house != r_house:
        return False

    if strict_street and not street_matches(search_address, result_address, strict=True):
        return False

    r_compact = compact_text(result_address)

    # Проверяем ключевые части адреса. Региональный префикс в Росреестре допускается.
    for token in important_address_tokens(search_address):
        if token and token not in r_compact:
            return False

    return True


def clean_reg_number(value: str) -> str:
    value = value or ""
    value = re.sub(r"(?i)\b(собственность|общая|долевая|совместная|право|вид|номер|дата|государственной|регистрации)\b", " ", value)
    value = value.replace("№", " ")
    value = re.sub(r"(?i)\bот\b.*$", "", value)
    value = re.sub(r"\s+", "", value)
    # Точку внутри/в конце регистрационного номера сохраняем: 26-01/...-192.3 и .2 — разные права.
    return value.strip(" ,:;№") or "."


def clean_reg_date(value: str) -> str:
    m = re.search(r"\b\d{2}\.\d{2}\.\d{4}\b", value or "")
    return m.group(0) if m else "."


def split_lines(text: str) -> list[str]:
    return [re.sub(r"\s+", " ", line).strip() for line in (text or "").splitlines() if line.strip()]


def find_line_value(lines: list[str], label: str, *, extra_label_part: str | None = None) -> str:
    """Ищет значение после строки-метки. Учитывает метки, разбитые на две строки."""
    label_n = normalize_text(label)
    extra_n = normalize_text(extra_label_part) if extra_label_part else None

    for i, line in enumerate(lines):
        ln = normalize_text(line)

        # Не ищем короткие метки как подстроки: "Этаж" не должен совпасть с "Количество этажей".
        matched_label = (
            ln == label_n
            or ln.startswith(label_n + ":")
            or (extra_n is not None and label_n in ln)
            or (label_n == "кадастровая стоимость" and ln.startswith(label_n))
        )
        if matched_label:
            j = i
            if extra_n:
                # Например: "Вид объекта" + следующая строка "недвижимости".
                for k in range(i + 1, min(i + 3, len(lines))):
                    if extra_n in normalize_text(lines[k]):
                        j = k
                        break
            # Если значение стоит в этой же строке после метки.
            tail = re.sub(re.escape(label), "", line, flags=re.I).strip(" :—-\t")
            if tail and tail != line and not extra_n and not tail.startswith("("):
                return tail
            # Иначе берём ближайшую следующую строку, которая не похожа на заголовок раздела.
            for k in range(j + 1, min(j + 5, len(lines))):
                candidate = lines[k].strip()
                c_norm = normalize_text(candidate)
                if not candidate:
                    continue
                if c_norm in {"общая информация", "характеристики объекта", "сведения о кадастровой стоимости"}:
                    continue
                if c_norm.startswith(("вид объекта", "статус объекта", "кадастровый номер", "дата присвоения")):
                    continue
                return candidate
    return ""


def parse_rights_from_text(text: str) -> list[RightInfo]:
    lines = split_lines(text)
    rights: list[RightInfo] = []

    # Берём только строки после блока прав, если блок найден. Если сайт меняет порядок — fallback по всему тексту.
    start_idx = 0
    for i, line in enumerate(lines):
        ln = normalize_text(line)
        if "вид, номер" in ln and "регистрации права" in ln:
            start_idx = i + 1
            break

    scan_lines = lines[start_idx:]
    for i, line in enumerate(scan_lines):
        ln = normalize_text(line)
        if "сведения о" in ln and "прав" not in ln and i > 0:
            break
        if "вид, номер" in ln:
            continue
        if "собствен" not in ln:
            continue

        # Каждая строка "Общая долевая собственность" / "Общая совместная собственность"
        # считается отдельной записью права. Не удаляем дубли: Росреестр часто показывает один
        # и тот же номер регистрации несколько раз — это разные собственники/доли.
        combined = " ".join(scan_lines[i : min(i + 5, len(scan_lines))])
        m = re.search(r"№\s*([^\s]+)\s*от\s*(\d{2}\.\d{2}\.\d{4})", combined, flags=re.I)
        if not m:
            # Иногда символ № может отсутствовать в innerText. Важно не терять суффиксы .2/.3.
            m = re.search(r"\b(\d{2}(?:[:\-]\d{1,2}){1,3}[0-9А-Яа-яA-Za-z:/\-.]*(?:[-/]\d+(?:\.\d+)?)+)\s+от\s*(\d{2}\.\d{2}\.\d{4})", combined, flags=re.I)
        if not m:
            continue

        reg_number = clean_reg_number(m.group(1))
        reg_date = clean_reg_date(m.group(2))

        if "совмест" in ln:
            right_type = "Общая совместная собственность" if "общ" in ln else "Совместная собственность"
        elif "общ" in ln and "долев" in ln:
            right_type = "Общая долевая собственность"
        else:
            right_type = "Собственность"

        rights.append(RightInfo(right_type=right_type, reg_number=reg_number, reg_date=reg_date))

    return rights

def detect_card_object_type(text: str, lines: list[str]) -> str:
    object_type = find_line_value(lines, "Вид объекта", extra_label_part="недвижимости")
    t = normalize_text(text)

    def explicit_type() -> str:
        if "машино-место" in t or "машино место" in t:
            return "Машино-место"
        if "земельный участок" in t:
            return "Земельный участок"
        if re.search(r"(^|\n|\s)здание(\s|\n|$)", t, re.I):
            return "Здание"
        if "помещение" in t:
            return "Помещение"
        return ""

    if not object_type:
        return explicit_type()

    # Иногда find_line_value цепляет служебную подпись, проверяем по явным словам карточки.
    ot = normalize_text(object_type)
    known = any(x in ot for x in ["земель", "участ", "здание", "помещ", "машино"])
    if not known:
        return explicit_type() or object_type

    if "машино" in ot:
        return "Машино-место"
    if "земель" in ot or "участ" in ot:
        return "Земельный участок"
    if "здание" in ot:
        return "Здание"
    if "помещ" in ot:
        return "Помещение"
    return object_type


FULL_INFO_SECTIONS = {
    normalize_text("Общая информация"),
    normalize_text("Характеристики объекта"),
    normalize_text("Сведения о кадастровой стоимости"),
    normalize_text("Ранее присвоенные номера"),
    normalize_text("Сведения о правах и ограничениях (обременениях)"),
}

FULL_INFO_LABELS = [
    "Дата обновления информации",
    "Вид объекта недвижимости",
    "Статус объекта",
    "Кадастровый номер",
    "Дата присвоения кадастрового номера",
    "Форма собственности",
    "Адрес (местоположение)",
    "Адрес",
    "Площадь, кв.м",
    "Площадь, кв. м",
    "Площадь",
    "Назначение",
    "Этаж",
    "Кадастровая стоимость (руб)",
    "Кадастровая стоимость",
    "Дата определения кадастровой стоимости",
    "Дата внесения кадастровой стоимости",
    "Дата определения",
    "Дата внесения",
    "Инвентарный номер",
    "Условный номер",
    "Категория земель",
    "Вид разрешенного использования",
    "Количество этажей",
    "Количество подземных этажей",
    "Материал наружных стен",
    "Год завершения строительства",
    "Вид, номер и дата государственной регистрации права",
    "Ограничение прав и обременение объекта недвижимости",
]


def full_label_matches(lines: list[str], index: int, label: str) -> tuple[bool, int]:
    label_n = normalize_text(label)
    line_n = normalize_text(lines[index] if index < len(lines) else "")
    next_n = normalize_text(lines[index + 1] if index + 1 < len(lines) else "")
    combined_n = normalize_text((lines[index] if index < len(lines) else "") + " " + (lines[index + 1] if index + 1 < len(lines) else ""))

    if line_n == label_n:
        return True, 1
    if combined_n == label_n:
        return True, 2

    # Некоторые подписи в DOM могут быть "Кадастровая стоимость (руб)" или "Дата обновления информации: 29.12.2023".
    if line_n.startswith(label_n + ":"):
        return True, 1
    return False, 0


def is_full_section_line(line: str) -> bool:
    return normalize_text(line) in FULL_INFO_SECTIONS


def is_any_full_label_start(lines: list[str], index: int) -> bool:
    if index >= len(lines):
        return False
    for label in FULL_INFO_LABELS:
        ok, _consumed = full_label_matches(lines, index, label)
        if ok:
            return True
    return False


def add_full_field(
    result: list[tuple[str, str, str, int]],
    seen: set[tuple[str, str, str]],
    section: str,
    field: str,
    value: Any,
    order: int,
) -> int:
    value_s = str(value or "").strip()
    if not value_s or value_s == ".":
        return order
    key = (section, field, value_s)
    if key not in seen:
        result.append((section, field, value_s, order))
        seen.add(key)
        order += 1
    return order


def ensure_structured_full_fields(text: str, result: list[tuple[str, str, str, int]], seen: set[tuple[str, str, str]], order: int) -> int:
    """
    Добавляет важные поля из структурированного разбора, если общий full-разбор их не поймал.
    Особенно важно для кадастровой стоимости и объектов разных типов.
    """
    lines = split_lines(text)
    info = parse_card_text(text, fallback_flat=".", fallback_cadastral="")
    section_common = "Общая информация"
    section_char = "Характеристики объекта"
    section_cost = "Сведения о кадастровой стоимости"

    order = add_full_field(result, seen, section_common, "Вид объекта недвижимости", info.object_type, order)
    order = add_full_field(result, seen, section_common, "Статус объекта", info.status, order)
    order = add_full_field(result, seen, section_common, "Кадастровый номер", info.cadastral_number, order)
    order = add_full_field(result, seen, section_common, "Форма собственности", info.ownership_form, order)

    cad_assign_date = find_line_value(lines, "Дата присвоения", extra_label_part="кадастрового номера")
    order = add_full_field(result, seen, section_common, "Дата присвоения кадастрового номера", cad_assign_date, order)

    order = add_full_field(result, seen, section_char, "Адрес (местоположение)", info.address, order)
    order = add_full_field(result, seen, section_char, "Площадь, кв.м", info.area, order)
    order = add_full_field(result, seen, section_char, "Назначение / ВРИ / Категория", info.purpose, order)
    order = add_full_field(result, seen, section_char, "Этаж", info.floor, order)

    for label, extra in [
        ("Категория земель", None),
        ("Вид разрешенного", "использования"),
        ("Количество этажей", None),
        ("Количество подземных", "этажей"),
        ("Материал наружных стен", None),
        ("Год завершения", "строительства"),
    ]:
        value = find_line_value(lines, label, extra_label_part=extra)
        nice_label = label if extra is None else f"{label} {extra}"
        order = add_full_field(result, seen, section_char, nice_label, value, order)

    # Кадастровая стоимость — важное поле для полного парсинга.
    cost = info.cadastral_cost or find_line_value(lines, "Кадастровая стоимость", extra_label_part="руб")
    order = add_full_field(result, seen, section_cost, "Кадастровая стоимость (руб)", cost, order)
    order = add_full_field(result, seen, section_cost, "Дата определения", find_line_value(lines, "Дата определения"), order)
    order = add_full_field(result, seen, section_cost, "Дата внесения", find_line_value(lines, "Дата внесения"), order)

    for label in ["Инвентарный номер", "Условный номер"]:
        order = add_full_field(result, seen, "Ранее присвоенные номера", label, find_line_value(lines, label), order)

    return order


def extract_full_info_fields(text: str) -> list[tuple[str, str, str, int]]:
    """
    Возвращает строки полного парсинга: раздел, поле, значение, порядок.
    Это не заменяет основной разбор, а сохраняет максимум видимых сведений карточки.
    """
    lines = split_lines(text)
    result: list[tuple[str, str, str, int]] = []
    seen: set[tuple[str, str, str]] = set()
    section = ""
    i = 0
    order = 1

    while i < len(lines):
        line = lines[i]

        if is_full_section_line(line):
            section = line
            i += 1
            continue

        # Строки вида "Дата обновления информации: 29.12.2023".
        m = re.match(r"^\s*(Дата\s+обновления\s+информации)\s*:\s*(.+?)\s*$", line, flags=re.I)
        if m:
            key = (section, m.group(1), m.group(2))
            if key not in seen:
                result.append((section, m.group(1), m.group(2), order))
                seen.add(key); order += 1
            i += 1
            continue

        matched = False
        # Длинные подписи проверяем первыми.
        for label in sorted(FULL_INFO_LABELS, key=len, reverse=True):
            ok, consumed = full_label_matches(lines, i, label)
            if not ok:
                continue

            value_parts: list[str] = []
            # Если значение в этой же строке после двоеточия.
            if ":" in line and consumed == 1:
                after = line.split(":", 1)[1].strip()
                if after:
                    value_parts.append(after)

            j = i + consumed
            while j < len(lines):
                if is_full_section_line(lines[j]) or is_any_full_label_start(lines, j):
                    break
                value_parts.append(lines[j])
                j += 1

            value = " | ".join(p for p in value_parts if p).strip()
            if value:
                key = (section, label, value)
                if key not in seen:
                    result.append((section, label, value, order))
                    seen.add(key); order += 1
            i = max(j, i + consumed)
            matched = True
            break

        if matched:
            continue
        i += 1

    # Добавляем контрольные структурированные поля, если общий разбор их не поймал.
    order = ensure_structured_full_fields(text, result, seen, order)

    # Отдельно добавляем структурированные права, чтобы их было удобно фильтровать.
    for idx, right in enumerate(parse_rights_from_text(text), start=1):
        rows = [
            ("Сведения о правах и ограничениях (обременениях)", f"Право {idx}: вид", right.right_type, order),
            ("Сведения о правах и ограничениях (обременениях)", f"Право {idx}: номер регистрации", right.reg_number, order + 1),
            ("Сведения о правах и ограничениях (обременениях)", f"Право {idx}: дата регистрации", right.reg_date, order + 2),
        ]
        for section_name, field, value, ord_no in rows:
            if value and value != ".":
                key = (section_name, field, value)
                if key not in seen:
                    result.append((section_name, field, value, ord_no))
                    seen.add(key)
        order += 3

    return result


def parse_card_text(text: str, fallback_flat: str, fallback_cadastral: str = "") -> ObjectInfo:
    lines = split_lines(text)
    object_type = detect_card_object_type(text, lines)
    status = find_line_value(lines, "Статус объекта")

    # На странице в DOM иногда остаётся поисковая форма с подписью
    # "Адрес или кадастровый номер". Старый разбор мог принять текст
    # "Адрес или" за кадастровый номер. Поэтому кадастровый номер
    # считаем валидным только если он реально похож на 26:12:030108:985.
    raw_cadastral = find_line_value(lines, "Кадастровый номер")
    cadastral_number = raw_cadastral if re.fullmatch(r"\d{2}:\d{2}:\d{3,}:\d+", raw_cadastral or "") else ""
    if not cadastral_number and fallback_cadastral:
        cadastral_number = fallback_cadastral
    if not cadastral_number:
        cadastral_number = extract_cadastral(text)

    ownership_form = find_line_value(lines, "Форма собственности")
    address = find_line_value(lines, "Адрес (местоположение)") or find_line_value(lines, "Адрес")
    area = find_line_value(lines, "Площадь, кв.м") or find_line_value(lines, "Площадь, кв. м")
    purpose = (
        find_line_value(lines, "Назначение")
        or find_line_value(lines, "Вид разрешенного", extra_label_part="использования")
        or find_line_value(lines, "Категория земель")
    )
    floor = find_line_value(lines, "Этаж")
    cadastral_cost = find_line_value(lines, "Кадастровая стоимость", extra_label_part="руб")
    rights = parse_rights_from_text(text)

    flat_from_address = extract_flat_no(address)
    flat_no = flat_from_address or str(fallback_flat)

    return ObjectInfo(
        flat_no=flat_no,
        area=area or ".",
        cadastral_number=cadastral_number or ".",
        ownership_form=ownership_form or "",
        address=address or "",
        object_type=object_type or "",
        status=status or "",
        purpose=purpose or "",
        floor=floor or "",
        cadastral_cost=cadastral_cost or "",
        rights=rights,
        raw_text=text or "",
    )


def municipal_value(ownership_form: str) -> str:
    form = normalize_text(ownership_form)
    if "муницип" in form or "государ" in form:
        return "да"
    return ""


def is_building_or_land(object_type: str) -> bool:
    """Здание и земельный участок не пишем в основной лист собственников."""
    t = normalize_text(object_type)
    return "земель" in t or "участ" in t or "здание" in t


def aux_purpose_value(info: ObjectInfo) -> str:
    """Для вспомогательного листа: назначение здания или ВРИ/категория участка, если удалось распарсить."""
    return info.purpose or ""


def share_for_right(right: RightInfo, all_rights: list[RightInfo]) -> tuple[str, str]:
    """Возвращает (признак совместной, доля)."""
    rt = normalize_text(right.right_type)
    if "совмест" in rt:
        return "да", "1/1"
    if "общ" in rt and "долев" in rt:
        dolevaya_count = sum(1 for r in all_rights if "общ" in normalize_text(r.right_type) and "долев" in normalize_text(r.right_type))
        return "", f"1/{max(dolevaya_count, 1)}"
    return "", "1/1"



# ----------------------------- registry reconciliation helpers -----------------------------


def split_multiline_cell(value: Any) -> list[str]:
    raw = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    raw = re.sub(r"\s+\n", "\n", raw)
    parts = []
    for part in re.split(r"\n+|;+", raw):
        part = re.sub(r"\s+", " ", part).strip(" ,;")
        if part:
            parts.append(part)
    return parts


def split_fio_cell(value: Any) -> list[str]:
    """
    ФИО в старых реестрах бывают:
    - по строкам;
    - через точку с запятой;
    - через запятую в одной ячейке.
    Если после разделения по запятой каждая часть похожа на ФИО/организацию,
    считаем это отдельными правообладателями.
    """
    base_parts = split_multiline_cell(value)
    result: list[str] = []
    for part in base_parts:
        comma_parts = [re.sub(r"\s+", " ", p).strip(" ,;") for p in part.split(",") if p.strip()]
        if len(comma_parts) > 1 and all(looks_like_fio(p) for p in comma_parts):
            result.extend(comma_parts)
        else:
            result.append(part)
    return result


def clean_reg_date_any(value: Any) -> str:
    """Дата из текста или отдельной Excel-ячейки: 29.01.2002 / 2002-01-29 / datetime."""
    if value is None:
        return ""
    try:
        if hasattr(value, "strftime"):
            return value.strftime("%d.%m.%Y")
    except Exception:
        pass
    s = str(value or "").strip()
    if not s:
        return ""
    m = re.search(r"\b\d{2}\.\d{2}\.\d{4}\b", s)
    if m:
        return m.group(0)
    m = re.search(r"\b(\d{4})-(\d{2})-(\d{2})\b", s)
    if m:
        return f"{m.group(3)}.{m.group(2)}.{m.group(1)}"
    return ""


def looks_like_fio(value: str) -> bool:
    s = str(value or "").strip()
    if not s or s == ".":
        return False
    n = normalize_text(s)
    if any(x in n for x in ["собственность", "договор", "регистрац", "кадастр", "площад", "голос", "квартира", "помещ"]):
        return False
    if re.search(r"\b(ооо|оао|зао|пао|ао|ип|администрац|муницип|российская федерация)\b", n):
        return True
    words = [w for w in re.split(r"\s+", s) if w]
    cap_words = [w for w in words if re.match(r"^[А-ЯЁA-Z][а-яёa-zA-Z-]+$", w)]
    return len(cap_words) >= 2


def split_fio_or_org(value: str) -> tuple[str, str, str]:
    s = re.sub(r"\s+", " ", str(value or "").strip())
    if not s or s == ".":
        return ".", ".", "."
    n = normalize_text(s)
    org_markers = ["ооо", "оао", "зао", "пао", " ао ", "ип ", "администрац", "муницип", "российская федерация", "комитет", "управление"]
    if any(m in f" {n} " for m in org_markers):
        return s, ".", "."
    parts = s.split()
    # Иногда старый реестр склеивает фамилию и имя: ФамилияИмя Отчество.
    # Делим только по внутренней заглавной русской букве, чтобы не ломать обычные фамилии.
    if parts:
        m_joined = re.match(r"^([А-ЯЁ][а-яё-]{2,})([А-ЯЁ][а-яё-]{2,})$", parts[0])
        if m_joined:
            parts = [m_joined.group(1), m_joined.group(2)] + parts[1:]
    if len(parts) >= 3:
        return parts[0], parts[1], " ".join(parts[2:])
    if len(parts) == 2:
        return parts[0], parts[1], "."
    return s, ".", "."


def normalize_reg_key(value: Any) -> str:
    s = str(value or "").strip()
    if not s or s == ".":
        return ""
    cleaned = clean_reg_number(s)
    if cleaned == ".":
        return ""
    return compact_text(cleaned).upper()


def normalize_contract_key(value: Any) -> str:
    s = normalize_text(value)
    if not s:
        return ""
    if not any(w in s for w in ["договор", "приватизац", "дарени", "мена", "обмен", "купли", "продажи", "свидетельство"]):
        return ""
    date = clean_reg_date(s)
    m = re.search(r"№\s*([^,\n]+?)(?=\s+от\b|,|$)", str(value or ""), re.I)
    num = re.sub(r"\s+", "", m.group(1)).strip(" .,:;№") if m else ""
    kind = "договор"
    for k in ["приватизац", "дарени", "мена", "обмен", "купли", "продажи", "свидетельство"]:
        if k in s:
            kind = k
            break
    key = f"CONTRACT:{kind}:{num}:{date if date != '.' else ''}"
    if not num and date == ".":
        key = "CONTRACT:" + compact_text(s)[:80].upper()
    return key.upper()


def extract_share(value: Any) -> str:
    s = str(value or "")
    # Долю берём только если рядом есть слово "доля", чтобы не принять часть регномера 26-01/12-... за долю.
    m = re.search(r"(?i)дол[яи]\D{0,40}(\d+\s*/\s*\d+)", s)
    if m:
        return re.sub(r"\s+", "", m.group(1))
    return ""


def extract_share_any(value: Any) -> str:
    """Доля из отдельной колонки: принимает и чистое '1/4', и текст с 'доля 1/4'."""
    s = str(value or "")
    m = re.search(r"(\d+\s*/\s*\d+)", s)
    if m:
        return re.sub(r"\s+", "", m.group(1))
    return ""


def divide_share(share: str, count: int) -> str:
    """Делит общую долю из одной строки старого реестра между несколькими ФИО."""
    if not share or count <= 1:
        return share
    try:
        from fractions import Fraction
        m = re.fullmatch(r"(\d+)\s*/\s*(\d+)", str(share).strip())
        if not m:
            return share
        frac = Fraction(int(m.group(1)), int(m.group(2))) / count
        return f"{frac.numerator}/{frac.denominator}"
    except Exception:
        return share


def parse_right_line(value: Any) -> tuple[str, str, str, bool, str, str]:
    """
    Возвращает: (right_key, reg_number, reg_date, is_contract, contract_text, share)
    """
    raw = str(value or "").strip()
    if not raw:
        return "", "", "", False, "", ""
    n = normalize_text(raw)
    date = clean_reg_date(raw)
    if date == ".":
        m_iso = re.search(r"\b(\d{4})-(\d{2})-(\d{2})\b", raw)
        if m_iso:
            date = f"{m_iso.group(3)}.{m_iso.group(2)}.{m_iso.group(1)}"
    share = extract_share(raw)

    is_contract = any(w in n for w in ["договор", "приватизац", "дарени", "мена", "обмен", "купли", "продажи", "свидетельство"])
    if is_contract and not re.search(r"\b\d{2}[:\-]\d{2}[:/\-]", raw):
        key = normalize_contract_key(raw)
        return key, "", "" if date == "." else date, True, raw, share

    reg = ""
    m = re.search(r"№\s*([^,\n]+?)(?=\s+от\b|,|доля|$)", raw, flags=re.I)
    if m:
        reg = m.group(1).strip(" .,:;№")
    if not reg:
        # Регномер без №:
        # 26-26-12/059/2008-526
        # 26:12:030101:1363-26/001/2017-2
        # 26-0-1-117/2003/2013-71
        m = re.search(r"\b\d{2}(?:[:\-]\d{1,2}){1,3}[0-9:/\-.]*(?:[-/]\d+(?:\.\d+)*)+\b", raw)
        if m:
            reg = m.group(0).strip(" .,:;№")
    if reg:
        key = normalize_reg_key(reg)
        return key, reg, "" if date == "." else date, False, "", share

    if is_contract:
        key = normalize_contract_key(raw)
        return key, "", "" if date == "." else date, True, raw, share

    return "", "", "" if date == "." else date, False, "", share


def format_contract_for_output(rec: OldRegistryRecord) -> str:
    """Очищает старый договор для записи в колонку номера права."""
    raw = str(rec.contract_text or rec.raw_right or "").strip()
    if not raw:
        return "."
    raw = re.sub(r"(?i)^\s*(?:общая\s+)?долевая\s+собственность\s*,?\s*", "", raw)
    raw = re.sub(r"(?i)^\s*собственность\s*,?\s*", "", raw)
    raw = re.sub(r"\s+", " ", raw).strip(" ,;")
    # Номер права должен быть без "от дата"; дата пишется в отдельную колонку.
    raw = re.sub(r"(?i)\s+от\s+\d{4}-\d{2}-\d{2}\b", "", raw).strip(" ,;")
    raw = re.sub(r"(?i)\s+от\s+\d{2}\.\d{2}\.\d{4}\b", "", raw).strip(" ,;")
    # Приводим "Договор" к "договор" для единообразия, если это первое слово.
    raw = re.sub(r"(?i)^договор", "договор", raw)
    return raw or "."


def apply_inferred_shares(records: list[OldRegistryRecord]) -> None:
    """
    Если в старом реестре одна и та же долевая запись/договор повторена для нескольких ФИО,
    а доли явно не указаны, проставляем равные доли.
    """
    groups: dict[tuple[str, str], list[OldRegistryRecord]] = {}
    for rec in records:
        if not rec.right_key:
            continue
        key = (rec.flat_no.upper(), rec.right_key)
        groups.setdefault(key, []).append(rec)
    for (_flat, _key), items in groups.items():
        if len(items) <= 1:
            continue
        if any(rec.share for rec in items):
            continue
        joined = normalize_text("\n".join(rec.raw_right or rec.contract_text for rec in items))
        if "долев" in joined or any(rec.is_contract for rec in items):
            inferred = f"1/{len(items)}"
            for rec in items:
                rec.share = inferred


def all_same_nonempty_flat(records: list[OldRegistryRecord], flat: str) -> bool:
    if not records or not flat:
        return False
    f = flat.upper()
    return all((r.flat_no or "").upper() == f for r in records)


def parse_fraction_share(value: Any):
    try:
        from fractions import Fraction
        s = str(value or "").strip()
        m = re.fullmatch(r"(\d+)\s*/\s*(\d+)", s)
        if not m:
            return None
        den = int(m.group(2))
        if den == 0:
            return None
        return Fraction(int(m.group(1)), den)
    except Exception:
        return None


def should_skip_old_share_override(ws, header_row: int, cols: dict[str, int | None], target_row: int, flat: str, old_share: str, *, expanded: bool = False) -> bool:
    """
    Не переносим устаревшую долю 1/1 из старого реестра, если в новом Росреестре
    по этому помещению уже есть несколько действующих прав/долей.

    Типичный случай: раньше в старом реестре у человека была 1/1,
    затем появилась новая регистрация доли. Росреестр уже верно дал 1/2 + 1/2,
    поэтому старую 1/1 нельзя перетирать поверх 1/2.

    При разворачивании нескольких старых собственников правило не применяется:
    там старые доли нужны для корректного разбиения.
    """
    if expanded:
        return False
    share_col = cols.get("share")
    flat_col = cols.get("flat")
    reg_col = cols.get("reg")
    if not share_col or not flat_col or not reg_col or not flat:
        return False

    old_frac = parse_fraction_share(old_share)
    current_frac = parse_fraction_share(ws.cell(target_row, share_col).value)
    if old_frac is None or current_frac is None:
        return False

    try:
        from fractions import Fraction
        if old_frac != Fraction(1, 1):
            return False
        if current_frac == Fraction(1, 1):
            return False

        same_flat_right_rows = 0
        for r in range(header_row + 1, ws.max_row + 1):
            row_flat = clean_unit_number(ws.cell(r, flat_col).value)
            if row_flat.upper() != flat.upper():
                continue
            reg = str(ws.cell(r, reg_col).value or "").strip()
            if reg and reg != ".":
                same_flat_right_rows += 1
        return same_flat_right_rows >= 2
    except Exception:
        return False


def resolve_old_registry_paths(value: Any) -> list[Path]:
    """Список старых реестров: можно указать один файл или несколько через ;."""
    raw = str(value or "").strip()
    if not raw:
        return []
    parts = [p.strip().strip('"') for p in re.split(r";+", raw) if p.strip()]
    return [Path(p) for p in parts]


def dedupe_old_records(records: list[OldRegistryRecord]) -> list[OldRegistryRecord]:
    """
    Убирает только абсолютно одинаковые записи, которые пришли из одного и того же места.

    Важно: одинаковые ФИО с одинаковым правом в разных строках старого реестра НЕ удаляем.
    Иногда это не ошибка файла, а несколько долей одного и того же правообладателя.
    Пример: по квартире 1 четыре доли по 1/4, где один собственник встречается дважды.
    """
    result: list[OldRegistryRecord] = []
    seen: set[tuple[str, str, int, str, str, str, str, str]] = set()
    for r in records:
        key = (
            r.source_file,
            r.sheet,
            int(r.row or 0),
            (r.right_key or "").upper(),
            (r.flat_no or "").upper(),
            normalize_text(r.fio_original),
            normalize_text(r.share),
            normalize_text(r.reg_date),
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(r)
    return result


def is_empty_person_cells(ws, row_idx: int, cols: dict[str, int | None]) -> bool:
    vals = []
    for key in ["last", "first", "patr"]:
        col = cols.get(key)
        vals.append(str(ws.cell(row_idx, col).value if col else "").strip())
    # Если фамилия/имя/отчество точками или пустые — строка считается незаполненной.
    normalized = [v for v in vals if v and v != "."]
    return not normalized


def highlight_reconcile_rows(ws, rows: list[int], *, share_col: int | None = None, reason: str = "") -> None:
    """Подсвечивает спорные строки сопоставления жёлтым цветом."""
    try:
        from openpyxl.styles import PatternFill
        from openpyxl.comments import Comment
    except Exception:
        return
    fill = PatternFill(fill_type="solid", fgColor="FFF2CC")
    max_col = ws.max_column
    for row_idx in rows:
        if not isinstance(row_idx, int) or row_idx <= 0:
            continue
        for c in range(1, max_col + 1):
            ws.cell(row_idx, c).fill = fill
        if share_col:
            cell = ws.cell(row_idx, share_col)
            if reason:
                try:
                    cell.comment = Comment(reason, "RosreestrParser")
                except Exception:
                    pass


def append_share_sum_warnings(ws, header_row: int, cols: dict[str, int | None], ambiguous: list[list[Any]]) -> int:
    """
    Пишет предупреждения, если по помещению сумма долей не равна 1.

    Важно: для общей совместной собственности доли 1/1 у каждого собственника — это нормально.
    Такие помещения не проверяем как обычную долевую собственность.
    """
    flat_col = cols.get("flat")
    share_col = cols.get("share")
    joint_col = cols.get("joint")
    if not flat_col or not share_col:
        return 0
    from fractions import Fraction
    groups: dict[str, list[tuple[int, Fraction | None, str, bool]]] = {}
    for r in range(header_row + 1, ws.max_row + 1):
        flat = clean_unit_number(ws.cell(r, flat_col).value)
        if not flat:
            continue
        raw_share = str(ws.cell(r, share_col).value or "").strip()
        frac = parse_fraction_share(raw_share)
        joint_yes = False
        if joint_col:
            joint_yes = normalize_text(ws.cell(r, joint_col).value) == "да"
        groups.setdefault(flat.upper(), []).append((r, frac, raw_share, joint_yes))
    warning_count = 0
    for flat, items in groups.items():
        if not items:
            continue
        if any(joint_yes for _r, _frac, _raw, joint_yes in items):
            # Совместная собственность: у каждого может быть 1/1, сумму не складываем.
            continue
        if any(frac is None for _r, frac, _raw, _joint in items):
            continue
        total = sum((frac for _r, frac, _raw, _joint in items if frac is not None), Fraction(0, 1))
        if total != Fraction(1, 1):
            rows = [r for r, _f, _raw, _joint in items]
            rows_text = ", ".join(str(r) for r in rows)
            shares = ", ".join(raw for _r, _f, raw, _joint in items)
            reason = f"Сумма долей по помещению = {total}, не равна 1/1"
            ambiguous.append([rows_text, flat, "", "", reason, shares])
            highlight_reconcile_rows(ws, rows, share_col=share_col, reason=reason)
            warning_count += 1
    return warning_count


def header_score(value: Any, keywords: list[str]) -> int:
    n = normalize_text(value)
    return sum(1 for kw in keywords if kw in n)


def find_header_row(ws, max_rows: int = 30) -> tuple[int, dict[str, int]]:
    """Ищет строку заголовков по смыслу колонок."""
    best_row = 1
    best_score = -1
    best_map: dict[str, int] = {}
    for r in range(1, min(max_rows, ws.max_row) + 1):
        row_map: dict[str, int] = {}
        score = 0
        for c in range(1, ws.max_column + 1):
            val = ws.cell(r, c).value
            n = normalize_text(val)
            if not n:
                continue
            if any(k in n for k in ["фио", "правооблад", "собственник", "владелец"]):
                row_map.setdefault("fio", c); score += 3
            if "дата" in n and ("регистрац" in n or "права" in n or "право" in n):
                row_map.setdefault("date", c); score += 2
            if (
                "правооблад" not in n
                and not ("дата" in n and ("регистрац" in n or "права" in n or "право" in n))
                and (
                    n == "права"
                    or "сведения о прав" in n
                    or "правоустанавлива" in n
                    or "регистрац" in n
                    or "договор" in n
                    or ("право" in n and "фио" not in n)
                    or "документ" in n
                )
            ):
                row_map.setdefault("right", c); score += 3
            if any(k in n for k in ["номер помещения", "номер пом", "№ помещения", "помещение", "квартира", "номер кв"]):
                row_map.setdefault("flat", c); score += 2
            if "кадастров" in n:
                row_map.setdefault("cad", c); score += 2
            if any(k in n for k in ["доля", "голос"]):
                row_map.setdefault("share", c); score += 1
        if score > best_score:
            best_score = score
            best_row = r
            best_map = row_map
    return best_row, best_map


def find_new_data_sheet(wb):
    """Ищет лист нового реестра по колонке номера регистрации и ФИО."""
    for ws in wb.worksheets:
        header_row, hmap = find_header_row(ws)
        headers = [normalize_text(ws.cell(header_row, c).value) for c in range(1, ws.max_column + 1)]
        has_reg = any("номер регистрации" in h or ("регистрац" in h and "номер" in h) for h in headers)
        has_fio = any("фамилия" in h or "наименование организации" in h for h in headers)
        has_cad = any("кадастров" in h for h in headers)
        if has_reg and has_fio and has_cad:
            return ws, header_row
    return wb.active, 1


def get_col_by_header(ws, header_row: int, patterns: list[str]) -> int | None:
    for c in range(1, ws.max_column + 1):
        n = normalize_text(ws.cell(header_row, c).value)
        if all(p in n for p in patterns):
            return c
    return None


def get_new_header_columns(ws, header_row: int) -> dict[str, int | None]:
    return {
        "flat": get_col_by_header(ws, header_row, ["номер", "помещ"]) or get_col_by_header(ws, header_row, ["номер", "кварт"]),
        "cad": get_col_by_header(ws, header_row, ["кадастров"]),
        "reg": get_col_by_header(ws, header_row, ["номер", "регистрац"]),
        "reg_date": get_col_by_header(ws, header_row, ["дата", "регистрац"]),
        "share": get_col_by_header(ws, header_row, ["доля"]),
        "joint": get_col_by_header(ws, header_row, ["совмест"]),
        "room_type": get_col_by_header(ws, header_row, ["тип", "помещ"]) or get_col_by_header(ws, header_row, ["вид", "объект"]),
        "right_type": get_col_by_header(ws, header_row, ["вид", "права"]),
        "last": get_col_by_header(ws, header_row, ["фамилия"]),
        "first": get_col_by_header(ws, header_row, ["имя"]),
        "patr": get_col_by_header(ws, header_row, ["отчество"]),
    }


def read_text_docx(path: Path) -> list[list[str]]:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("Для чтения .docx нужен пакет python-docx. Запустите 00_SETUP.bat после обновления.") from e
    doc = Document(str(path))
    rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
    # Если таблиц нет, берём абзацы как одну колонку.
    if not rows:
        for p in doc.paragraphs:
            if p.text.strip():
                rows.append([p.text.strip()])
    return rows


def extract_old_records_from_rows(rows: list[list[Any]], source_file: str, sheet_name: str) -> list[OldRegistryRecord]:
    records: list[OldRegistryRecord] = []
    if not rows:
        return records

    # Превращаем список в псевдотаблицу и ищем заголовок вручную.
    best_idx = 0
    best_score = -1
    hmap: dict[str, int] = {}
    for i, row in enumerate(rows[:30]):
        local: dict[str, int] = {}
        score = 0
        for c, val in enumerate(row):
            n = normalize_text(val)
            if "фамил" in n or "наименование организации" in n:
                local.setdefault("last", c); score += 3
            elif n == "имя" or "имя (" in n:
                local.setdefault("first", c); score += 2
            elif "отчеств" in n:
                local.setdefault("patr", c); score += 2
            elif any(k in n for k in ["фио", "правооблад", "собственник"]):
                local.setdefault("fio", c); score += 3
            if "адрес" in n:
                local.setdefault("address", c); score += 3
            if "дата" in n and ("регистрац" in n or "права" in n or "право" in n):
                local.setdefault("date", c); score += 2
            if (
                "правооблад" not in n
                and not ("дата" in n and ("регистрац" in n or "права" in n or "право" in n))
                and (
                    n == "права"
                    or "сведения о прав" in n
                    or "правоустанавлива" in n
                    or "регистрац" in n
                    or "договор" in n
                    or ("право" in n and "фио" not in n)
                )
            ):
                local.setdefault("right", c); score += 3
            if any(k in n for k in ["номер помещения", "номер пом", "помещение", "квартира"]):
                local.setdefault("flat", c); score += 2
            if "кадастров" in n:
                local.setdefault("cad", c); score += 2
            if any(k in n for k in ["доля", "голос"]):
                local.setdefault("share", c); score += 1
        if score > best_score:
            best_score = score; best_idx = i; hmap = local

    for ri, row in enumerate(rows[best_idx + 1:], start=best_idx + 2):
        def cell(key: str) -> str:
            idx = hmap.get(key)
            return str(row[idx]).strip() if idx is not None and idx < len(row) and row[idx] is not None else ""

        address_text = cell("address")
        flat = clean_unit_number(cell("flat"))
        if not flat and address_text:
            flat = clean_unit_number(extract_flat_no(address_text))
        cad = extract_cadastral(cell("cad")) or cell("cad")
        fio_text = cell("fio")
        # Старые/новые реестры часто хранят ФИО раздельно: Фамилия | Имя | Отчество.
        last_part = cell("last")
        first_part = cell("first")
        patr_part = cell("patr")
        if not fio_text and (last_part or first_part or patr_part):
            fio_text = " ".join(p for p in [last_part, first_part, patr_part] if p).strip()
        right_text = cell("right")
        date_text = cell("date")
        share_text = cell("share")

        if not right_text:
            # если колонка не определилась, ищем ячейку с правом/договором в строке
            for val in row:
                txt = str(val or "")
                ntxt = normalize_text(txt)
                if "правооблад" in ntxt:
                    continue
                if any(k in ntxt for k in ["собственность", "договор", "регистрац", "правоустанавлива"]):
                    right_text = txt
                    break
        if not fio_text:
            for val in row:
                txt = str(val or "")
                if looks_like_fio(txt):
                    fio_text = txt
                    break

        fio_lines = split_fio_cell(fio_text)
        right_lines = split_multiline_cell(right_text)
        share_from_column_total = extract_share_any(share_text)

        if not fio_lines and not right_lines:
            continue
        if not right_lines:
            right_lines = [""]
        if not fio_lines:
            fio_lines = ["."]

        pair_count = max(len(fio_lines), len(right_lines))
        for idx in range(pair_count):
            fio = fio_lines[idx] if idx < len(fio_lines) else (fio_lines[0] if len(fio_lines) == 1 else ".")
            right_raw = right_lines[idx] if idx < len(right_lines) else (right_lines[0] if len(right_lines) == 1 else "")
            key, reg, reg_date, is_contract, contract_text, share_from_right = parse_right_line(right_raw)
            separate_date = clean_reg_date_any(date_text)
            if not reg_date and separate_date:
                reg_date = separate_date
            if share_from_right:
                share = share_from_right
            elif share_from_column_total:
                # Если в одной старой строке несколько ФИО и одна общая доля 1/1,
                # делим её между ФИО. Если ФИО один — берём долю как есть.
                if len(fio_lines) > 1 and len(right_lines) == 1:
                    share = divide_share(share_from_column_total, len(fio_lines))
                else:
                    share = share_from_column_total
            else:
                share = ""
            last, first, patr = split_fio_or_org(fio)
            fio_from_named_column = bool(fio_text.strip() and hmap.get("fio") is not None)
            address_surname_record = bool(not key and fio_from_named_column and (flat or address_text))
            if not key and not looks_like_fio(fio) and not address_surname_record:
                continue
            if key and fio and fio != ".":
                confidence = "высокая"
            elif address_surname_record:
                confidence = "адрес+фамилия"
            else:
                confidence = "низкая"
            records.append(OldRegistryRecord(
                source_file=source_file,
                sheet=sheet_name,
                row=ri,
                flat_no=flat,
                cadastral_number=str(cad or "").strip(),
                right_key=key,
                reg_number=reg,
                reg_date=reg_date,
                is_contract=is_contract,
                contract_text=contract_text,
                fio_original=fio,
                last_name=last,
                first_name=first,
                patronymic=patr,
                share=share,
                raw_right=right_raw,
                confidence=confidence,
                old_address=address_text,
            ))
    return records


def read_xls_rows(path: Path) -> list[tuple[str, list[list[Any]]]]:
    """
    Читает старые .xls через xlrd.
    Используется для реестров вида LicInfo.xls: Адрес / Фамилия.
    """
    try:
        import xlrd
    except Exception as e:
        raise RuntimeError("Для чтения .xls нужен пакет xlrd. Запустите 00_SETUP.bat после обновления.") from e

    book = xlrd.open_workbook(str(path))
    result: list[tuple[str, list[list[Any]]]] = []
    for sheet in book.sheets():
        rows: list[list[Any]] = []
        for r in range(sheet.nrows):
            row: list[Any] = []
            for c in range(sheet.ncols):
                cell = sheet.cell(r, c)
                value: Any = cell.value
                if cell.ctype == xlrd.XL_CELL_DATE:
                    try:
                        value = xlrd.xldate.xldate_as_datetime(value, book.datemode)
                    except Exception:
                        pass
                row.append(value)
            rows.append(row)
        result.append((sheet.name, rows))
    return result


def extract_old_registry_records(path: Path) -> list[OldRegistryRecord]:
    ext = path.suffix.lower()
    records: list[OldRegistryRecord] = []
    if ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        for ws in wb.worksheets:
            rows = [[ws.cell(r, c).value for c in range(1, ws.max_column + 1)] for r in range(1, ws.max_row + 1)]
            records.extend(extract_old_records_from_rows(rows, path.name, ws.title))
    elif ext == ".xls":
        for sheet_name, rows in read_xls_rows(path):
            records.extend(extract_old_records_from_rows(rows, path.name, sheet_name))
    elif ext == ".docx":
        rows = read_text_docx(path)
        records.extend(extract_old_records_from_rows(rows, path.name, "DOCX"))
    else:
        raise RuntimeError(f"Формат старого реестра пока не поддержан: {ext}. Поддерживаются .xlsx/.xlsm/.xls/.docx")
    apply_inferred_shares(records)
    return records


def copy_row_style(ws, src_row: int, dst_row: int) -> None:
    for c in range(1, ws.max_column + 1):
        src = ws.cell(src_row, c)
        dst = ws.cell(dst_row, c)
        dst._style = copy(src._style)
        if src.number_format:
            dst.number_format = src.number_format
        if src.alignment:
            dst.alignment = copy(src.alignment)
        if src.font:
            dst.font = copy(src.font)
        if src.fill:
            dst.fill = copy(src.fill)
        if src.border:
            dst.border = copy(src.border)


def set_cell(ws, row_idx: int, col_idx: int | None, value: Any) -> None:
    if col_idx:
        ws.cell(row_idx, col_idx).value = value


def dedupe_ambiguous_rows(rows: list[list[Any]]) -> list[list[Any]]:
    """
    Убирает повторы в листе «Неоднозначные».
    Ключ: номер помещения + номер права + причина.
    """
    result: list[list[Any]] = []
    seen: set[tuple[str, str, str]] = set()
    for row in rows:
        flat = str(row[1] if len(row) > 1 else "").strip()
        reg = str(row[2] if len(row) > 2 else "").strip()
        reason = normalize_text(row[4] if len(row) > 4 else "")
        key = (flat.upper(), compact_text(reg).upper(), reason)
        if key in seen:
            continue
        seen.add(key)
        result.append(row)
    return result


def is_org_name(value: str) -> bool:
    t = normalize_text(value)
    return any(k in t for k in ["ооо", "ао", "пао", "зао", "нко", "тсж", "жск", "администрац", "муницип", "собственность рф", "российская федерация"])


def fraction_to_roskvartal_share(value: Any, joint: bool = False) -> str:
    if joint:
        return "1"
    s = str(value or "").strip()
    if not s or s == ".":
        return ""
    if s == "1/1":
        return "1"
    return s


def roskvartal_document_type(row: dict[str, Any]) -> str:
    """Тип документа/вида права для Roskvartal."""
    raw_right = str(row.get("right_type") or "").strip()
    t = normalize_text(raw_right)
    reg = str(row.get("reg") or "").strip()
    share = str(row.get("share") or "").strip()
    if reg == "." or "договор" in normalize_text(reg) or "договор" in t or "старое право" in t:
        return "Договор / старое право"
    if row.get("joint") or "совмест" in t:
        return "Общая совместная собственность"
    if "долев" in t:
        return "Общая долевая собственность"
    if share and share not in {"1", "1/1", "."}:
        return "Общая долевая собственность"
    return "Собственность"


def roskvartal_person_name(last: Any, first: Any, patr: Any) -> tuple[str, str]:
    parts = [str(last or "").strip(), str(first or "").strip(), str(patr or "").strip()]
    clean = [p for p in parts if p and p != "."]
    if not clean:
        return ".", ""
    if len(clean) == 1 and is_org_name(clean[0]):
        return "", clean[0]
    return " ".join(clean), ""


def infer_roskvartal_room_type(row: dict[str, Any]) -> str:
    # Roskvartal принимает тип так, как он указан в справочной информации Росреестра.
    for key in ["type", "Тип помещения", "Тип объекта", "Вид объекта", "Назначение"]:
        val = str(row.get(key) or "").strip()
        if val and val != ".":
            return val
    return "Помещение"


def build_roskvartal_sheet_from_data(wb, *, source_sheet: str = "Данные", target_sheet: str = "Roskvartal") -> None:
    from openpyxl.styles import Font, Alignment
    from openpyxl.utils import get_column_letter

    if source_sheet not in wb.sheetnames:
        return
    ws = wb[source_sheet]
    headers = {normalize_text(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1) if ws.cell(1, c).value}

    def col(*names: str) -> int | None:
        for name in names:
            c = headers.get(normalize_text(name))
            if c:
                return c
        return None

    col_flat = col("Номер помещения (обяз.)", "Номер помещения")
    col_area = col("Площадь помещения, м2 (обяз.)", "Площадь помещения")
    col_cad = col("Кадастровый номер")
    col_joint = col('Признак совместной собс. ("да" для сов. собс.)', "Признак совместной собственности")
    col_share = col("Доля (1/1, 1/2, 1/3 и т.д.) (обяз., сумма в помещении должна быть 1/1)", "Доля")
    col_reg = col("Номер регистрации (обяз.)", "Номер регистрации")
    col_date = col("Дата регистрации")
    col_last = col("Фамилия / Наименование организации (обяз.)", "Фамилия", "ФИО собственника")
    col_first = col("Имя (обяз. для физ лиц)", "Имя")
    col_patr = col("Отчество")
    col_type = col("Тип помещения", "Вид объекта", "Тип объекта", "Назначение")
    col_right_type = col("Вид права", "Тип документа", "Вид права Росреестра")

    # Если источник содержит полный лист «Все сведения», подхватываем оттуда вид объекта по КН.
    full_type_by_cad: dict[str, str] = {}
    if "Все сведения" in wb.sheetnames:
        full_ws = wb["Все сведения"]
        full_headers = {normalize_text(full_ws.cell(1, c).value): c for c in range(1, full_ws.max_column + 1) if full_ws.cell(1, c).value}
        cad_c = full_headers.get(normalize_text("Кадастровый номер"))
        field_c = full_headers.get(normalize_text("Поле"))
        val_c = full_headers.get(normalize_text("Значение"))
        if cad_c and field_c and val_c:
            for rr in range(2, full_ws.max_row + 1):
                field_name = normalize_text(full_ws.cell(rr, field_c).value)
                if "вид объекта" in field_name:
                    cad_val = str(full_ws.cell(rr, cad_c).value or "").strip()
                    val = str(full_ws.cell(rr, val_c).value or "").strip()
                    if cad_val and val and val != ".":
                        full_type_by_cad.setdefault(cad_val, val)

    rows: list[dict[str, Any]] = []
    for r in range(2, ws.max_row + 1):
        def get(c: int | None) -> Any:
            return ws.cell(r, c).value if c else ""
        flat = str(get(col_flat) or "").strip()
        if not flat:
            continue
        rows.append({
            "row_idx": r,
            "flat": flat,
            "area": str(get(col_area) or "").strip(),
            "cad": str(get(col_cad) or "").strip(),
            "joint": normalize_text(get(col_joint)) == "да",
            "share": str(get(col_share) or "").strip(),
            "reg": str(get(col_reg) or "").strip(),
            "date": str(get(col_date) or "").strip(),
            "last": str(get(col_last) or "").strip(),
            "first": str(get(col_first) or "").strip(),
            "patr": str(get(col_patr) or "").strip(),
            "type": str(get(col_type) or full_type_by_cad.get(str(get(col_cad) or "").strip(), "")).strip(),
            "right_type": str(get(col_right_type) or "").strip(),
        })

    if target_sheet in wb.sheetnames:
        del wb[target_sheet]
    out = wb.create_sheet(target_sheet)
    for i, h in enumerate(ROSKVARTAL_HEADERS, start=1):
        out.cell(1, i).value = h
        out.cell(1, i).font = Font(bold=True)
        out.cell(1, i).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    widths = {1: 16, 2: 16, 3: 18, 4: 38, 5: 30, 6: 12, 7: 18, 8: 28, 9: 18, 10: 24, 11: 18, 12: 18}
    for idx, width in widths.items():
        out.column_dimensions[get_column_letter(idx)].width = width
    try:
        out.freeze_panes = "A2"
    except Exception:
        pass

    distinct_cads_by_flat: dict[str, set[str]] = {}
    for row in rows:
        flat_key = clean_unit_number(row.get("flat")) or str(row.get("flat") or "").upper()
        cad = str(row.get("cad") or "").strip()
        if flat_key and cad and cad != ".":
            distinct_cads_by_flat.setdefault(flat_key, set()).add(cad)

    used_row_ids: set[int] = set()
    output_rows: list[list[Any]] = []

    def append_ros_row(base: dict[str, Any], fio: str, org: str, share: str, doc_type: str | None = None) -> None:
        flat_key = clean_unit_number(base.get("flat")) or str(base.get("flat") or "").upper()
        unique = "да" if len(distinct_cads_by_flat.get(flat_key, set())) > 1 else "нет"
        output_rows.append([
            base.get("flat") or ".",
            base.get("area") or ".",
            infer_roskvartal_room_type(base),
            fio or ".",
            org or "",
            share,
            doc_type or roskvartal_document_type(base),
            base.get("reg") or ".",
            base.get("date") or "",
            "",
            "",
            unique,
        ])

    joint_groups: dict[tuple[str, str, str], list[dict[str, Any]]] = {}
    for row in rows:
        if row.get("joint"):
            key = (
                (clean_unit_number(row.get("flat")) or str(row.get("flat") or "")).upper(),
                str(row.get("cad") or ""),
                normalize_reg_key(str(row.get("reg") or "")) or str(row.get("reg") or ""),
            )
            joint_groups.setdefault(key, []).append(row)

    for _key, group in joint_groups.items():
        for row in group:
            used_row_ids.add(int(row["row_idx"]))
        base = group[0]
        fio_parts: list[str] = []
        org_parts: list[str] = []
        for row in group:
            fio, org = roskvartal_person_name(row.get("last"), row.get("first"), row.get("patr"))
            if fio and fio != ".":
                fio_parts.append(fio)
            if org:
                org_parts.append(org)
        append_ros_row(base, ", ".join(fio_parts) if fio_parts else ".", ", ".join(org_parts), "1", "Общая совместная собственность")

    for row in rows:
        if int(row["row_idx"]) in used_row_ids:
            continue
        fio, org = roskvartal_person_name(row.get("last"), row.get("first"), row.get("patr"))
        append_ros_row(row, fio, org, fraction_to_roskvartal_share(row.get("share"), joint=False))

    for r_idx, row in enumerate(output_rows, start=2):
        for c_idx, val in enumerate(row, start=1):
            cell = out.cell(r_idx, c_idx)
            cell.value = val
            if c_idx in {1, 8, 9}:
                cell.number_format = "@"


def append_sheet_rows(wb, name: str, headers: list[str], rows: list[list[Any]]) -> None:
    if name in wb.sheetnames:
        del wb[name]
    ws = wb.create_sheet(name)
    for c, h in enumerate(headers, start=1):
        ws.cell(1, c).value = h
    for r, row in enumerate(rows, start=2):
        for c, v in enumerate(row, start=1):
            ws.cell(r, c).value = v
    for c in range(1, len(headers) + 1):
        ws.column_dimensions[ws.cell(1, c).column_letter].width = 24
    try:
        ws.freeze_panes = "A2"
    except Exception:
        pass


def split_old_search_result_sheets(wb) -> None:
    """
    Если новый реестр был создан старой версией, в листе «Спорные результаты»
    могли лежать и реальные спорные случаи, и обычные отклонённые строки выдачи.
    Разносим их по двум листам.
    """
    source_name = "Спорные результаты"
    if source_name not in wb.sheetnames:
        return
    src = wb[source_name]
    rows = list(src.iter_rows(values_only=True))
    if not rows:
        return

    headers = [str(v or "") for v in rows[0]]
    rejected_rows: list[list[Any]] = []
    disputed_rows: list[list[Any]] = []

    for row in rows[1:]:
        values = list(row)
        joined = normalize_text(" ".join(str(v or "") for v in values))
        # Обычные отклонения выдачи обычно содержат решения/причины "отклонено", "не совпал", "не подходит".
        is_rejected = any(x in joined for x in [
            "отклон",
            "не совпал",
            "не совпадает",
            "не подходит",
            "другой дом",
            "другая квартира",
            "не точное совпадение",
        ])
        if is_rejected:
            rejected_rows.append(values)
        else:
            disputed_rows.append(values)

    del wb[source_name]
    append_sheet_rows(wb, "Отклонённые результаты поиска", headers, rejected_rows)
    append_sheet_rows(wb, "Спорные результаты", headers, disputed_rows)


def remove_redundant_placeholder_rows(ws, header_row: int, cols: dict[str, int | None], ambiguous: list[list[Any]], match_log: list[list[Any]]) -> int:
    """
    Удаляет лишние строки-заготовки с правом "." после сопоставления,
    если остальные строки этого помещения уже покрывают долю 1/1.

    Это исправляет ситуацию:
    - парсер создал строку "." под старую приватизацию;
    - сопоставление развернуло старых собственников;
    - строка "." осталась лишней и ломает сумму долей.
    """
    flat_col = cols.get("flat")
    share_col = cols.get("share")
    reg_col = cols.get("reg")
    if not flat_col or not share_col or not reg_col:
        return 0

    groups: dict[str, list[int]] = {}
    for r in range(header_row + 1, ws.max_row + 1):
        flat = clean_unit_number(ws.cell(r, flat_col).value)
        if flat:
            groups.setdefault(flat.upper(), []).append(r)

    rows_to_delete: list[int] = []
    try:
        from fractions import Fraction
        for flat, rows in groups.items():
            placeholder_rows: list[int] = []
            real_fracs: list[Any] = []
            for r in rows:
                reg = str(ws.cell(r, reg_col).value or "").strip()
                share = str(ws.cell(r, share_col).value or "").strip()
                is_placeholder = (not reg or reg == ".") and is_empty_person_cells(ws, r, cols)
                frac = parse_fraction_share(share)
                if is_placeholder:
                    placeholder_rows.append(r)
                elif frac is not None:
                    real_fracs.append(frac)
            if placeholder_rows and real_fracs:
                total = sum(real_fracs, Fraction(0, 1))
                if total == Fraction(1, 1):
                    rows_to_delete.extend(placeholder_rows)
                    ambiguous.append([
                        ", ".join(str(r) for r in placeholder_rows),
                        flat,
                        ".",
                        "",
                        "Удалена лишняя строка '.' после полного сопоставления долей 1/1",
                        "",
                    ])
                    for r in placeholder_rows:
                        match_log.append([r, flat, ".", "удалено", "лишняя строка '.'", "", "", "Остальные строки помещения покрывают 1/1"])
    except Exception:
        return 0

    for r in sorted(set(rows_to_delete), reverse=True):
        ws.delete_rows(r, 1)
    return len(set(rows_to_delete))


def build_reconcile_validation_rows(ws, header_row: int, cols: dict[str, int | None]) -> list[list[Any]]:
    """Проверка итогового файла после сопоставления."""
    rows: list[list[Any]] = []
    now = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    flat_col = cols.get("flat")
    share_col = cols.get("share")
    joint_col = cols.get("joint")

    def add(level: str, row_idx: Any, field: str, message: str) -> None:
        rows.append([now, level, ws.title, row_idx, field, message])

    # Доли по помещению.
    if flat_col and share_col:
        from fractions import Fraction
        groups: dict[str, list[tuple[int, Any, str, bool]]] = {}
        for r in range(header_row + 1, ws.max_row + 1):
            flat = clean_unit_number(ws.cell(r, flat_col).value)
            if not flat:
                continue
            raw_share = str(ws.cell(r, share_col).value or "").strip()
            frac = parse_fraction_share(raw_share)
            joint_yes = False
            if joint_col:
                joint_yes = normalize_text(ws.cell(r, joint_col).value) == "да"
            groups.setdefault(flat.upper(), []).append((r, frac, raw_share, joint_yes))

        for flat, items in groups.items():
            if any(joint_yes for _r, _frac, _raw, joint_yes in items):
                # Совместная собственность: 1/1 у каждого совместного собственника допустимо.
                continue
            if any(frac is None for _r, frac, raw, _joint in items if raw and raw != "."):
                bad = ", ".join(f"{r}: {raw}" for r, frac, raw, _joint in items if raw and raw != "." and frac is None)
                add("Ошибка", ", ".join(str(r) for r, _f, _raw, _joint in items), "Доля", f"Некорректная доля по помещению {flat}: {bad}")
                continue
            fracs = [frac for _r, frac, raw, _joint in items if frac is not None]
            if not fracs:
                continue
            total = sum(fracs, Fraction(0, 1))
            if total != Fraction(1, 1):
                shares = ", ".join(raw for _r, _f, raw, _joint in items)
                add("Ошибка", ", ".join(str(r) for r, _f, _raw, _joint in items), "Доля", f"Сумма долей по помещению {flat} = {total}, не равна 1/1. Доли: {shares}")

    # Шумные предупреждения по одинаковому КН + номеру права намеренно не пишем.
    # Для долевой/совместной собственности это часто нормальная ситуация.
    return rows


def reconcile_registries(
    new_path: Path,
    old_path: Path | str | list[Path],
    output_path: Path,
    *,
    transfer_shares: bool = True,
    allow_contracts_for_empty_right: bool = True,
    expand_empty_right_rows: bool = True,
    preserve_existing_fio: bool = True,
    allow_address_surname_match: bool = False,
    export_format: str = "burmistr",
) -> dict[str, Any]:
    from openpyxl import load_workbook

    if isinstance(old_path, list):
        old_paths = old_path
    else:
        old_paths = resolve_old_registry_paths(old_path)
        if not old_paths and old_path:
            old_paths = [Path(str(old_path))]

    old_records: list[OldRegistryRecord] = []
    for one_old_path in old_paths:
        old_records.extend(extract_old_registry_records(Path(one_old_path)))
    old_records = dedupe_old_records(old_records)

    reg_index: dict[str, list[OldRegistryRecord]] = {}
    contract_by_flat: dict[str, list[OldRegistryRecord]] = {}
    surname_by_flat: dict[str, list[OldRegistryRecord]] = {}
    for rec in old_records:
        if rec.right_key:
            if rec.is_contract:
                if rec.flat_no:
                    contract_by_flat.setdefault(rec.flat_no.upper(), []).append(rec)
            else:
                reg_index.setdefault(rec.right_key, []).append(rec)
        elif rec.flat_no and rec.last_name and rec.last_name != ".":
            # Низкоуверенный реестр без прав: Адрес + Фамилия.
            surname_by_flat.setdefault(rec.flat_no.upper(), []).append(rec)

    wb = load_workbook(new_path)

    split_old_search_result_sheets(wb)

    # Старые файлы парсинга могли содержать шумные служебные листы.
    # При сопоставлении пересобираем их заново, чтобы не переносить старый шум в итоговый файл.
    for noisy_sheet in ["Проверка", "Дубли", "Неоднозначные", "Сопоставление", "Распознанный старый реестр", "Сводка сопоставления"]:
        if noisy_sheet in wb.sheetnames:
            del wb[noisy_sheet]

    ws, header_row = find_new_data_sheet(wb)
    cols = get_new_header_columns(ws, header_row)

    new_rows_by_flat: dict[str, int] = {}
    new_rows_by_flat_reg: dict[tuple[str, str], list[int]] = {}
    for r in range(header_row + 1, ws.max_row + 1):
        f = clean_unit_number(ws.cell(r, cols["flat"]).value if cols.get("flat") else "")
        if f:
            f_key = f.upper()
            new_rows_by_flat[f_key] = new_rows_by_flat.get(f_key, 0) + 1
            reg_val = str(ws.cell(r, cols["reg"]).value if cols.get("reg") else "").strip()
            reg_key_val = normalize_reg_key(reg_val)
            if reg_key_val:
                new_rows_by_flat_reg.setdefault((f_key, reg_key_val), []).append(r)

    recognized_rows = []
    for rec in old_records:
        recognized_rows.append([
            rec.source_file, rec.sheet, rec.row, rec.flat_no, rec.cadastral_number,
            rec.old_address, rec.right_key, rec.reg_number, rec.reg_date, "да" if rec.is_contract else "",
            rec.fio_original, rec.last_name, rec.first_name, rec.patronymic,
            rec.share, rec.raw_right, rec.confidence,
        ])

    match_log: list[list[Any]] = []
    ambiguous: list[list[Any]] = []
    stats = {"matched_by_reg": 0, "matched_by_contract": 0, "matched_by_address_surname": 0, "expanded_rows": 0, "ambiguous": 0, "not_found": 0, "skipped_existing": 0}

    def try_address_surname_match(row_idx: int, flat: str, reg_raw: str) -> bool:
        """Низкоуверенное сопоставление по реестру Адрес+Фамилия."""
        if not flat:
            return False
        records = surname_by_flat.get(flat.upper(), [])
        if not records:
            return False

        variants = "; ".join(
            f"{r.fio_original} [{r.old_address or 'без адреса'}; {r.source_file}:{r.row}]"
            for r in records
        )

        if not allow_address_surname_match:
            stats["ambiguous"] += 1
            reason = "Адрес+Фамилия: слабое сопоставление отключено, ФИО не перенесено"
            ambiguous.append([row_idx, flat, reg_raw, "", reason, variants])
            highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)
            match_log.append([row_idx, flat, reg_raw, "не перенесено", "адрес + фамилия отключено", "", "", variants])
            return True

        if len(records) == 1 and new_rows_by_flat.get(flat.upper(), 0) == 1:
            rec = records[0]
            if preserve_existing_fio and not is_empty_person_cells(ws, row_idx, cols):
                stats["skipped_existing"] += 1
                match_log.append([row_idx, flat, reg_raw, "пропущено", "Адрес+Фамилия: ФИО уже заполнено", rec.fio_original, "", f"{rec.source_file} / {rec.sheet} / строка {rec.row}"])
                return True
            set_cell(ws, row_idx, cols.get("last"), rec.last_name or ".")
            set_cell(ws, row_idx, cols.get("first"), rec.first_name or ".")
            set_cell(ws, row_idx, cols.get("patr"), rec.patronymic or ".")
            stats["matched_by_address_surname"] += 1
            reason = "Адрес+Фамилия: перенесено с низкой уверенностью, проверить вручную"
            ambiguous.append([row_idx, flat, reg_raw, "", reason, variants])
            stats["ambiguous"] += 1
            highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)
            match_log.append([
                row_idx, flat, reg_raw, "найдено", "адрес + фамилия (низкая уверенность, подсвечено)",
                rec.fio_original, "", f"{rec.source_file} / {rec.sheet} / строка {rec.row}; адрес: {rec.old_address}",
            ])
            return True

        stats["ambiguous"] += 1
        reason = "Адрес+Фамилия: неоднозначное сопоставление"
        if len(records) > 1:
            reason += f"; в старом реестре несколько строк по помещению {flat}"
        if new_rows_by_flat.get(flat.upper(), 0) > 1:
            reason += f"; в новом реестре несколько строк по помещению {flat}"
        ambiguous.append([row_idx, flat, reg_raw, "", reason, variants])
        highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)
        match_log.append([row_idx, flat, reg_raw, "неоднозначно", "адрес + фамилия", "", "", variants])
        return True

    # Идём снизу вверх, чтобы можно было вставлять строки для разворачивания договоров.
    for row_idx in range(ws.max_row, header_row, -1):
        flat = clean_unit_number(ws.cell(row_idx, cols["flat"]).value if cols.get("flat") else "")
        reg_raw = str(ws.cell(row_idx, cols["reg"]).value if cols.get("reg") else "").strip()
        share_current = str(ws.cell(row_idx, cols["share"]).value if cols.get("share") else "").strip()
        reg_key = normalize_reg_key(reg_raw)

        chosen: list[OldRegistryRecord] = []
        match_type = ""
        comment = ""

        should_expand = False
        if reg_key:
            records = reg_index.get(reg_key, [])
            if len(records) == 1:
                chosen = records
                match_type = "номер регистрации"
            elif len(records) > 1:
                by_flat = [r for r in records if flat and r.flat_no.upper() == flat.upper()]
                if len(by_flat) == 1:
                    chosen = by_flat
                    match_type = "номер регистрации + помещение"
                elif len(by_flat) > 1 and expand_empty_right_rows:
                    # В старых реестрах один номер регистрации может повторяться на нескольких ФИО
                    # по одному помещению. Если в новом реестре есть несколько строк с этим правом,
                    # привязываем по порядку. Если в Росреестре совместная собственность одной строкой —
                    # разворачиваем в несколько строк с признаком совместной собственности.
                    same_new_rows = sorted(new_rows_by_flat_reg.get((flat.upper(), reg_key), [])) if flat else []
                    joint_col = cols.get("joint")
                    is_joint_row = False
                    if joint_col:
                        is_joint_row = normalize_text(ws.cell(row_idx, joint_col).value) == "да"

                    if is_joint_row and len(same_new_rows) <= 1:
                        chosen = by_flat
                        match_type = "совместная собственность, развёрнуто по старому реестру"
                        should_expand = True
                        stats["ambiguous"] += 1
                        reason = "Совместная собственность развёрнута по старому реестру, проверить вручную"
                        ambiguous.append([row_idx, flat, reg_raw, "", reason, "; ".join(r.fio_original for r in by_flat)])
                        highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)
                    elif same_new_rows:
                        ordinal = same_new_rows.index(row_idx) if row_idx in same_new_rows else 0
                        if ordinal < len(by_flat):
                            chosen = [by_flat[ordinal]]
                            should_expand = False
                            if len(same_new_rows) < len(by_flat):
                                match_type = "номер регистрации + помещение, по порядку, часть старых собственников не перенесена"
                                stats["ambiguous"] += 1
                                reason = f"Старых записей по праву больше ({len(by_flat)}), чем строк Росреестра с этим правом ({len(same_new_rows)}); перенесены только строки с сохранившимся правом"
                                ambiguous.append([row_idx, flat, reg_raw, "", reason, "; ".join(r.fio_original for r in by_flat)])
                                highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)
                            else:
                                match_type = "номер регистрации + помещение, по порядку"
                        else:
                            stats["ambiguous"] += 1
                            reason = f"В новом реестре строк по праву больше ({len(same_new_rows)}), чем записей в старом ({len(by_flat)}); требуется проверка"
                            ambiguous.append([row_idx, flat, reg_raw, "", reason, "; ".join(r.fio_original for r in by_flat)])
                            highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)
                            continue
                    else:
                        chosen = by_flat
                        match_type = "номер регистрации + помещение, развёрнуто"
                        should_expand = True
                else:
                    stats["ambiguous"] += 1
                    reason = "Один номер права найден у нескольких ФИО, но не удалось однозначно привязать к помещению"
                    ambiguous.append([row_idx, flat, reg_raw, "", reason, "; ".join(r.fio_original for r in records)])
                    highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)
                    continue
            else:
                if try_address_surname_match(row_idx, flat, reg_raw):
                    continue
                stats["not_found"] += 1
                match_log.append([row_idx, flat, reg_raw, "не найдено", "новое право / нет в старом реестре", "", "", "Нет совпадения по номеру регистрации"])
                continue
        else:
            # Только если в новом реестре номер регистрации "."/пустой — можно сопоставлять старые договоры.
            if allow_contracts_for_empty_right and flat:
                records = contract_by_flat.get(flat.upper(), [])
                if records:
                    chosen = records
                    match_type = "договор при пустом новом праве"
                    should_expand = len(records) > 1 and expand_empty_right_rows
                else:
                    if try_address_surname_match(row_idx, flat, reg_raw):
                        continue
                    stats["not_found"] += 1
                    match_log.append([row_idx, flat, reg_raw, "не найдено", "пустое право", "", "", "Пустое право, договоры по помещению не найдены"])
                    continue
            else:
                continue

        if not chosen:
            continue

        if len(chosen) > 1 and should_expand:
            # Текущую строку заполняем первой записью, остальные вставляем ниже.
            base_values = [ws.cell(row_idx, c).value for c in range(1, ws.max_column + 1)]
            for extra_i in range(len(chosen) - 1):
                ws.insert_rows(row_idx + 1)
                copy_row_style(ws, row_idx, row_idx + 1)
                for c, v in enumerate(base_values, start=1):
                    ws.cell(row_idx + 1, c).value = v
            target_rows = [row_idx + i for i in range(len(chosen))]
            stats["expanded_rows"] += max(0, len(chosen) - 1)
        else:
            target_rows = [row_idx]
            if len(chosen) > 1:
                # Не разворачиваем — берём первую, остальные пишем в неоднозначные.
                stats["ambiguous"] += 1
                reason = "Найдено несколько старых записей, строка не развёрнута"
                ambiguous.append([row_idx, flat, reg_raw, "", reason, "; ".join(r.fio_original for r in chosen)])
                highlight_reconcile_rows(ws, [row_idx], share_col=cols.get("share"), reason=reason)

        for target_row, rec in zip(target_rows, chosen):
            if preserve_existing_fio and not is_empty_person_cells(ws, target_row, cols):
                stats["skipped_existing"] += 1
                match_log.append([
                    target_row, flat, reg_raw, "пропущено", "ФИО уже заполнено",
                    rec.fio_original, rec.share, f"{rec.source_file} / {rec.sheet} / строка {rec.row}",
                ])
                continue

            set_cell(ws, target_row, cols.get("last"), rec.last_name or ".")
            set_cell(ws, target_row, cols.get("first"), rec.first_name or ".")
            set_cell(ws, target_row, cols.get("patr"), rec.patronymic or ".")
            if "совместная собственность" in match_type:
                set_cell(ws, target_row, cols.get("joint"), "да")
                set_cell(ws, target_row, cols.get("share"), "1/1")
            if rec.is_contract and (not reg_raw or reg_raw == "."):
                # В номер регистрации нового формата пишем очищенный договор без "Собственность/Долевая собственность" и без даты.
                # Дату договора пишем в отдельную колонку даты.
                visible_contract = format_contract_for_output(rec)
                set_cell(ws, target_row, cols.get("reg"), visible_contract)
                if cols.get("right_type"):
                    set_cell(ws, target_row, cols.get("right_type"), "Договор / старое право")
                if cols.get("reg_date") and rec.reg_date:
                    set_cell(ws, target_row, cols.get("reg_date"), rec.reg_date)
            elif cols.get("right_type") and "совместная собственность" in match_type:
                set_cell(ws, target_row, cols.get("right_type"), "Общая совместная собственность")
            share_note = ""
            if transfer_shares and rec.share and cols.get("share") and "совместная собственность" not in match_type:
                if should_skip_old_share_override(ws, header_row, cols, target_row, flat, rec.share, expanded=bool(should_expand and len(chosen) > 1)):
                    share_note = f"; доля из старого реестра {rec.share} не перенесена, сохранена доля Росреестра {ws.cell(target_row, cols.get('share')).value}"
                else:
                    set_cell(ws, target_row, cols.get("share"), rec.share)

            if rec.is_contract:
                stats["matched_by_contract"] += 1
            else:
                stats["matched_by_reg"] += 1
            match_log.append([
                target_row, flat, reg_raw, "найдено", match_type,
                rec.fio_original, rec.share, f"{rec.source_file} / {rec.sheet} / строка {rec.row}{share_note}",
            ])

    removed_placeholders = remove_redundant_placeholder_rows(ws, header_row, cols, ambiguous, match_log)
    stats["removed_placeholders"] = removed_placeholders

    share_warning_count = append_share_sum_warnings(ws, header_row, cols, ambiguous)
    stats["ambiguous"] += share_warning_count

    validation_rows = build_reconcile_validation_rows(ws, header_row, cols)

    append_sheet_rows(
        wb,
        "Распознанный старый реестр",
        ["Файл", "Лист", "Строка", "Номер помещения", "Кадастровый номер", "Адрес старого реестра", "Ключ права", "Номер регистрации", "Дата", "Договор", "ФИО исходное", "Фамилия/Орг", "Имя", "Отчество", "Доля", "Исходный текст права", "Уверенность"],
        recognized_rows,
    )
    append_sheet_rows(
        wb,
        "Сопоставление",
        ["Строка нового", "Номер помещения", "Номер права новый", "Статус", "Тип совпадения", "ФИО перенесено", "Доля", "Источник"],
        match_log,
    )
    ambiguous = dedupe_ambiguous_rows(ambiguous)
    stats["ambiguous"] = len(ambiguous)

    append_sheet_rows(
        wb,
        "Неоднозначные",
        ["Строка нового", "Номер помещения", "Номер права", "Кадастровый номер", "Причина", "Варианты"],
        ambiguous,
    )
    append_sheet_rows(
        wb,
        "Проверка",
        VALIDATION_HEADERS,
        validation_rows,
    )
    append_sheet_rows(
        wb,
        "Сводка сопоставления",
        ["Показатель", "Значение"],
        [
            ["Старых записей распознано", len(old_records)],
            ["Совпадений по номеру регистрации", stats["matched_by_reg"]],
            ["Совпадений по договору при пустом праве", stats["matched_by_contract"]],
            ["Совпадений по адресу и фамилии", stats["matched_by_address_surname"]],
            ["Добавлено строк при разворачивании", stats["expanded_rows"]],
            ["Неоднозначных случаев", stats["ambiguous"]],
            ["Замечаний проверки", len(validation_rows)],
            ["Не найдено", stats["not_found"]],
            ["Пропущено, потому что ФИО уже было заполнено", stats.get("skipped_existing", 0)],
            ["Удалено лишних строк с точкой", stats.get("removed_placeholders", 0)],
            ["Старых реестров обработано", len(old_paths)],
            ["Перенос долей", "да" if transfer_shares else "нет"],
            ["Договоры при пустом праве", "да" if allow_contracts_for_empty_right else "нет"],
            ["Разворачивать пустые строки", "да" if expand_empty_right_rows else "нет"],
            ["Слабое сопоставление Адрес+Фамилия", "да" if allow_address_surname_match else "нет"],
        ],
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if str(export_format or "burmistr") == "roskvartal":
        build_roskvartal_sheet_from_data(wb)
        result_sheets = {"Roskvartal"}
    else:
        result_sheets = RESULT_SHEETS
    log_path = save_result_and_log_workbooks(wb, output_path, result_sheets=result_sheets, log_kind="matched")
    return {"old_records": len(old_records), **stats, "output": str(output_path), "log_output": str(log_path)}



# ----------------------------- standalone Excel validation -----------------------------


def validate_existing_registry_excel(input_path: Path, output_path: Path) -> dict[str, Any]:
    """Проверяет уже готовый Excel без запуска парсинга. Создаёт/перезаписывает листы Проверка и Дубли."""
    from openpyxl import load_workbook

    wb = load_workbook(input_path)
    ws, header_row = find_new_data_sheet(wb)
    cols = get_new_header_columns(ws, header_row)

    def make_sheet(name: str, headers: list[str]):
        if name in wb.sheetnames:
            del wb[name]
        sheet = wb.create_sheet(name)
        for c, h in enumerate(headers, start=1):
            sheet.cell(1, c).value = h
            sheet.column_dimensions[sheet.cell(1, c).column_letter].width = 24
        try:
            sheet.freeze_panes = "A2"
        except Exception:
            pass
        return sheet

    check_ws = make_sheet("Проверка", VALIDATION_HEADERS)
    dup_ws = make_sheet("Дубли", DUPLICATE_HEADERS)

    def add_check(level: str, sheet: str, row_idx: Any, field: str, message: str) -> None:
        r = check_ws.max_row + 1
        vals = [datetime.now().strftime("%d.%m.%Y %H:%M:%S"), level, sheet, row_idx, field, message]
        for c, v in enumerate(vals, start=1):
            check_ws.cell(r, c).value = v

    def add_dup(dup_type: str, key: str, where: str, cad: str = "", reg: str = "", comment: str = "") -> None:
        r = dup_ws.max_row + 1
        vals = [datetime.now().strftime("%d.%m.%Y %H:%M:%S"), dup_type, key, where, cad, reg, comment]
        for c, v in enumerate(vals, start=1):
            dup_ws.cell(r, c).value = v

    area_col = get_col_by_header(ws, header_row, ["площад"])
    data_cads: set[str] = set()
    aux_cads: set[str] = set()
    seen_reg_pairs: set[tuple[str, str]] = set()
    flat_groups: dict[str, list[tuple[int, Any, str]]] = {}
    reg_to_people: dict[str, set[str]] = {}

    for row_idx in range(header_row + 1, ws.max_row + 1):
        flat = clean_unit_number(ws.cell(row_idx, cols["flat"]).value if cols.get("flat") else "")
        cad = str(ws.cell(row_idx, cols["cad"]).value if cols.get("cad") else "").strip()
        reg = str(ws.cell(row_idx, cols["reg"]).value if cols.get("reg") else "").strip()
        reg_date = str(ws.cell(row_idx, cols["reg_date"]).value if cols.get("reg_date") else "").strip()
        share = str(ws.cell(row_idx, cols["share"]).value if cols.get("share") else "").strip()
        area = str(ws.cell(row_idx, area_col).value if area_col else "").strip()
        last = str(ws.cell(row_idx, cols["last"]).value if cols.get("last") else "").strip()
        first = str(ws.cell(row_idx, cols["first"]).value if cols.get("first") else "").strip()
        patr = str(ws.cell(row_idx, cols["patr"]).value if cols.get("patr") else "").strip()

        if not cad or cad == ".":
            add_check("Ошибка", ws.title, row_idx, "Кадастровый номер", "Пустой кадастровый номер")
        if not flat or flat == ".":
            add_check("Предупреждение", ws.title, row_idx, "Номер помещения", "Пустой номер помещения")
        if not area or area == ".":
            add_check("Предупреждение", ws.title, row_idx, "Площадь", "Пустая площадь")
        if reg == "." and reg_date:
            add_check("Ошибка", ws.title, row_idx, "Дата регистрации", "Если номер регистрации '.', дата должна быть пустой")
        if reg and reg != "." and not reg_date and "договор" not in normalize_text(reg):
            add_check("Предупреждение", ws.title, row_idx, "Дата регистрации", "Есть номер регистрации, но дата регистрации пустая")

        frac = parse_fraction_share(share)
        if share and share != "." and frac is None:
            add_check("Ошибка", ws.title, row_idx, "Доля", f"Некорректный формат доли: {share}")
        if flat and flat != ".":
            flat_groups.setdefault(flat.upper(), []).append((row_idx, frac, share))

        if cad and cad != ".":
            data_cads.add(cad)
        if cad and cad != "." and reg:
            pair = (cad, reg)
            if pair in seen_reg_pairs:
                add_dup("Кадастровый номер + регистрация", f"{cad}|{reg}", f"{ws.title}!{row_idx}", cad, reg, "Повтор в основном листе")
            seen_reg_pairs.add(pair)

        if reg and reg != ".":
            person_key = normalize_text(" ".join([last, first, patr])).strip()
            if person_key and person_key != ".":
                reg_to_people.setdefault(normalize_reg_key(reg) or reg, set()).add(person_key)

    try:
        from fractions import Fraction
        for flat, items in flat_groups.items():
            if not items or any(frac is None for _r, frac, raw in items if raw and raw != "."):
                continue
            fractions = [frac for _r, frac, _raw in items if frac is not None]
            if not fractions:
                continue
            total = sum(fractions, Fraction(0, 1))
            if total != Fraction(1, 1):
                rows = ", ".join(str(r) for r, _f, _raw in items)
                shares = ", ".join(str(raw) for _r, _f, raw in items)
                add_check("Ошибка", ws.title, rows, "Доля", f"Сумма долей по помещению {flat} = {total}, не равна 1/1. Доли: {shares}")
    except Exception as e:
        add_check("Предупреждение", ws.title, "", "Доля", f"Не удалось проверить суммы долей: {e}")

    for reg_key, people in reg_to_people.items():
        if len(people) > 1:
            add_check("Предупреждение", ws.title, "", "Номер регистрации", f"Один номер права встречается у разных ФИО: {reg_key}; {', '.join(sorted(people)[:5])}")

    if "Здания и участки" in wb.sheetnames:
        aux = wb["Здания и участки"]
        # В этом листе КН обычно в 4 колонке.
        for row_idx in range(2, aux.max_row + 1):
            cad = str(aux.cell(row_idx, 4).value or "").strip()
            if cad and cad != ".":
                if cad in aux_cads:
                    add_dup("Кадастровый номер", cad, f"Здания и участки!{row_idx}", cad, "", "Повтор на листе зданий/участков")
                if cad in data_cads:
                    add_dup("КН в основном и вспомогательном листе", cad, f"Здания и участки!{row_idx}", cad, "", "Один КН есть и в основном листе, и в зданиях/участках")
                aux_cads.add(cad)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)
    return {"checks": max(0, check_ws.max_row - 1), "duplicates": max(0, dup_ws.max_row - 1), "output": str(output_path)}


# ----------------------------- Excel writer -----------------------------


def collect_snt_not_found_plots_from_workbook(path: Path) -> list[str]:
    """Читает result-файл СНТ и возвращает номера участков со строками «не найдено»."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
    except Exception:
        return []
    if "Земельные участки" not in wb.sheetnames:
        return []
    ws = wb["Земельные участки"]
    headers = {normalize_text(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1) if ws.cell(1, c).value}
    plot_col = headers.get(normalize_text("Номер участка"))
    cad_col = headers.get(normalize_text("Кадастровый номер"))
    comment_col = headers.get(normalize_text("Комментарий"))
    status_col = headers.get(normalize_text("Статус объекта"))
    if not plot_col:
        return []
    result: list[str] = []
    seen: set[str] = set()
    for r in range(2, ws.max_row + 1):
        plot = normalize_plot_number(ws.cell(r, plot_col).value)
        if not plot:
            continue
        cad = str(ws.cell(r, cad_col).value or "").strip() if cad_col else ""
        comment = normalize_text(ws.cell(r, comment_col).value) if comment_col else ""
        status = normalize_text(ws.cell(r, status_col).value) if status_col else ""
        is_not_found = (
            "не найден" in comment
            or (cad in {"", "."} and status in {"", "."})
        )
        if is_not_found and plot not in seen:
            seen.add(plot)
            result.append(plot)
    return result


def snt_result_totals_from_sheet(ws) -> dict[str, int]:
    """Итоги по всему листу «Земельные участки», считая и строки, и уникальные номера участков."""
    headers = {normalize_text(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1) if ws.cell(1, c).value}
    plot_col = headers.get(normalize_text("Номер участка"))
    cad_col = headers.get(normalize_text("Кадастровый номер"))
    status_col = headers.get(normalize_text("Статус объекта"))
    comment_col = headers.get(normalize_text("Комментарий"))
    if not plot_col:
        return {
            "Строк в result-файле по земельным участкам": 0,
            "Уникальных земельных участков в result-файле": 0,
            "Актуальных земельных участков, уникально": 0,
            "Погашенных земельных участков, уникально": 0,
            "Не найдено участков, уникально": 0,
            "Участков с несколькими статусами": 0,
        }
    all_plots: set[str] = set()
    actual: set[str] = set()
    redeemed: set[str] = set()
    not_found: set[str] = set()
    status_by_plot: dict[str, set[str]] = {}
    data_rows = 0
    for r in range(2, ws.max_row + 1):
        row_has_data = any(str(ws.cell(r, c).value or "").strip() for c in range(1, ws.max_column + 1))
        if not row_has_data:
            continue
        data_rows += 1
        plot = normalize_plot_number(ws.cell(r, plot_col).value)
        if not plot:
            continue
        all_plots.add(plot)
        cad = str(ws.cell(r, cad_col).value or "").strip() if cad_col else ""
        status = normalize_text(ws.cell(r, status_col).value) if status_col else ""
        comment = normalize_text(ws.cell(r, comment_col).value) if comment_col else ""
        if "погаш" in status:
            redeemed.add(plot)
            status_by_plot.setdefault(plot, set()).add("Погашено")
        elif "не найден" in comment or (cad in {"", "."} and status in {"", "."}):
            not_found.add(plot)
            status_by_plot.setdefault(plot, set()).add("Не найдено")
        else:
            actual.add(plot)
            status_by_plot.setdefault(plot, set()).add("Актуально")
    not_found = {p for p in not_found if p not in actual and p not in redeemed}
    multi_status = {p for p, statuses in status_by_plot.items() if len(statuses - {"Не найдено"}) > 1}
    return {
        "Строк в result-файле по земельным участкам": data_rows,
        "Уникальных земельных участков в result-файле": len(all_plots),
        "Актуальных земельных участков, уникально": len(actual),
        "Погашенных земельных участков, уникально": len(redeemed),
        "Не найдено участков, уникально": len(not_found),
        "Участков с несколькими статусами": len(multi_status),
    }


def snt_building_totals_from_sheet(ws) -> dict[str, int]:
    """Итоги по листу «Здания» для режима СНТ."""
    headers = {normalize_text(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1) if ws.cell(1, c).value}
    plot_col = headers.get(normalize_text("Номер участка/дома"))
    status_col = headers.get(normalize_text("Статус объекта"))
    cad_col = headers.get(normalize_text("Кадастровый номер"))
    if not plot_col:
        return {"Зданий в result-файле, строк": 0, "Зданий в result-файле, уникально": 0}
    rows = 0
    unique: set[str] = set()
    actual: set[str] = set()
    redeemed: set[str] = set()
    for r in range(2, ws.max_row + 1):
        row_has_data = any(str(ws.cell(r, c).value or "").strip() for c in range(1, ws.max_column + 1))
        if not row_has_data:
            continue
        rows += 1
        plot = normalize_plot_number(ws.cell(r, plot_col).value) or str(ws.cell(r, plot_col).value or "").strip()
        cad = str(ws.cell(r, cad_col).value or "").strip() if cad_col else ""
        key = f"{plot}|{cad}" if cad else plot
        if key:
            unique.add(key)
        status = normalize_text(ws.cell(r, status_col).value) if status_col else ""
        if "погаш" in status:
            redeemed.add(key)
        elif key:
            actual.add(key)
    return {
        "Зданий в result-файле, строк": rows,
        "Зданий в result-файле, уникально": len(unique),
        "Актуальных/действующих зданий в result-файле": len(actual),
        "Погашенных зданий в result-файле": len(redeemed),
    }

class ExcelOutput:
    def __init__(self, template_path: Path, output_path: Path, log: Callable[[str], None], append_existing: bool = False, export_format: str = "burmistr"):
        self.template_path = template_path
        self.output_path = output_path
        self.log = log
        self.append_existing = append_existing
        self.export_format = str(export_format or "burmistr")
        self.log_path = build_log_output_path(output_path, "result")

        try:
            from openpyxl import Workbook, load_workbook
        except Exception as e:
            raise RuntimeError("Не установлен openpyxl. Запустите setup.bat") from e

        if append_existing and output_path.exists():
            self.wb = load_workbook(output_path)
            self.ws = self.wb.active
            self._merge_existing_log_workbook()
        elif template_path.exists():
            self.wb = load_workbook(template_path)
            self.ws = self.wb.active
        else:
            self.wb = Workbook()
            self.ws = self.wb.active
            self.ws.title = "Данные"

        self.ws.title = self.ws.title or "Данные"
        self._seen_cadastral: set[str] = set()
        self._seen_reg_keys: set[tuple[str, str]] = set()

        if self.export_format == "snt":
            # СНТ не должен подготавливать листы ОСС/Burmistr.
            # При дозаписи в существующий result-файл старые строки сохраняем.
            if "Земельные участки" in self.wb.sheetnames:
                self.snt_ws = self.wb["Земельные участки"]
                normalize_snt_land_sheet_headers(self.snt_ws)
                if not append_existing and worksheet_has_data(self.snt_ws):
                    self.snt_ws.delete_rows(2, self.snt_ws.max_row - 1)
                    normalize_snt_land_sheet_headers(self.snt_ws)
            else:
                self.snt_ws = self._make_sheet("Земельные участки", SNT_LAND_HEADERS, SNT_LAND_WIDTHS)
                normalize_snt_land_sheet_headers(self.snt_ws)
            if "Здания" in self.wb.sheetnames:
                self.snt_building_ws = self.wb["Здания"]
                if not append_existing and worksheet_has_data(self.snt_building_ws):
                    self.snt_building_ws.delete_rows(2, self.snt_building_ws.max_row - 1)
            else:
                self.snt_building_ws = None
            self._prepare_errors_sheet()
            self._prepare_extra_sheets()
        else:
            self._prepare_data_sheet()
            self._prepare_errors_sheet()
            self._prepare_building_land_sheet()
            self._prepare_extra_sheets()

    def _merge_existing_log_workbook(self) -> None:
        """При продолжении подмешивает старые служебные листы из log_result_ в текущую книгу."""
        log_path = build_log_output_path(self.output_path, "result")
        if not log_path or not Path(log_path).exists():
            return
        try:
            from openpyxl import load_workbook
            log_wb = load_workbook(log_path)
            for sheet_name in log_wb.sheetnames:
                if sheet_name in self.wb.sheetnames:
                    continue
                src = log_wb[sheet_name]
                dst = self.wb.create_sheet(sheet_name)
                for row in src.iter_rows():
                    for cell in row:
                        dst.cell(cell.row, cell.column).value = cell.value
                try:
                    dst.freeze_panes = src.freeze_panes
                except Exception:
                    pass
                for key, dim in src.column_dimensions.items():
                    try:
                        dst.column_dimensions[key].width = dim.width
                    except Exception:
                        pass
            self.log(f"Продолжение: старый log-файл подмешан: {log_path}")
        except Exception as e:
            self.log(f"Не удалось подмешать старый log-файл: {e}")

    def remove_snt_not_found_rows(self, plot_no: str) -> int:
        """Удаляет старые строки «не найдено» по участку перед перепроверкой или успешной записью."""
        if "Земельные участки" not in self.wb.sheetnames:
            return 0
        ws = self.wb["Земельные участки"]
        headers = {normalize_text(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1) if ws.cell(1, c).value}
        plot_col = headers.get(normalize_text("Номер участка"))
        cad_col = headers.get(normalize_text("Кадастровый номер"))
        status_col = headers.get(normalize_text("Статус объекта"))
        comment_col = headers.get(normalize_text("Комментарий"))
        if not plot_col:
            return 0
        target = normalize_plot_number(plot_no)
        to_delete: list[int] = []
        for r in range(2, ws.max_row + 1):
            plot = normalize_plot_number(ws.cell(r, plot_col).value)
            if plot != target:
                continue
            cad = str(ws.cell(r, cad_col).value or "").strip() if cad_col else ""
            status = normalize_text(ws.cell(r, status_col).value) if status_col else ""
            comment = normalize_text(ws.cell(r, comment_col).value) if comment_col else ""
            if "не найден" in comment or (cad in {"", "."} and status in {"", "."}):
                to_delete.append(r)
        for r in reversed(to_delete):
            ws.delete_rows(r, 1)
        return len(to_delete)

    def snt_result_totals(self) -> dict[str, int]:
        totals: dict[str, int] = {}
        if "Земельные участки" in self.wb.sheetnames:
            totals.update(snt_result_totals_from_sheet(self.wb["Земельные участки"]))
        if "Здания" in self.wb.sheetnames:
            totals.update(snt_building_totals_from_sheet(self.wb["Здания"]))
        return totals

    def _prepare_data_sheet(self) -> None:
        ws = self.ws
        # Сохраняем стили второй строки, если они были в шаблоне.
        self.row_style = []
        if ws.max_row >= 2:
            for cell in ws[2]:
                self.row_style.append(copy(cell._style))
        else:
            self.row_style = []

        # Первая строка — заголовки. Если в шаблоне не хватает ФИО-колонок, добавляем их.
        self.data_headers = list(REQUIRED_HEADERS)
        if self.export_format == "roskvartal":
            self.data_headers += [h for h in ROSKVARTAL_INTERNAL_DATA_HEADERS if h not in self.data_headers]
        existing = [ws.cell(1, col).value for col in range(1, max(ws.max_column, len(self.data_headers)) + 1)]
        existing_norm = [normalize_text(v) for v in existing]
        col = 1
        for header in self.data_headers:
            h_norm = normalize_text(header)
            if h_norm in existing_norm:
                continue
            while ws.cell(1, col).value not in (None, ""):
                col += 1
            ws.cell(1, col).value = header
            existing_norm.append(h_norm)
            col += 1

        # Перестраиваем карту заголовков после добавления.
        self.header_map: dict[str, int] = {}
        for c in range(1, ws.max_column + 1):
            value = ws.cell(1, c).value
            if value:
                self.header_map[normalize_text(value)] = c

        # Удаляем примеры/старые данные, оставляем только заголовок.
        # При продолжении с места старые строки результата сохраняем и дописываем новые.
        if ws.max_row > 1 and not self.append_existing:
            ws.delete_rows(2, ws.max_row - 1)

        # Ширины колонок — чтобы файл был читаемый.
        widths = {
            1: 18,
            2: 24,
            3: 22,
            4: 28,
            5: 42,
            6: 28,
            7: 32,
            8: 16,
            9: 34,
            10: 22,
            11: 22,
            12: 22,
            13: 22,
            14: 22,
        }
        for col_idx, width in widths.items():
            ws.column_dimensions[ws.cell(1, col_idx).column_letter].width = width

        try:
            ws.freeze_panes = "A2"
        except Exception:
            pass

    def _prepare_errors_sheet(self) -> None:
        if "Ошибки" in self.wb.sheetnames:
            self.err_ws = self.wb["Ошибки"]
            if not self.append_existing:
                self.err_ws.delete_rows(1, self.err_ws.max_row)
        else:
            self.err_ws = self.wb.create_sheet("Ошибки")
        headers = snt_log_headers(ERROR_HEADERS) if self.export_format == "snt" else ERROR_HEADERS
        for idx, header in enumerate(headers, start=1):
            self.err_ws.cell(1, idx).value = header
            self.err_ws.column_dimensions[self.err_ws.cell(1, idx).column_letter].width = 25
        from openpyxl.utils import get_column_letter
        self.err_ws.column_dimensions[get_column_letter(3)].width = 55
        self.err_ws.column_dimensions[get_column_letter(5)].width = 55
        self.err_ws.column_dimensions[get_column_letter(7)].width = 55
        try:
            self.err_ws.freeze_panes = "A2"
        except Exception:
            pass

    def _prepare_building_land_sheet(self) -> None:
        sheet_name = "Здания и участки"
        if sheet_name in self.wb.sheetnames:
            self.aux_ws = self.wb[sheet_name]
            if not self.append_existing:
                self.aux_ws.delete_rows(1, self.aux_ws.max_row)
        else:
            self.aux_ws = self.wb.create_sheet(sheet_name)
        for idx, header in enumerate(BUILDING_LAND_HEADERS, start=1):
            self.aux_ws.cell(1, idx).value = header
        widths = {
            1: 20,
            2: 22,
            3: 18,
            4: 24,
            5: 60,
            6: 18,
            7: 28,
            8: 22,
            9: 45,
        }
        for idx, width in widths.items():
            self.aux_ws.column_dimensions[self.aux_ws.cell(1, idx).column_letter].width = width
        try:
            self.aux_ws.freeze_panes = "A2"
        except Exception:
            pass

    def append_building_land(self, info: ObjectInfo, source: str = "") -> None:
        row_idx = self.aux_ws.max_row + 1
        row = [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            info.object_type or ".",
            info.status or ".",
            info.cadastral_number or ".",
            info.address or ".",
            info.area or ".",
            aux_purpose_value(info) or ".",
            info.cadastral_cost or ".",
            source or "",
        ]
        for idx, val in enumerate(row, start=1):
            cell = self.aux_ws.cell(row_idx, idx)
            cell.value = val
            if idx in {4}:
                cell.number_format = "@"

    def _make_sheet(self, name: str, headers: list[str], widths: dict[int, int] | None = None):
        if name in self.wb.sheetnames:
            ws = self.wb[name]
            if not self.append_existing:
                ws.delete_rows(1, ws.max_row)
        else:
            ws = self.wb.create_sheet(name)
        for idx, header in enumerate(headers, start=1):
            ws.cell(1, idx).value = header
        for idx, width in (widths or {}).items():
            ws.column_dimensions[ws.cell(1, idx).column_letter].width = width
        try:
            ws.freeze_panes = "A2"
        except Exception:
            pass
        return ws

    def _prepare_extra_sheets(self) -> None:
        if self.export_format == "snt":
            search_headers = snt_log_headers(SEARCH_RESULTS_HEADERS)
            rejected_headers = snt_log_headers(REJECTED_HEADERS)
            disputed_headers = snt_log_headers(DISPUTED_HEADERS)
            cad_headers = snt_log_headers(CADASTRAL_ONLY_HEADERS)
            full_headers = snt_log_headers(FULL_INFO_HEADERS)
        else:
            search_headers = SEARCH_RESULTS_HEADERS
            rejected_headers = REJECTED_HEADERS
            disputed_headers = DISPUTED_HEADERS
            cad_headers = CADASTRAL_ONLY_HEADERS
            full_headers = FULL_INFO_HEADERS
        self.search_ws = self._make_sheet("Все результаты поиска", search_headers, {2: 55, 4: 24, 5: 80, 6: 22, 7: 55})
        self.rejected_ws = self._make_sheet("Отклонённые результаты поиска", rejected_headers, {2: 55, 4: 24, 5: 80, 6: 24, 7: 55, 8: 24})
        self.disputed_ws = self._make_sheet("Спорные результаты", disputed_headers, {2: 55, 4: 24, 5: 80, 6: 24, 7: 55, 8: 24})
        self.duplicates_ws = self._make_sheet("Дубли", DUPLICATE_HEADERS, {2: 24, 3: 40, 4: 32, 5: 24, 6: 32, 7: 55})
        self.summary_ws = self._make_sheet("Сводка", SUMMARY_HEADERS, {1: 42, 2: 80})
        self.validation_ws = self._make_sheet("Проверка", VALIDATION_HEADERS, {2: 14, 3: 22, 4: 12, 5: 28, 6: 75})
        self.cad_only_ws = self._make_sheet("Кадастровые номера", cad_headers, {2: 55, 4: 24, 5: 80, 6: 24, 7: 55})
        self.full_info_ws = self._make_sheet("Все сведения", full_headers, {2: 55, 3: 18, 4: 24, 5: 65, 6: 22, 8: 38, 9: 42, 10: 90})

    def _append_row(self, ws, row: list[Any]) -> None:
        row_idx = ws.max_row + 1
        for idx, val in enumerate(row, start=1):
            ws.cell(row_idx, idx).value = val

    def append_search_result(self, search_address: str, flat_no: str, cad: str, row_text: str, decision: str, reason: str) -> None:
        self._append_row(self.search_ws, [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            search_address,
            flat_no,
            cad,
            row_text,
            decision,
            reason,
        ])

    def append_disputed(self, search_address: str, flat_no: str, cad: str, row_text: str, action: str, reason: str, chosen_cad: str = "") -> None:
        self._append_row(self.disputed_ws, [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            search_address,
            flat_no,
            cad,
            row_text,
            action,
            reason,
            chosen_cad,
        ])

    def append_rejected(self, search_address: str, flat_no: str, cad: str, row_text: str, action: str, reason: str, chosen_cad: str = "") -> None:
        self._append_row(self.rejected_ws, [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            search_address,
            flat_no,
            cad,
            row_text,
            action,
            reason,
            chosen_cad,
        ])

    def append_duplicate(self, dup_type: str, key: str, where: str, cad: str = "", reg_number: str = "", comment: str = "") -> None:
        self._append_row(self.duplicates_ws, [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            dup_type,
            key,
            where,
            cad,
            reg_number,
            comment,
        ])

    def append_cadastral_only(self, search_address: str, flat_no: str, cad: str, row_text: str, status: str, reason: str) -> None:
        self._append_row(self.cad_only_ws, [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            search_address,
            flat_no,
            cad,
            row_text,
            status,
            reason,
        ])

    def append_full_object_info(self, info: ObjectInfo, search_address: str = "") -> int:
        rows = extract_full_info_fields(info.raw_text or "")
        if not rows:
            # fallback, если сайт изменит DOM и общий разбор не найдёт пары.
            rows = [
                ("Общая информация", "Вид объекта недвижимости", info.object_type or ".", 1),
                ("Общая информация", "Статус объекта", info.status or ".", 2),
                ("Общая информация", "Кадастровый номер", info.cadastral_number or ".", 3),
                ("Характеристики объекта", "Адрес (местоположение)", info.address or ".", 4),
                ("Характеристики объекта", "Площадь, кв.м", info.area or ".", 5),
                ("Характеристики объекта", "Назначение", info.purpose or ".", 6),
                ("Характеристики объекта", "Этаж", info.floor or ".", 7),
                ("Сведения о кадастровой стоимости", "Кадастровая стоимость", info.cadastral_cost or ".", 8),
            ]

        count = 0
        for section, field, value, order in rows:
            self._append_row(self.full_info_ws, [
                datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                search_address,
                info.flat_no or ".",
                info.cadastral_number or ".",
                info.address or ".",
                info.object_type or ".",
                info.status or ".",
                section,
                field,
                value,
                order,
            ])
            count += 1
        return count

    def ensure_snt_sheet(self):
        if hasattr(self, "snt_ws") and self.snt_ws is not None:
            normalize_snt_land_sheet_headers(self.snt_ws)
            return self.snt_ws
        if "Земельные участки" in self.wb.sheetnames:
            self.snt_ws = self.wb["Земельные участки"]
            normalize_snt_land_sheet_headers(self.snt_ws)
            return self.snt_ws
        self.snt_ws = self._make_sheet("Земельные участки", SNT_LAND_HEADERS, SNT_LAND_WIDTHS)
        normalize_snt_land_sheet_headers(self.snt_ws)
        return self.snt_ws

    def ensure_snt_building_sheet(self):
        if hasattr(self, "snt_building_ws") and self.snt_building_ws is not None:
            return self.snt_building_ws
        if "Здания" in self.wb.sheetnames:
            self.snt_building_ws = self.wb["Здания"]
            return self.snt_building_ws
        self.snt_building_ws = self._make_sheet("Здания", SNT_BUILDING_HEADERS, SNT_BUILDING_WIDTHS)
        return self.snt_building_ws

    def append_snt_building(self, info: ObjectInfo, plot_no: str, search_address: str = "", comment: str = "") -> int:
        """Пишет здание/дом, найденный в СНТ, на отдельный лист result-файла."""
        ws = self.ensure_snt_building_sheet()
        raw = info.raw_text or ""
        floors = card_field(raw, "Количество этажей") or ""
        walls = card_field(raw, "Материал наружных стен") or ""
        build_year = card_field(raw, "Год завершения", "строительства") or ""
        cost_date = card_field(raw, "Дата определения") or ""
        cost_input_date = card_field(raw, "Дата внесения") or ""
        rights = info.rights or [RightInfo(".", ".", "")]
        count = 0
        for right in rights:
            self._append_row(ws, [
                datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                plot_no or info.flat_no or ".",
                info.cadastral_number or ".",
                info.status or ".",
                info.object_type or ".",
                info.address or ".",
                info.area or ".",
                info.purpose or ".",
                floors or ".",
                walls or ".",
                build_year or ".",
                info.cadastral_cost or ".",
                cost_date or "",
                cost_input_date or "",
                right.right_type or ".",
                right.reg_number or ".",
                right.reg_date or "",
                search_address,
                comment or "Здание/дом, найденный по номеру участка. Связь с земельным участком требует проверки по кадастровой карте.",
            ])
            count += 1
        return count

    def append_snt_land(self, info: ObjectInfo, plot_no: str, search_address: str = "", comment: str = "") -> int:
        """Пишет земельный участок СНТ. Если прав несколько — отдельная строка на каждое право."""
        self.remove_snt_not_found_rows(plot_no)
        ws = self.ensure_snt_sheet()
        raw = info.raw_text or ""
        category = card_field(raw, "Категория земель") or ""
        vri = card_field(raw, "Вид разрешенного", "использования") or ""
        cost_date = card_field(raw, "Дата определения") or ""
        cost_input_date = card_field(raw, "Дата внесения") or ""
        rights = info.rights or [RightInfo(".", ".", "")]
        count = 0
        for right in rights:
            self._append_row(ws, [
                datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                plot_no or info.flat_no or ".",
                info.cadastral_number or ".",
                info.status or ".",
                info.address or ".",
                info.area or ".",
                category or ".",
                vri or ".",
                info.cadastral_cost or ".",
                cost_date or "",
                cost_input_date or "",
                right.right_type or ".",
                right.reg_number or ".",
                right.reg_date or "",
                search_address,
                comment or "",
            ])
            count += 1
        return count

    def append_snt_not_found(self, plot_no: str, search_address: str, comment: str) -> None:
        ws = self.ensure_snt_sheet()
        self._append_row(ws, [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            plot_no,
            ".",
            ".",
            ".",
            ".",
            ".",
            ".",
            ".",
            "",
            "",
            ".",
            ".",
            "",
            search_address,
            comment,
        ])

    def write_summary(self, summary: dict[str, Any]) -> None:
        self.summary_ws.delete_rows(1, self.summary_ws.max_row)
        for idx, header in enumerate(SUMMARY_HEADERS, start=1):
            self.summary_ws.cell(1, idx).value = header
        for key, value in summary.items():
            self._append_row(self.summary_ws, [key, value])
        try:
            self.summary_ws.freeze_panes = "A2"
        except Exception:
            pass

    def append_validation(self, level: str, sheet: str, row_idx: Any, field: str, message: str) -> None:
        self._append_row(self.validation_ws, [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            level,
            sheet,
            row_idx,
            field,
            message,
        ])

    def validate_result(self) -> None:
        # Очищаем старые результаты проверки, оставляя заголовок.
        self.validation_ws.delete_rows(2, self.validation_ws.max_row)
        headers = {normalize_text(self.ws.cell(1, c).value): c for c in range(1, self.ws.max_column + 1) if self.ws.cell(1, c).value}

        def get(row_idx: int, header: str) -> Any:
            col = headers.get(normalize_text(header))
            return self.ws.cell(row_idx, col).value if col else ""

        validation_seen_reg_keys: set[tuple[str, str]] = set()
        flat_groups: dict[str, list[tuple[int, Any, str]]] = {}
        reg_to_people: dict[str, set[str]] = {}

        for row_idx in range(2, self.ws.max_row + 1):
            cad = str(get(row_idx, "Кадастровый номер") or "").strip()
            flat = str(get(row_idx, "Номер помещения (обяз.)") or "").strip()
            area = str(get(row_idx, "Площадь помещения, м2 (обяз.)") or "").strip()
            reg = str(get(row_idx, "Номер регистрации (обяз.)") or "").strip()
            reg_date = str(get(row_idx, "Дата регистрации") or "").strip()
            joint = str(get(row_idx, 'Признак совместной собс. ("да" для сов. собс.)') or "").strip()
            share = str(get(row_idx, "Доля (1/1, 1/2, 1/3 и т.д.) (обяз., сумма в помещении должна быть 1/1)") or "").strip()
            last = str(get(row_idx, "Фамилия / Наименование организации (обяз.)") or "").strip()
            first = str(get(row_idx, "Имя (обяз. для физ лиц)") or "").strip()
            patr = str(get(row_idx, "Отчество") or "").strip()

            if not cad or cad == ".":
                self.append_validation("Ошибка", "Данные", row_idx, "Кадастровый номер", "Пустой кадастровый номер")
            if not flat or flat == ".":
                self.append_validation("Предупреждение", "Данные", row_idx, "Номер помещения", "Пустой номер помещения")
            if not area or area == ".":
                self.append_validation("Предупреждение", "Данные", row_idx, "Площадь", "Пустая площадь")
            if reg == "." and reg_date:
                self.append_validation("Ошибка", "Данные", row_idx, "Дата регистрации", "Если номер регистрации '.', дата должна быть пустой")
            if reg and reg != "." and not reg_date and "договор" not in normalize_text(reg):
                self.append_validation("Предупреждение", "Данные", row_idx, "Дата регистрации", "Есть номер регистрации, но дата регистрации пустая")
            if joint == "да" and share != "1/1":
                self.append_validation("Ошибка", "Данные", row_idx, "Доля", "Для совместной собственности доля должна быть 1/1")

            frac = parse_fraction_share(share)
            if share and share != "." and frac is None:
                self.append_validation("Ошибка", "Данные", row_idx, "Доля", f"Некорректный формат доли: {share}")
            if flat and flat != ".":
                flat_groups.setdefault(flat.upper(), []).append((row_idx, frac, share))

            if cad and cad != "." and reg:
                key = (cad, reg)
                if key in validation_seen_reg_keys:
                    self.append_duplicate("Кадастровый номер + регистрация", f"{cad}|{reg}", f"Данные!{row_idx}", cad, reg, "Повтор в основном листе")
                validation_seen_reg_keys.add(key)

            if reg and reg != ".":
                person_key = normalize_text(" ".join([last, first, patr])).strip()
                if person_key and person_key != ".":
                    reg_to_people.setdefault(normalize_reg_key(reg) or reg, set()).add(person_key)

        try:
            from fractions import Fraction
            for flat, items in flat_groups.items():
                if not items or any(frac is None for _r, frac, raw in items if raw and raw != "."):
                    continue
                fractions = [frac for _r, frac, _raw in items if frac is not None]
                if not fractions:
                    continue
                total = sum(fractions, Fraction(0, 1))
                if total != Fraction(1, 1):
                    rows = ", ".join(str(r) for r, _f, _raw in items)
                    shares = ", ".join(str(raw) for _r, _f, raw in items)
                    self.append_validation("Ошибка", "Данные", rows, "Доля", f"Сумма долей по помещению {flat} = {total}, не равна 1/1. Доли: {shares}")
        except Exception as e:
            self.append_validation("Предупреждение", "Данные", "", "Доля", f"Не удалось проверить суммы долей: {e}")

        for reg_key, people in reg_to_people.items():
            if len(people) > 1:
                self.append_validation("Предупреждение", "Данные", "", "Номер регистрации", f"Один номер права встречается у разных ФИО: {reg_key}; {', '.join(sorted(people)[:5])}")

        # Земля/здания не должны попадать в основной лист. Проверяем хотя бы дубли между основным листом и листом "Здания и участки".
        data_cads = set()
        cad_col = headers.get(normalize_text("Кадастровый номер"))
        if cad_col:
            for row_idx in range(2, self.ws.max_row + 1):
                cad = str(self.ws.cell(row_idx, cad_col).value or "").strip()
                if cad and cad != ".":
                    data_cads.add(cad)

        aux_cads = set()
        for row_idx in range(2, self.aux_ws.max_row + 1):
            cad = str(self.aux_ws.cell(row_idx, 4).value or "").strip()
            if cad and cad != ".":
                if cad in aux_cads:
                    self.append_duplicate("Кадастровый номер", cad, f"Здания и участки!{row_idx}", cad, "", "Повтор на листе зданий/участков")
                if cad in data_cads:
                    self.append_duplicate("КН в основном и вспомогательном листе", cad, f"Здания и участки!{row_idx}", cad, "", "Один КН есть и в основном листе, и в зданиях/участках")
                aux_cads.add(cad)

    def append_data(self, values: dict[str, Any]) -> None:
        row_idx = self.ws.max_row + 1
        cad_for_dupe = str(values.get("Кадастровый номер", "") or "").strip()
        reg_for_dupe = str(values.get("Номер регистрации (обяз.)", "") or "").strip()
        # Один кадастровый номер может законно повторяться по нескольким правам/долям.
        # Дублем считаем только повтор той же пары КН + номер регистрации.
        if cad_for_dupe and reg_for_dupe:
            key = (cad_for_dupe, reg_for_dupe)
            if key in self._seen_reg_keys:
                self.append_duplicate("Кадастровый номер + регистрация", f"{cad_for_dupe}|{reg_for_dupe}", f"Данные!{row_idx}", cad_for_dupe, reg_for_dupe, "Пара КН+регистрация уже встречалась раньше")
            self._seen_reg_keys.add(key)

        headers_to_write = getattr(self, "data_headers", REQUIRED_HEADERS)
        for header in headers_to_write:
            col_idx = self.header_map.get(normalize_text(header))
            if not col_idx:
                continue
            val = values.get(header, "")
            cell = self.ws.cell(row_idx, col_idx)
            cell.value = val
            if self.row_style and col_idx - 1 < len(self.row_style):
                cell._style = copy(self.row_style[col_idx - 1])
            # Номера регистрации, кадастровые номера и даты прав держим как текст.
            if header in {"Кадастровый номер", "Номер регистрации (обяз.)", "Дата регистрации", "ИНН (обяз. для юр лиц)", "ОГРН (ОГРНИП)", "Снилс"}:
                cell.number_format = "@"

    def append_error(self, flat_no: str, search_address: str, status: str, reason: str, cadastral: str = "", rr_address: str = "") -> None:
        row_idx = self.err_ws.max_row + 1
        row = [
            datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
            flat_no,
            search_address,
            status,
            reason,
            cadastral,
            rr_address,
        ]
        for idx, val in enumerate(row, start=1):
            self.err_ws.cell(row_idx, idx).value = val

    def save(self) -> None:
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        if self.export_format == "roskvartal":
            build_roskvartal_sheet_from_data(self.wb)
            result_sheets = {"Roskvartal"}
        elif self.export_format == "snt":
            result_sheets = {"Земельные участки", "Здания"}
        else:
            result_sheets = RESULT_SHEETS
        self.log_path = save_result_and_log_workbooks(self.wb, self.output_path, result_sheets=result_sheets, log_kind="result")


# ----------------------------- Playwright worker -----------------------------


class ParserWorker(threading.Thread):
    def __init__(self, cfg: dict[str, Any], log_queue: queue.Queue[str], ready_event: threading.Event, pause_event: threading.Event, stop_event: threading.Event):
        super().__init__(daemon=True)
        self.cfg = cfg
        self.log_queue = log_queue
        self.ready_event = ready_event
        self.pause_event = pause_event
        self.stop_event = stop_event
        self.playwright = None
        self.context = None
        self.page = None
        self.excel: ExcelOutput | None = None
        self.next_task_event = threading.Event()
        self.idle = False
        self.completed_task_id = 0
        self.metrics: dict[str, int] = {}
        self.run_started_at = ""

    def log(self, msg: str) -> None:
        self.log_queue.put(msg)

    def parser_mode(self) -> str:
        mode = str(self.cfg.get("parser_mode") or "").strip()
        if not mode:
            mode = "cadastral" if bool(self.cfg.get("collect_cadastral_only", False)) else "oss"
        return mode

    def is_snt_mode(self) -> bool:
        return self.parser_mode() == "snt"

    def is_cadastral_mode(self) -> bool:
        return self.parser_mode() == "cadastral"

    def is_full_info_mode(self) -> bool:
        return self.parser_mode() == "full"

    def wait_control(self) -> bool:
        while not self.stop_event.is_set() and not self.pause_event.is_set():
            time.sleep(0.2)
        return not self.stop_event.is_set()

    def wait_manual_continue(self, reason: str) -> bool:
        self.ready_event.clear()
        self.log(reason)
        self.log("Нажмите кнопку «Продолжить после входа/капчи», когда можно идти дальше.")
        while not self.stop_event.is_set():
            if self.ready_event.wait(timeout=0.5):
                if self.stop_event.is_set():
                    raise StopRequested("Остановка во время ожидания входа/капчи")
                return True
        raise StopRequested("Остановка во время ожидания входа/капчи")

    def human_delay(self) -> None:
        lo = float(self.cfg.get("pause_min_seconds", 3) or 0)
        hi = float(self.cfg.get("pause_max_seconds", 7) or lo)
        if hi < lo:
            hi = lo
        delay = random.uniform(lo, hi)
        self.controlled_sleep(delay)

    def controlled_sleep(self, seconds: float, *, log_every: int = 30, reason: str = "") -> bool:
        """Сон, который уважает Пауза/Стоп. Возвращает False при Stop."""
        end = time.time() + max(0.0, float(seconds))
        next_log = time.time() + log_every
        while time.time() < end:
            if not self.wait_control():
                return False
            now = time.time()
            if reason and now >= next_log:
                left = int(max(0, end - now))
                self.log(f"  {reason}: осталось примерно {left} сек.")
                next_log = now + log_every
            time.sleep(0.2)
        return not self.stop_event.is_set()

    def build_progress_path(self, base: Path, address_prefix: str, start_flat: int, end_flat: int) -> Path:
        if self.parser_mode() == "snt":
            city = str(self.cfg.get("snt_city") or "").strip()
            name = str(self.cfg.get("snt_name") or "").strip()
            quarters = "_".join(parse_cadastral_quarters(self.cfg.get("snt_cadastral_quarters") or "")) or "all_quarters"
            slug = filename_slug(f"snt_{city}_{name}_{quarters}") or "snt"
            return base / "state" / f"progress_{slug}.json"
        house = house_address_from_prefix(address_prefix) or address_prefix or "house"
        slug = filename_slug(f"{house}_kv_{start_flat}-{end_flat}") or "progress"
        return base / "state" / f"progress_{slug}.json"

    def load_progress_state(self, path: Path) -> dict[str, Any]:
        if not path.exists():
            return {}
        try:
            with path.open("r", encoding="utf-8") as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}
        except Exception as e:
            self.log(f"Не удалось прочитать прогресс {path}: {e}")
            return {}

    def save_progress_state(self) -> None:
        path = getattr(self, "progress_path", None)
        if not path:
            return
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            state = getattr(self, "progress_state", {}) or {}
            state["updated_at"] = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            state["done_units"] = sorted(getattr(self, "done_units", set()))
            state["done_cadastrals"] = sorted(getattr(self, "processed_cadastrals", set()))
            if self.excel:
                state["output_path"] = str(self.excel.output_path)
            with path.open("w", encoding="utf-8") as f:
                json.dump(state, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.log(f"Не удалось сохранить прогресс: {e}")

    def mark_unit_done(self, unit_no: str, status: str = "done") -> None:
        if not bool(self.cfg.get("resume_enabled", True)):
            return
        key = clean_unit_number(unit_no) or str(unit_no)
        if not hasattr(self, "done_units"):
            self.done_units = set()
        self.done_units.add(key.upper())
        state = getattr(self, "progress_state", {}) or {}
        state["last_unit"] = key
        state["last_status"] = status
        self.progress_state = state
        self.save_progress_state()

    def run(self) -> None:
        """
        Рабочий поток живёт между задачами, если браузер оставлен открытым.
        Это позволяет запустить следующий дом/объект без повторной авторизации:
        пользователь меняет настройки и снова нажимает «Старт».
        """
        while not self.stop_event.is_set():
            self.idle = False
            try:
                self._run_inner()
            except Exception:
                self.log("Критическая ошибка:\n" + traceback.format_exc())
            finally:
                try:
                    if self.excel:
                        self.excel.save()
                        self.log(f"Excel сохранён: {self.excel.output_path}")
                        self.log(f"Лог Excel сохранён: {self.excel.log_path}") if self.excel.log_path else None
                except Exception as e:
                    self.log(f"Не удалось сохранить Excel: {e}")
                self.excel = None

            self.completed_task_id += 1

            if self.stop_event.is_set():
                break

            if self.cfg.get("close_browser_after_finish", False):
                break

            self.idle = True
            self.log(
                "Задача завершена. Браузер пока оставлен открытым. "
                "Можно изменить настройки и нажать «Старт» — следующая задача пойдёт в этом же браузере без новой авторизации. "
                "Чтобы закрыть браузер, нажмите «Стоп» или выберите закрытие в появившемся вопросе."
            )

            # Ждём либо следующую задачу, либо команду Стоп.
            while not self.stop_event.is_set():
                if self.next_task_event.wait(timeout=0.5):
                    self.next_task_event.clear()
                    break

            if self.stop_event.is_set():
                break

            self.log("Запускаю следующую задачу в уже открытом браузере.")
            # цикл продолжится и вызовет _run_inner() с тем же self.context/page

        self.idle = False
        self.close_browser()
        self.log("Работа завершена.")

    def request_next_task(self, cfg: dict[str, Any]) -> None:
        self.cfg = cfg.copy()
        self.next_task_event.set()

    def close_browser(self) -> None:
        try:
            if self.context:
                self.context.close()
        except Exception:
            pass
        self.context = None
        self.page = None
        try:
            if self.playwright:
                self.playwright.stop()
        except Exception:
            pass
        self.playwright = None

    def _run_inner(self) -> None:
        try:
            from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
            from playwright.sync_api import sync_playwright
        except Exception as e:
            raise RuntimeError("Не установлен Playwright. Запустите setup.bat") from e

        base = app_dir()
        template_path = Path(self.cfg.get("template_path") or "template_burmistr.xlsx")
        if self.parser_mode() == "snt" and template_path.name.lower() in {"template.xlsx", "template_burmistr.xlsx", ""}:
            template_path = Path("template_snt.xlsx")
        elif self.parser_mode() == "oss" and str(self.cfg.get("oss_export_format") or "burmistr") == "burmistr" and template_path.name.lower() in {"template.xlsx", ""}:
            template_path = Path("template_burmistr.xlsx")
        elif self.parser_mode() == "oss" and str(self.cfg.get("oss_export_format") or "burmistr") == "roskvartal" and template_path.name.lower() in {"template.xlsx", "template_burmistr.xlsx", ""}:
            template_path = Path("template_roskvartal.xlsx")
        if not template_path.is_absolute():
            # Сначала рядом с программой, потом как ресурс PyInstaller.
            p1 = base / template_path
            p2 = resource_path(str(template_path))
            template_path = p1 if p1.exists() else p2
        output_path = Path(self.cfg.get("output_path") or "output/result.xlsx")
        if not output_path.is_absolute():
            output_path = base / output_path

        start_flat = int(self.cfg.get("start_flat") or 1)
        end_flat = int(self.cfg.get("end_flat") or start_flat)
        address_prefix = str(self.cfg.get("address_prefix") or "").strip()
        if self.parser_mode() == "snt":
            address_prefix = f"{self.cfg.get('snt_city') or ''} {self.cfg.get('snt_name') or ''}".strip() or address_prefix

        resume_enabled = bool(self.cfg.get("resume_enabled", True))
        self.progress_path = self.build_progress_path(base, address_prefix, start_flat, end_flat)
        self.progress_state = self.load_progress_state(self.progress_path) if resume_enabled else {}
        append_existing = False

        if resume_enabled and self.progress_state.get("output_path") and Path(str(self.progress_state.get("output_path"))).exists():
            output_path = Path(str(self.progress_state.get("output_path")))
            append_existing = True
            self.log(f"Продолжение с места: использую прежний Excel: {output_path}")
        elif bool(self.cfg.get("auto_output_filename", True)):
            output_path = build_auto_output_path(output_path, address_prefix, start_flat, end_flat)
            self.log(f"Автоимя итогового файла: {output_path}")

        if self.parser_mode() == "snt" and bool(self.cfg.get("snt_append_existing", True)) and output_path.exists() and not append_existing and not bool(self.cfg.get("auto_output_filename", True)):
            append_existing = True
            self.log(f"СНТ: дозапись в существующий Excel: {output_path}")

        if self.parser_mode() == "snt":
            export_format = "snt"
        elif self.parser_mode() == "oss":
            export_format = str(self.cfg.get("oss_export_format") or "burmistr")
        else:
            export_format = "burmistr"
        self.excel = ExcelOutput(template_path, output_path, self.log, append_existing=append_existing, export_format=export_format)
        self.run_started_at = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        self.metrics = {
            "Заданий всего": 0,
            "Участков в диапазоне": 0,
            "Пропущено как уже обработанные": 0,
            "Найдено помещений/квартир": 0,
            "Не найдено": 0,
            "Ошибок": 0,
            "Погашенных": 0,
            "Лимитов Росреестра": 0,
            "Зданий/участков": 0,
            "Кадастровых номеров без открытия карточек": 0,
            "Спорных результатов": 0,
            "Отклонённых строк поиска": 0,
            "Найдено зданий СНТ": 0,
        }

        self.done_units: set[str] = set(str(x).upper() for x in (self.progress_state.get("done_units", []) if resume_enabled else []))
        self.processed_cadastrals: set[str] = set(str(x) for x in (self.progress_state.get("done_cadastrals", []) if resume_enabled else []))
        if resume_enabled:
            self.progress_state.update({
                "version": APP_TITLE,
                "address_prefix": address_prefix,
                "start_flat": start_flat,
                "end_flat": end_flat,
                "output_path": str(output_path),
                "started_at": self.progress_state.get("started_at") or self.run_started_at,
            })
            self.save_progress_state()
            if self.done_units:
                self.log(f"Продолжение с места включено. Уже обработано номеров: {len(self.done_units)}. Файл прогресса: {self.progress_path}")

        profile_dir = Path(self.cfg.get("browser_profile_dir") or "profile_rosreestr")
        if not profile_dir.is_absolute():
            profile_dir = base / profile_dir
        profile_dir.mkdir(parents=True, exist_ok=True)

        if self.context is None or self.page is None or self.playwright is None:
            self.log("Открываю браузер. Авторизуйтесь вручную, если сайт попросит вход через Госуслуги/Росреестр.")
            self.playwright = sync_playwright().start()

            launch_kwargs = {
                "headless": False,
                "slow_mo": int(self.cfg.get("slow_mo_ms") or 0),
                "args": ["--disable-blink-features=AutomationControlled"],
            }

            channel = (self.cfg.get("browser_channel") or "").strip()
            try:
                if channel:
                    self.context = self.playwright.chromium.launch_persistent_context(str(profile_dir), channel=channel, **launch_kwargs)
                else:
                    self.context = self.playwright.chromium.launch_persistent_context(str(profile_dir), **launch_kwargs)
            except Exception as e:
                self.log(f"Не удалось открыть Chrome channel={channel!r}: {e}")
                self.log("Пробую открыть Playwright Chromium. Если ошибка повторится — запустите setup.bat.")
                self.context = self.playwright.chromium.launch_persistent_context(str(profile_dir), **launch_kwargs)

            self.page = self.context.pages[0] if self.context.pages else self.context.new_page()
            self.page.set_default_timeout(30000)
            self.page.goto(self.cfg.get("login_url"), wait_until="domcontentloaded", timeout=60000)

            if not self.wait_manual_continue("После успешной авторизации откройте/дождитесь личного кабинета и нажмите продолжение в этой программе."):
                return
        else:
            self.log("Использую уже открытый браузер и текущую авторизованную сессию.")
            try:
                self.page.set_default_timeout(30000)
            except Exception:
                pass

        self.open_service_page()

        if self.is_snt_mode():
            self.process_snt_mode(start_flat, end_flat)
            if self.excel:
                try:
                    self.write_run_summary(address_prefix, start_flat, end_flat)
                except Exception as e:
                    self.log(f"Не удалось выполнить сводку СНТ: {e}")
                self.excel.save()
                self.log(f"Готово: {self.excel.output_path}")
                self.log(f"Лог-файл: {self.excel.log_path}") if self.excel.log_path else None
            return

        search_numbered = bool(self.cfg.get("search_numbered_units", True))
        try_variants = bool(self.cfg.get("try_address_variants", True))

        manual_numbers = str(self.cfg.get("manual_unit_numbers") or "")
        manual_file_path = str(self.cfg.get("manual_unit_file_path") or "")
        extra_numbers = str(self.cfg.get("extra_unit_numbers") or "")
        unit_queue = build_full_unit_queue(start_flat, end_flat, extra_numbers, manual_numbers, manual_file_path, include_range=search_numbered)

        if unit_queue:
            if search_numbered:
                self.log(f"Этап 1: поиск квартир/помещений. Очередь: {len(unit_queue)} номеров.")
            else:
                self.log(f"Этап 1: диапазон выключен, используется ручная очередь. Очередь: {len(unit_queue)} номеров.")
            if extra_numbers.strip() and search_numbered:
                self.log("  Дополнительные номера: " + ", ".join(parse_extra_unit_numbers(extra_numbers)))
            manual_from_text = parse_extra_unit_numbers(manual_numbers)
            if manual_from_text:
                self.log("  Ручная очередь из поля: " + ", ".join(manual_from_text))
            if manual_file_path.strip():
                try:
                    manual_from_file = read_manual_units_from_file(manual_file_path)
                    self.log(f"  Ручная очередь из файла: {len(manual_from_file)} номеров — " + ", ".join(manual_from_file[:30]) + (" ..." if len(manual_from_file) > 30 else ""))
                except Exception as e:
                    self.log(f"  Не удалось прочитать файл ручной очереди: {e}")
            for idx, flat in enumerate(unit_queue, start=1):
                    if not self.wait_control():
                        break
                    flat_key = (clean_unit_number(flat) or str(flat)).upper()
                    if bool(self.cfg.get("resume_enabled", True)) and bool(self.cfg.get("skip_done_units", True)) and flat_key in getattr(self, "done_units", set()):
                        self.log(f"[{idx}/{len(unit_queue)}] Пропуск помещения {flat}: уже обработано по файлу прогресса.")
                        continue
                    search_addresses = build_search_addresses(address_prefix, str(flat), enabled=try_variants)
                    primary_address = search_addresses[0] if search_addresses else f"{address_prefix} {flat}".strip()
                    self.log(f"[{idx}/{len(unit_queue)}] Ищу помещение {flat}: {primary_address}")
                    if try_variants and len(search_addresses) > 1:
                        self.log("  Варианты поиска: " + " | ".join(search_addresses))
                    try:
                        self.metrics["Заданий всего"] = self.metrics.get("Заданий всего", 0) + 1
                        self.process_flat(str(flat), search_addresses)
                    except Exception:
                        self.log(f"Ошибка по помещению {flat}:\n{traceback.format_exc()}")
                        if self.excel:
                            self.excel.append_error(str(flat), primary_address, "Ошибка", "Исключение при обработке. Подробности в логе.")
                    if not self.stop_event.is_set():
                        self.mark_unit_done(str(flat))
                    if self.excel and self.cfg.get("save_after_each_flat", True):
                        self.excel.save()
                    self.human_delay()
        else:
            if search_numbered:
                self.log("Очередь квартир/помещений пустая. Этап квартир пропущен.")
            else:
                self.log("Этап квартир/помещений по номерам выключен, ручная очередь пустая.")

        if bool(self.cfg.get("include_unnumbered_house_objects", False)) and not self.stop_event.is_set():
            if bool(self.cfg.get("resume_enabled", True)) and "__UNNUMBERED__" in getattr(self, "done_units", set()):
                self.log("Этап 2 пропущен: поиск помещений без номеров уже был выполнен по файлу прогресса.")
            else:
                self.log("Этап 2: поиск помещений без номеров по голому адресу дома.")
                self.process_unnumbered_house_objects()
                if not self.stop_event.is_set():
                    self.mark_unit_done("__UNNUMBERED__")
        elif not self.stop_event.is_set():
            self.log("Этап поиска помещений без номеров выключен.")

        if self.excel:
            try:
                self.excel.validate_result()
                self.write_run_summary(address_prefix, start_flat, end_flat)
            except Exception as e:
                self.log(f"Не удалось выполнить проверку/сводку: {e}")
            self.excel.save()
            self.log(f"Готово: {self.excel.output_path}")
            self.log(f"Лог-файл: {self.excel.log_path}") if self.excel.log_path else None

    def write_run_summary(self, address_prefix: str, start_flat: int, end_flat: int) -> None:
        if not self.excel:
            return
        if self.is_snt_mode():
            summary = {
                "Версия программы": APP_TITLE,
                "Дата/время начала": self.run_started_at,
                "Дата/время окончания": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                "Режим парсинга": "Реестр земельных участков / СНТ",
                "Город / населённый пункт": str(self.cfg.get("snt_city") or "").strip(),
                "Название товарищества": str(self.cfg.get("snt_name") or "").strip(),
                "Кадастровые кварталы": str(self.cfg.get("snt_cadastral_quarters") or "").strip(),
                "Диапазон участков": f"{start_flat}–{end_flat}",
                "Дополнительные участки": str(self.cfg.get("extra_unit_numbers") or "").strip(),
                "Ручная очередь участков": str(self.cfg.get("manual_unit_numbers") or "").strip(),
                "Файл ручной очереди": str(self.cfg.get("manual_unit_file_path") or "").strip(),
                "Перебирать участки по диапазону": "да" if self.cfg.get("search_numbered_units") else "нет",
                "Пробовать варианты адреса": "да" if self.cfg.get("snt_try_address_variants") else "нет",
                "Записывать актуальные участки": "да" if self.cfg.get("snt_include_actual") else "нет",
                "Записывать погашенные участки": "да" if self.cfg.get("snt_include_redeemed") else "нет",
                "Записывать найденные здания": "да" if self.cfg.get("snt_write_buildings", True) else "нет",
                "Добавлять не найденные участки": "да" if self.cfg.get("snt_write_not_found") else "нет",
                "Пресет пауз": self.cfg.get("pause_preset", ""),
                "=== Текущий запуск ===": "",
                "Участков в диапазоне": self.metrics.get("Участков в диапазоне", self.metrics.get("Заданий всего", 0)),
                "Пропущено как уже обработанные": self.metrics.get("Пропущено как уже обработанные", 0),
                "Обработано в текущем запуске": max(0, int(self.metrics.get("Участков в диапазоне", self.metrics.get("Заданий всего", 0)) or 0) - int(self.metrics.get("Пропущено как уже обработанные", 0) or 0)),
                "Найдено актуальных/действующих участков за запуск": self.metrics.get("Найдено земельных участков СНТ", 0),
                "Найдено погашенных участков за запуск": self.metrics.get("Погашенных", 0),
                "Не найдено участков за запуск": self.metrics.get("Не найдено", 0),
                "Найдено зданий за запуск": self.metrics.get("Найдено зданий СНТ", 0),
                "Отклонённых строк поиска за запуск": self.metrics.get("Отклонённых строк поиска", 0),
                "Спорных участков за запуск": self.metrics.get("Спорных результатов", 0),
                "Ошибок за запуск": self.metrics.get("Ошибок", 0),
                "Лимитов Росреестра за запуск": self.metrics.get("Лимитов Росреестра", 0),
                "=== Весь result-файл ===": "",
            }
            try:
                summary.update(self.excel.snt_result_totals() if self.excel else {})
            except Exception as e:
                summary["Итог result-файла"] = f"Не удалось посчитать: {e}"
        else:
            summary = {
                "Версия программы": APP_TITLE,
                "Дата/время начала": self.run_started_at,
                "Дата/время окончания": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
                "Адрес до номера": address_prefix,
                "Режим парсинга": PARSER_MODE_LABELS.get(str(self.cfg.get("parser_mode") or "oss"), str(self.cfg.get("parser_mode") or "oss")),
                "Формат выгрузки ОСС": OSS_EXPORT_FORMAT_LABELS.get(str(self.cfg.get("oss_export_format") or "burmistr"), str(self.cfg.get("oss_export_format") or "burmistr")),
                "Диапазон": f"{start_flat}–{end_flat}",
                "Дополнительные номера": str(self.cfg.get("extra_unit_numbers") or "").strip(),
                "Ручная очередь": str(self.cfg.get("manual_unit_numbers") or "").strip(),
                "Файл ручной очереди": str(self.cfg.get("manual_unit_file_path") or "").strip(),
                "Искать квартиры/помещения по диапазону": "да" if self.cfg.get("search_numbered_units") else "нет",
                "Поиск помещений без номеров": "да" if self.cfg.get("include_unnumbered_house_objects") else "нет",
                "Только собрать кадастровые номера": "да" if self.cfg.get("collect_cadastral_only") else "нет",
                "Пресет пауз": self.cfg.get("pause_preset", ""),
            }
            summary.update(self.metrics)
        self.excel.write_summary(summary)

    def detect_captcha(self) -> bool:
        try:
            text = self.page.inner_text("body", timeout=3000) if self.page else ""
        except Exception:
            return False
        return any(word in normalize_text(text) for word in ["капча", "captcha", "введите символы", "проверочный код"])

    def detect_rate_limit(self) -> bool:
        """Росреестр иногда показывает toast: превышен лимит обращений."""
        try:
            text = self.page.inner_text("body", timeout=3000) if self.page else ""
        except Exception:
            return False
        t = normalize_text(text)
        return (
            "превышен лимит обращений" in t
            or "попробуйте позже" in t and "список объектов" in t
            or "не удалось получить список объектов недвижимости" in t
        )

    def is_login_page(self) -> bool:
        """Понимаем, что сайт вернул на авторизацию. Это не ошибка: ждём пользователя."""
        assert self.page is not None
        url = (self.page.url or "").lower()
        if "/login" in url or "esia" in url or "gosuslugi" in url:
            return True
        try:
            text = normalize_text(self.page.inner_text("body", timeout=3000))
        except Exception:
            text = ""
        login_markers = [
            "войти",
            "авторизац",
            "госуслуг",
            "единая система идентификации",
            "логин",
            "пароль",
        ]
        service_markers = [
            "справочная информация по объектам недвижимости",
            "адрес или кадастровый номер",
            "выберите тип поиска",
        ]
        if any(m in text for m in service_markers):
            return False
        return any(m in text for m in login_markers)

    def service_page_ready(self) -> bool:
        assert self.page is not None
        try:
            text = normalize_text(self.page.inner_text("body", timeout=5000))
        except Exception:
            return False
        return (
            "справочная информация по объектам недвижимости" in text
            or "адрес или кадастровый номер" in text
            or "выберите тип поиска" in text
        )

    def safe_goto(self, url: str, timeout: int = 60000) -> None:
        """Переход без падения на типовом редиректе Росреестра service -> login."""
        assert self.page is not None
        try:
            self.page.goto(url, wait_until="domcontentloaded", timeout=timeout)
        except Exception as e:
            msg = str(e)
            if "interrupted by another navigation" in msg or "NS_BINDING_ABORTED" in msg:
                self.log("  Переход был прерван редиректом сайта, продолжаю проверку текущей страницы.")
                try:
                    self.page.wait_for_load_state("domcontentloaded", timeout=15000)
                except Exception:
                    pass
            else:
                raise
        try:
            self.page.wait_for_load_state("networkidle", timeout=15000)
        except Exception:
            pass

    def open_service_page(self) -> None:
        assert self.page is not None
        service_url = self.cfg.get("service_url")

        # Росреестр часто редиректит service_url на /login, если авторизация ещё не завершена
        # или если ЕСИА открыла промежуточную страницу. Это не критическая ошибка — даём
        # пользователю вручную завершить вход и нажать «Продолжить».
        for step in range(1, 6):
            self.safe_goto(service_url, timeout=60000)

            if self.detect_captcha():
                if not self.wait_manual_continue("Похоже, появилась капча. Введите её вручную в браузере."):
                    return
                continue

            if self.service_page_ready():
                return

            if self.is_login_page():
                if not self.wait_manual_continue(
                    "Сайт вернул на страницу входа. Завершите авторизацию в браузере и нажмите «Продолжить»."
                ):
                    return
                continue

            # Иногда после входа ЛК открывает главную страницу, а не нужный сервис.
            # Повторяем прямой переход несколько раз; если не получится — просим открыть сервис вручную.
            if step < 5:
                self.log("  Сервисная страница пока не открылась, повторяю переход.")
                time.sleep(2)
                continue

            if not self.wait_manual_continue(
                "Не удалось автоматически открыть страницу сервиса. Откройте её вручную в браузере и нажмите «Продолжить»."
            ):
                return
            if self.service_page_ready():
                return

        if not self.service_page_ready():
            raise RuntimeError("Не удалось открыть страницу сервиса справочной информации")

    def process_unnumbered_house_objects(self) -> None:
        """Дополнительный проход: поиск по голому адресу дома. Помещения без номера идут в основной лист, здания/участки — в отдельный лист."""
        assert self.page is not None
        assert self.excel is not None

        address_prefix = str(self.cfg.get("address_prefix") or "").strip()
        house_address = house_address_from_prefix(address_prefix)
        if not house_address:
            self.log("Поиск помещений без номеров пропущен: не задан адрес дома.")
            return

        self.log(f"Поиск помещений без номеров по голому адресу: {house_address}")
        max_pages = int(self.cfg.get("unnumbered_max_pages") or 10)
        rate_limit_retry_count = int(self.cfg.get("rate_limit_retry_count") or 12)
        rate_limit_hits = 0

        while True:
            try:
                self.open_service_page()
                self.fill_search_address(house_address)
                self.click_find()
                self.wait_results()
                break
            except RateLimitError as e:
                rate_limit_hits += 1
                if rate_limit_hits > rate_limit_retry_count:
                    self.excel.append_error(".", house_address, "Лимит", f"Лимит обращений при поиске помещений без номеров: {e}")
                    return
                pause = float(self.cfg.get("rate_limit_pause_seconds") or 300)
                self.log(f"  Лимит при голом адресе. Жду {int(pause)} сек. ({rate_limit_hits}/{rate_limit_retry_count}).")
                if not self.controlled_sleep(pause, reason="Ожидание лимита Росреестра"):
                    return
            except Exception as e:
                self.excel.append_error(".", house_address, "Ошибка", f"Не удалось выполнить поиск по голому адресу: {e}")
                return

        total_added = 0
        page_no = 1
        while page_no <= max_pages and not self.stop_event.is_set():
            if not self.wait_control():
                return
            try:
                candidates = self.collect_house_candidates(house_address)
                self.log(f"  Голый адрес, страница {page_no}: кандидатов без 'кв' найдено {len(candidates)}")
                for cand in candidates:
                    if not self.wait_control():
                        return
                    cad = cand.get("cad") or ""
                    if cad and cad in getattr(self, "processed_cadastrals", set()):
                        continue
                    info = self.open_candidate_and_parse(cand, ".", house_address)
                    if not info:
                        self.excel.append_error(".", house_address, "Ошибка", "Не удалось открыть карточку помещения без номера", cad)
                        self.back_to_results()
                        continue
                    if "погаш" in normalize_text(info.status):
                        self.log(f"  Пропуск погашенного объекта без номера: {info.cadastral_number}")
                        self.excel.append_error(info.flat_no or ".", house_address, "Пропущено", "Погашенный объект", info.cadastral_number, info.address)
                        self.back_to_results()
                        continue
                    if self.is_full_info_mode():
                        rows_added = self.excel.append_full_object_info(info, search_address=house_address)
                        self.metrics["Полных карточек"] = self.metrics.get("Полных карточек", 0) + 1
                        self.log(f"  Полный парсинг по голому адресу: записано строк сведений {rows_added}; КН {info.cadastral_number}")
                    elif is_building_or_land(info.object_type):
                        self.metrics["Зданий/участков"] = self.metrics.get("Зданий/участков", 0) + 1
                        self.excel.append_building_land(info, source=f"Поиск по голому адресу: {house_address}")
                        self.log(f"  Здание/участок записан на отдельный лист: {info.object_type}; КН {info.cadastral_number}")
                    else:
                        self.write_object_rows(info)
                        self.metrics["Найдено помещений/квартир"] = self.metrics.get("Найдено помещений/квартир", 0) + 1
                        self.log(f"  Помещение без номера записано в основной лист: КН {info.cadastral_number}")
                    total_added += 1
                    if self.excel and self.cfg.get("save_after_each_flat", True):
                        self.excel.save()
                    self.back_to_results()
                    self.human_delay()

                if not self.click_next_results_page():
                    break
                page_no += 1
                self.page.wait_for_timeout(1500)
                if self.detect_rate_limit():
                    raise RateLimitError("Превышен лимит обращений Росреестра")
            except RateLimitError as e:
                rate_limit_hits += 1
                if rate_limit_hits > rate_limit_retry_count:
                    self.excel.append_error(".", house_address, "Лимит", f"Лимит обращений при перелистывании голого адреса: {e}")
                    return
                pause = float(self.cfg.get("rate_limit_pause_seconds") or 300)
                self.log(f"  Лимит при обработке голого адреса. Жду {int(pause)} сек. ({rate_limit_hits}/{rate_limit_retry_count}).")
                if not self.controlled_sleep(pause, reason="Ожидание лимита Росреестра"):
                    return
                self.open_service_page()
                self.fill_search_address(house_address)
                self.click_find()
                self.wait_results()
                page_no = 1
            except Exception as e:
                self.excel.append_error(".", house_address, "Ошибка", f"Ошибка при обработке помещений без номеров: {e}")
                return

        self.log(f"Поиск помещений без номеров завершён. Добавлено карточек: {total_added}")

    def collect_house_candidates(self, house_address: str) -> list[dict[str, Any]]:
        """Кандидаты по голому дому: тот же дом, без обычного маркера квартиры 'кв'. Здания/участки тоже собираем, но потом пишем отдельно."""
        assert self.page is not None
        rows = self.page.evaluate(
            r"""
            () => {
                const selectors = ['table tbody tr', 'tr', '[role=row]', '.table-row', '.results-row'];
                const found = [];
                const seen = new Set();
                for (const sel of selectors) {
                    for (const row of Array.from(document.querySelectorAll(sel))) {
                        const text = (row.innerText || '').trim();
                        if (!text || seen.has(text)) continue;
                        seen.add(text);
                        const rect = row.getBoundingClientRect();
                        if (rect.height <= 0 || rect.width <= 0) continue;
                        const cad = text.match(/\b\d{2}:\d{2}:\d{3,}:\d+\b/);
                        if (!cad) continue;
                        found.push({text, cad: cad[0]});
                    }
                }
                return found;
            }
            """
        )
        candidates: list[dict[str, Any]] = []
        seen: set[str] = set()
        processed = getattr(self, "processed_cadastrals", set())
        for row in rows:
            text = row.get("text") or ""
            cad = row.get("cad") or ""
            if not cad or cad in seen or cad in processed:
                continue
            # Этап "поиск помещений без номеров" работает наоборот:
            # игнорируем все строки с кв/пом/№ и также строки вида "д 7, 7".
            # Оставляем только голые объекты дома: здание, участок, помещение без номера.
            if has_unit_marker(text) or has_extra_trailing_unit_without_marker(text, extract_house_no(house_address)):
                continue
            if not house_address_matches(house_address, text, strict_street=bool(self.cfg.get("strict_street_match", False)), strict_house=bool(self.cfg.get("strict_house_match", True))):
                continue
            seen.add(cad)
            row["row_no"] = len(candidates)
            candidates.append(row)
        return candidates

    def click_next_results_page(self) -> bool:
        """Пытается перейти на следующую страницу результатов. Возвращает True, если клик выполнен."""
        assert self.page is not None
        try:
            clicked = self.page.evaluate(
                r"""
                () => {
                    function visible(el) {
                        const r = el.getBoundingClientRect();
                        const st = window.getComputedStyle(el);
                        return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
                    }
                    const bodyText = document.body.innerText || '';
                    let cur = null, total = null;
                    const m = bodyText.match(/Страница\s+(\d+)\s+из\s+(\d+)/i);
                    if (m) { cur = parseInt(m[1], 10); total = parseInt(m[2], 10); }
                    if (cur !== null && total !== null && cur >= total) return false;

                    const nodes = Array.from(document.querySelectorAll('button, a, li, [role=button]')).filter(visible);
                    const enabled = nodes.filter(el => !el.disabled && !el.classList.contains('disabled') && el.getAttribute('aria-disabled') !== 'true');

                    const nextByText = enabled.find(el => {
                        const t = ((el.innerText || el.textContent || el.getAttribute('aria-label') || el.getAttribute('title') || '')).trim().toLowerCase();
                        return t === '>' || t === '›' || t === '»' || t.includes('след') || t.includes('next');
                    });
                    if (nextByText) { nextByText.click(); return true; }

                    if (cur !== null) {
                        const nextNum = String(cur + 1);
                        const nextButton = enabled.find(el => ((el.innerText || el.textContent || '').trim() === nextNum));
                        if (nextButton) { nextButton.click(); return true; }
                    }
                    return false;
                }
                """
            )
            return bool(clicked)
        except Exception:
            return False

    def seed_snt_progress_from_excel(self) -> None:
        """Подхватывает уже записанные участки и КН из существующего result-файла СНТ."""
        if not self.excel or "Земельные участки" not in self.excel.wb.sheetnames:
            return
        try:
            ws = self.excel.wb["Земельные участки"]
            headers = {normalize_text(ws.cell(1, c).value): c for c in range(1, ws.max_column + 1) if ws.cell(1, c).value}
            plot_col = headers.get(normalize_text("Номер участка"))
            cad_col = headers.get(normalize_text("Кадастровый номер"))
            added_plots = 0
            added_cads = 0
            for r in range(2, ws.max_row + 1):
                if plot_col:
                    plot = normalize_plot_number(ws.cell(r, plot_col).value)
                    if plot and plot != "." and plot.upper() not in getattr(self, "done_units", set()):
                        self.done_units.add(plot.upper())
                        added_plots += 1
                if cad_col:
                    cad = str(ws.cell(r, cad_col).value or "").strip()
                    if re.fullmatch(r"\d{2}:\d{2}:\d{3,}:\d+", cad) and cad not in getattr(self, "processed_cadastrals", set()):
                        self.processed_cadastrals.add(cad)
                        added_cads += 1
            if added_plots or added_cads:
                self.log(f"СНТ: из существующего Excel подхвачено участков: {added_plots}, кадастровых номеров: {added_cads}.")
                self.save_progress_state()
        except Exception as e:
            self.log(f"СНТ: не удалось подхватить прогресс из Excel: {e}")

    def process_snt_mode(self, start_flat: int, end_flat: int) -> None:
        assert self.excel is not None
        city = str(self.cfg.get("snt_city") or "").strip()
        name = str(self.cfg.get("snt_name") or "").strip()
        if not city or not name:
            self.log("Режим СНТ: заполните город и название товарищества.")
            return

        search_numbered = bool(self.cfg.get("search_numbered_units", True))
        unit_queue = build_full_unit_queue(
            start_flat,
            end_flat,
            str(self.cfg.get("extra_unit_numbers") or ""),
            str(self.cfg.get("manual_unit_numbers") or ""),
            str(self.cfg.get("manual_unit_file_path") or ""),
            include_range=search_numbered,
        )
        self.metrics["Заданий всего"] = len(unit_queue)
        self.metrics["Участков в диапазоне"] = len(unit_queue)
        self.log(f"Режим СНТ: {city} / {name}. Участков в очереди: {len(unit_queue)}.")
        quarters = parse_cadastral_quarters(self.cfg.get("snt_cadastral_quarters") or "")
        if quarters:
            self.log("  Фильтр кадастровых кварталов: " + "; ".join(quarters))
        if not unit_queue:
            self.log("Очередь участков пустая.")
            return

        self.seed_snt_progress_from_excel()

        for idx, plot in enumerate(unit_queue, start=1):
            if not self.wait_control():
                break
            plot_key = normalize_plot_number(plot) or str(plot).upper()
            if bool(self.cfg.get("resume_enabled", True)) and bool(self.cfg.get("skip_done_units", True)) and plot_key in getattr(self, "done_units", set()):
                self.metrics["Пропущено как уже обработанные"] = self.metrics.get("Пропущено как уже обработанные", 0) + 1
                self.log(f"[{idx}/{len(unit_queue)}] Пропуск участка {plot}: уже обработан по файлу прогресса.")
                continue
            self.log(f"[{idx}/{len(unit_queue)}] Ищу участок {plot}: {city} {name}")
            try:
                self.process_snt_plot(str(plot))
            except Exception:
                self.log(f"Ошибка по участку {plot}:\n{traceback.format_exc()}")
                if self.excel:
                    self.excel.append_snt_not_found(str(plot), f"{city} {name} {plot}", "Ошибка при обработке, подробности в логе")
            if not self.stop_event.is_set():
                self.mark_unit_done(plot_key)
            if self.excel and self.cfg.get("save_after_each_flat", True):
                self.excel.save()
            self.human_delay()

    def collect_snt_candidates(self, search_address: str, plot_no: str) -> list[dict[str, Any]]:
        assert self.page is not None
        assert self.excel is not None
        rows = self.page.evaluate(
            r"""
            () => {
                const selectors = ['table tbody tr', 'tr', '[role=row]', '.table-row', '.results-row'];
                const found = [];
                const seen = new Set();
                for (const sel of selectors) {
                    for (const row of Array.from(document.querySelectorAll(sel))) {
                        const text = (row.innerText || '').trim();
                        if (!text || seen.has(text)) continue;
                        seen.add(text);
                        const rect = row.getBoundingClientRect();
                        if (rect.height <= 0 || rect.width <= 0) continue;
                        const cad = text.match(/\b\d{2}:\d{2}:\d{3,}:\d+\b/);
                        if (!cad) continue;
                        found.push({text, cad: cad[0]});
                    }
                }
                return found;
            }
            """
        )
        city = str(self.cfg.get("snt_city") or "").strip()
        name = str(self.cfg.get("snt_name") or "").strip()
        quarters = parse_cadastral_quarters(self.cfg.get("snt_cadastral_quarters") or "")
        candidates: list[dict[str, Any]] = []
        self.log(f"  Строк в выдаче с кадастровыми номерами: {len(rows)}")
        for row in rows:
            text = row.get("text") or ""
            cad = row.get("cad") or ""
            ok, reason = snt_candidate_match_reason(text, cad, city=city, name=name, plot_no=plot_no, quarters=quarters)
            if ok:
                row["row_no"] = len(candidates)
                candidates.append(row)
                self.excel.append_search_result(search_address, str(plot_no), cad, text, "Принят к открытию", reason)
                self.log(f"    [кандидат СНТ] {cad} | {text[:180]}")
            else:
                self.excel.append_search_result(search_address, str(plot_no), cad, text, "Отклонён", reason)
                self.excel.append_rejected(search_address, str(plot_no), cad, text, "Отклонён в СНТ", reason)
                self.metrics["Отклонённых строк поиска"] = self.metrics.get("Отклонённых строк поиска", 0) + 1
        self.log(f"  Найдено подходящих строк СНТ в выдаче: {len(candidates)}")
        return candidates

    def process_snt_plot(self, plot_no: str) -> None:
        assert self.page is not None
        assert self.excel is not None

        city = str(self.cfg.get("snt_city") or "").strip()
        name = str(self.cfg.get("snt_name") or "").strip()
        if bool(self.cfg.get("snt_recheck_not_found", False)) and self.excel:
            removed = self.excel.remove_snt_not_found_rows(plot_no)
            if removed:
                self.log(f"  СНТ: удалены старые строки «не найдено» по участку {plot_no}: {removed}")
        search_addresses = build_snt_search_addresses(city, name, plot_no, try_variants=bool(self.cfg.get("snt_try_address_variants", True)))
        if not bool(self.cfg.get("snt_try_address_variants", True)):
            self.log("  СНТ: перебор вариантов адреса выключен, используется один основной запрос.")
        retry_count = int(self.cfg.get("retry_count") or 3)
        rate_limit_retry_count = int(self.cfg.get("rate_limit_retry_count") or 12)
        rate_limit_hits = 0
        opened_any = False
        written = 0
        saw_redeemed = False

        for search_address in search_addresses:
            if not self.wait_control():
                return
            for attempt in range(1, retry_count + 1):
                try:
                    self.open_service_page()
                    if self.stop_event.is_set():
                        raise StopRequested("Остановка перед вводом поискового запроса")
                    self.fill_search_address(search_address)
                    self.click_find()
                    self.wait_results()
                    candidates = self.collect_snt_candidates(search_address, plot_no)
                    if not candidates:
                        break

                    for candidate in candidates:
                        if not self.wait_control():
                            return
                        cand_cad = candidate.get("cad") or ""
                        if bool(self.cfg.get("resume_enabled", True)) and bool(self.cfg.get("skip_done_cadastrals", True)) and cand_cad and cand_cad in getattr(self, "processed_cadastrals", set()):
                            self.log(f"  Пропуск КН {cand_cad}: уже был обработан ранее.")
                            continue

                        info = self.open_candidate_and_parse(candidate, str(plot_no), search_address)
                        opened_any = True
                        if not info:
                            self.back_to_results()
                            continue

                        status_n = normalize_text(info.status)
                        if "погаш" in status_n:
                            saw_redeemed = True
                            if not bool(self.cfg.get("snt_include_redeemed", True)):
                                self.log(f"  Погашенный участок пропущен по настройке: {info.cadastral_number}")
                                self.back_to_results()
                                continue
                            comment = "Погашенный участок"
                        elif "актуаль" in status_n or not info.status:
                            if not bool(self.cfg.get("snt_include_actual", True)):
                                self.log(f"  Актуальный участок пропущен по настройке: {info.cadastral_number}")
                                self.back_to_results()
                                continue
                            comment = "Актуальный участок"
                        else:
                            comment = f"Статус: {info.status}"

                        if not is_land_object(info.object_type):
                            if bool(self.cfg.get("snt_write_buildings", True)):
                                building_comment = "Здание/дом, найденный по номеру участка. Точная связь с земельным участком проверяется по кадастровой карте."
                                rows_added = self.excel.append_snt_building(info, str(plot_no), search_address=search_address, comment=building_comment)
                                written += rows_added
                                self.metrics["Найдено зданий СНТ"] = self.metrics.get("Найдено зданий СНТ", 0) + 1
                                if re.fullmatch(r"\d{2}:\d{2}:\d{3,}:\d+", info.cadastral_number or ""):
                                    self.processed_cadastrals.add(info.cadastral_number)
                                self.log(f"  СНТ: записано здание строк {rows_added}; участок/дом {plot_no}; КН {info.cadastral_number}; статус {info.status}")
                                self.back_to_results()
                                self.human_delay()
                                continue
                            self.excel.append_disputed(search_address, str(plot_no), info.cadastral_number or "", info.address or "", "Не записано в СНТ", f"Вид объекта не земельный участок: {info.object_type}", info.cadastral_number or "")
                            self.metrics["Спорных результатов"] = self.metrics.get("Спорных результатов", 0) + 1
                            self.log(f"  Не земельный участок, отправлено в спорные: {info.object_type}; {info.cadastral_number}")
                            self.back_to_results()
                            continue

                        rows_added = self.excel.append_snt_land(info, str(plot_no), search_address=search_address, comment=comment)
                        written += rows_added
                        if "погаш" in status_n:
                            self.metrics["Погашенных"] = self.metrics.get("Погашенных", 0) + 1
                        else:
                            self.metrics["Найдено земельных участков СНТ"] = self.metrics.get("Найдено земельных участков СНТ", 0) + 1
                        if re.fullmatch(r"\d{2}:\d{2}:\d{3,}:\d+", info.cadastral_number or ""):
                            self.processed_cadastrals.add(info.cadastral_number)
                        self.log(f"  СНТ: записано строк {rows_added}; участок {plot_no}; КН {info.cadastral_number}; статус {info.status}")
                        self.back_to_results()
                        self.human_delay()

                    if written:
                        return
                    break
                except StopRequested:
                    raise
                except RateLimitError as e:
                    rate_limit_hits += 1
                    if rate_limit_hits > rate_limit_retry_count:
                        self.metrics["Лимитов Росреестра"] = self.metrics.get("Лимитов Росреестра", 0) + 1
                        self.excel.append_snt_not_found(str(plot_no), search_address, f"Лимит обращений после {rate_limit_retry_count} ожиданий: {e}")
                        return
                    pause = float(self.cfg.get("rate_limit_pause_seconds") or 300)
                    self.log(f"  Лимит при поиске участка. Жду {int(pause)} сек. ({rate_limit_hits}/{rate_limit_retry_count}).")
                    if not self.controlled_sleep(pause, reason="Ожидание лимита Росреестра"):
                        return
                except Exception as e:
                    self.log(f"  СНТ участок {plot_no}, попытка {attempt}/{retry_count} не удалась: {e}")
                    if attempt < retry_count:
                        self.controlled_sleep(float(self.cfg.get("retry_delay_seconds") or 10))
                    else:
                        break

        if self.stop_event.is_set():
            raise StopRequested("Остановка во время обработки участка")

        if not written:
            self.metrics["Не найдено"] = self.metrics.get("Не найдено", 0) + 1
            if bool(self.cfg.get("snt_write_not_found", True)) and (normalize_plot_number(plot_no).upper() not in getattr(self, "done_units", set())):
                reason = "Не найдено подходящих земельных участков"
                if saw_redeemed and not bool(self.cfg.get("snt_include_redeemed", True)):
                    reason = "Найдены только погашенные, но они выключены настройкой"
                elif opened_any:
                    reason = "Карточки открывались, но подходящий земельный участок не записан"
                self.excel.append_snt_not_found(str(plot_no), search_addresses[0] if search_addresses else str(plot_no), reason)
            self.log(f"  СНТ: участок {plot_no} не найден/не записан.")

    def process_flat(self, flat_no: str, search_addresses: str | list[str]) -> None:
        assert self.page is not None
        assert self.excel is not None

        retry_count = int(self.cfg.get("retry_count") or 3)
        rate_limit_retry_count = int(self.cfg.get("rate_limit_retry_count") or 12)
        rate_limit_hits = 0
        attempt = 1

        if isinstance(search_addresses, str):
            variants = [search_addresses]
        else:
            variants = [a for a in search_addresses if str(a).strip()]
        if not variants:
            variants = [flat_no]
        primary_address = variants[0]

        while attempt <= retry_count:
            if not self.wait_control():
                return
            got_any_exact_candidate = False
            got_actual = False
            got_redeemed = False
            got_aux_object = False

            try:
                for variant_idx, search_address in enumerate(variants, start=1):
                    if not self.wait_control():
                        return
                    if variant_idx > 1:
                        self.log(f"  Пробую вариант адреса {variant_idx}/{len(variants)}: {search_address}")

                    self.open_service_page()
                    self.fill_search_address(search_address)
                    self.click_find()
                    self.wait_results()

                    if self.detect_captcha():
                        if not self.wait_manual_continue("Появилась капча. Введите её вручную в браузере."):
                            return
                        self.click_find()
                        self.wait_results()

                    candidates = self.collect_candidates(search_address, flat_no)
                    if not candidates:
                        continue

                    if self.is_cadastral_mode():
                        for cand in candidates:
                            if self.excel:
                                self.excel.append_cadastral_only(search_address, flat_no, cand.get("cad") or "", cand.get("text") or "", "Принят", "точное совпадение выдачи, карточка не открывалась")
                        self.metrics["Кадастровых номеров без открытия карточек"] = self.metrics.get("Кадастровых номеров без открытия карточек", 0) + len(candidates)
                        self.log(f"  Режим только КН: добавлено кадастровых номеров из выдачи: {len(candidates)}")
                        return

                    got_any_exact_candidate = True
                    for candidate in candidates:
                        if not self.wait_control():
                            return
                        cand_cad = candidate.get("cad") or ""
                        if bool(self.cfg.get("resume_enabled", True)) and bool(self.cfg.get("skip_done_cadastrals", True)) and cand_cad and cand_cad in getattr(self, "processed_cadastrals", set()):
                            self.log(f"  Пропуск КН {cand_cad}: уже был обработан ранее.")
                            continue
                        info = self.open_candidate_and_parse(candidate, flat_no, search_address)
                        if not info:
                            continue
                        if "погаш" in normalize_text(info.status):
                            got_redeemed = True
                            self.metrics["Погашенных"] = self.metrics.get("Погашенных", 0) + 1
                            self.log(f"  Пропуск погашенного объекта: {info.cadastral_number}")
                            self.back_to_results()
                            continue
                        if "актуаль" not in normalize_text(info.status) and info.status:
                            self.log(f"  Неожиданный статус объекта: {info.status}")

                        if self.is_full_info_mode():
                            rows_added = self.excel.append_full_object_info(info, search_address=search_address)
                            self.metrics["Полных карточек"] = self.metrics.get("Полных карточек", 0) + 1
                            self.log(f"  Полный парсинг: записано строк сведений {rows_added}; КН {info.cadastral_number}")
                            got_actual = True
                            return

                        if is_building_or_land(info.object_type):
                            got_aux_object = True
                            self.metrics["Зданий/участков"] = self.metrics.get("Зданий/участков", 0) + 1
                            self.excel.append_building_land(info, source=f"Попалось при поиске номера {flat_no}: {search_address}")
                            self.log(f"  Объект {info.object_type} перенесён на лист «Здания и участки»: {info.cadastral_number}")
                            self.back_to_results()
                            continue

                        self.write_object_rows(info)
                        self.metrics["Найдено помещений/квартир"] = self.metrics.get("Найдено помещений/квартир", 0) + 1
                        got_actual = True
                        return

                    # Если по этому варианту нашли точные строки, но все они погашены/не открылись,
                    # всё равно пробуем следующие варианты: иногда один формат адреса выдаёт другой актуальный КН.

                if got_redeemed and not got_actual:
                    self.excel.append_error(flat_no, primary_address, "Пропущено", "Найден только погашенный объект")
                    return

                if got_any_exact_candidate:
                    if got_aux_object:
                        self.excel.append_error(flat_no, primary_address, "Не найдено", "В выдаче были только здание/земельный участок или неподходящие объекты; подходящее помещение не найдено")
                    else:
                        self.excel.append_error(flat_no, primary_address, "Ошибка", "Не удалось открыть подходящую актуальную карточку")
                    return

                self.log("  Нет точного совпадения адреса по всем вариантам.")
                self.excel.append_error(flat_no, primary_address, "Не найдено", "Нет точного совпадения адреса по всем вариантам")
                return

            except RateLimitError as e:
                rate_limit_hits += 1
                if rate_limit_hits > rate_limit_retry_count:
                    self.metrics["Лимитов Росреестра"] = self.metrics.get("Лимитов Росреестра", 0) + 1
                    self.excel.append_error(flat_no, primary_address, "Лимит", f"Лимит обращений после {rate_limit_retry_count} ожиданий: {e}")
                    return
                pause = float(self.cfg.get("rate_limit_pause_seconds") or 300)
                self.log(f"  Росреестр сообщил о лимите обращений. Жду {int(pause)} сек. и повторяю эту же квартиру ({rate_limit_hits}/{rate_limit_retry_count}).")
                if not self.controlled_sleep(pause, reason="Ожидание лимита Росреестра"):
                    return
                # Попытку не увеличиваем: это не ошибка объекта, а временный лимит сайта.
                continue
            except StopRequested:
                raise
            except Exception as e:
                self.log(f"  Попытка {attempt}/{retry_count} не удалась: {e}")
                if attempt < retry_count:
                    self.log(f"  Повтор через {self.cfg.get('retry_delay_seconds')} сек.")
                    self.controlled_sleep(float(self.cfg.get("retry_delay_seconds") or 10))
                else:
                    self.metrics["Ошибок"] = self.metrics.get("Ошибок", 0) + 1
                    self.excel.append_error(flat_no, primary_address, "Ошибка", f"Ошибка после {retry_count} попыток: {e}")
                    return
                attempt += 1

    def fill_search_address(self, search_address: str) -> None:
        assert self.page is not None
        # На странице может быть несколько input. Ищем поле по placeholder/aria/окружению.
        idx = self.page.evaluate(
            """
            () => {
                const nodes = Array.from(document.querySelectorAll('input, textarea'));
                const good = nodes.findIndex(el => {
                    const s = ((el.placeholder || '') + ' ' + (el.getAttribute('aria-label') || '') + ' ' + (el.name || '') + ' ' + (el.id || '')).toLowerCase();
                    return s.includes('адрес') || s.includes('кадастр') || s.includes('address');
                });
                if (good >= 0) return good;
                // fallback: берём видимое текстовое поле с наибольшей шириной
                let best = -1, bestW = 0;
                nodes.forEach((el, i) => {
                    const r = el.getBoundingClientRect();
                    const st = window.getComputedStyle(el);
                    if (r.width > bestW && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none' && !el.disabled && !el.readOnly) {
                        best = i; bestW = r.width;
                    }
                });
                return best;
            }
            """
        )
        if idx is None or int(idx) < 0:
            raise RuntimeError("Не найдено поле ввода адреса/кадастрового номера")
        loc = self.page.locator("input, textarea").nth(int(idx))
        loc.click(timeout=10000)
        try:
            loc.press("Control+A")
            loc.fill(search_address)
        except Exception:
            self.page.evaluate(
                """
                ([idx, value]) => {
                    const el = Array.from(document.querySelectorAll('input, textarea'))[idx];
                    el.focus();
                    el.value = value;
                    el.dispatchEvent(new Event('input', {bubbles: true}));
                    el.dispatchEvent(new Event('change', {bubbles: true}));
                }
                """,
                [int(idx), search_address],
            )

    def click_find(self) -> None:
        assert self.page is not None
        for selector in [
            "button:has-text('Найти')",
            "input[type=submit][value*='Найти']",
            "button:has-text('НАЙТИ')",
        ]:
            try:
                self.page.locator(selector).first.click(timeout=7000)
                return
            except Exception:
                pass
        clicked = self.page.evaluate(
            """
            () => {
                const nodes = Array.from(document.querySelectorAll('button, input[type=button], input[type=submit], a'));
                const el = nodes.find(n => ((n.innerText || n.value || '')).trim().toLowerCase().includes('найти'));
                if (el) { el.click(); return true; }
                return false;
            }
            """
        )
        if not clicked:
            raise RuntimeError("Не найдена кнопка Найти")

    def wait_results(self) -> None:
        assert self.page is not None
        # Небольшая пауза нужна, чтобы не принять старую таблицу результатов за ответ нового поиска.
        self.page.wait_for_timeout(1200)
        if self.detect_rate_limit():
            raise RateLimitError("Превышен лимит обращений Росреестра")
        try:
            self.page.wait_for_function(
                """
                () => {
                    const t = document.body.innerText.toLowerCase();
                    return t.includes('превышен лимит обращений')
                        || t.includes('не удалось получить список объектов недвижимости')
                        || t.includes('найдено результатов')
                        || t.includes('не найдено')
                        || t.includes('сведения об объекте');
                }
                """,
                timeout=45000,
            )
        except Exception:
            # Некоторые версии сайта не пишут явный статус, ждём таблицу.
            self.page.wait_for_timeout(3000)
        if self.detect_rate_limit():
            raise RateLimitError("Превышен лимит обращений Росреестра")

    def collect_candidates(self, search_address: str, flat_no: str | None = None) -> list[dict[str, Any]]:
        assert self.page is not None
        rows = self.page.evaluate(
            r"""
            () => {
                const selectors = ['table tbody tr', 'tr', '[role=row]', '.table-row', '.results-row'];
                const found = [];
                const seen = new Set();
                for (const sel of selectors) {
                    for (const row of Array.from(document.querySelectorAll(sel))) {
                        const text = (row.innerText || '').trim();
                        if (!text || seen.has(text)) continue;
                        seen.add(text);
                        const rect = row.getBoundingClientRect();
                        if (rect.height <= 0 || rect.width <= 0) continue;
                        const cad = text.match(/\b\d{2}:\d{2}:\d{3,}:\d+\b/);
                        if (!cad) continue;
                        found.push({text, cad: cad[0]});
                    }
                }
                return found;
            }
            """
        )

        candidates: list[dict[str, Any]] = []
        q_house = extract_house_no(search_address)
        self.log(f"  Строк в выдаче с кадастровыми номерами: {len(rows)}")
        for row in rows:
            text = row.get("text") or ""
            cad = row.get("cad") or ""
            try:
                ok = address_matches(search_address, text, flat_no, strict_street=bool(self.cfg.get("strict_street_match", False)), strict_house=bool(self.cfg.get("strict_house_match", True)), allow_unit_without_marker=False)
            except TypeError:
                ok = address_matches(search_address, text, flat_no, strict_street=bool(self.cfg.get("strict_street_match", False)), strict_house=bool(self.cfg.get("strict_house_match", True)))
            if ok:
                reason = "адрес и номер помещения совпали"
                row["priority"] = numbered_candidate_priority(text, flat_no or "", q_house)
                row["row_no"] = len(candidates)
                candidates.append(row)
                if self.excel:
                    self.excel.append_search_result(search_address, flat_no or "", cad, text, "Принят к открытию", reason)
                self.log(f"    [кандидат] {cad} | {text[:180]}")
            else:
                reason = "не совпал адрес/дом/номер помещения"
                if self.excel:
                    self.excel.append_search_result(search_address, flat_no or "", cad, text, "Отклонён", reason)
                    self.excel.append_rejected(search_address, flat_no or "", cad, text, "Отклонён в выдаче", reason)
                self.metrics["Отклонённых строк поиска"] = self.metrics.get("Отклонённых строк поиска", 0) + 1
                self.log(f"    [отклонено] {cad} | {reason}")

        candidates.sort(key=lambda r: (int(r.get("priority", 9)), len(str(r.get("text") or ""))))
        for idx, row in enumerate(candidates):
            row["row_no"] = idx

        self.log(f"  Найдено точных строк в выдаче: {len(candidates)}")
        if candidates:
            preview = " | ".join(f"{c.get('cad')} p={c.get('priority')}" for c in candidates[:5])
            self.log(f"  Очередь открытия результатов: {preview}")
        return candidates

    def open_candidate_and_parse(self, candidate: dict[str, Any], flat_no: str, search_address: str = "") -> ObjectInfo | None:
        assert self.page is not None
        cad = candidate.get("cad") or ""
        for attempt in range(1, int(self.cfg.get("retry_count") or 3) + 1):
            try:
                # После возврата назад сайт иногда теряет таблицу результатов.
                # Перед кликом по следующему КН убеждаемся, что нужная строка снова есть на странице.
                if cad and cad not in (self.page.inner_text("body", timeout=5000) or ""):
                    self.open_service_page()
                    if search_address:
                        self.fill_search_address(search_address)
                        self.click_find()
                        self.wait_results()
                self.click_candidate(cad)
                self.page.wait_for_function(
                    """
                    () => {
                        const t = document.body.innerText.toLowerCase();
                        return t.includes('сведения об объекте') || t.includes('общая информация') || t.includes('статус объекта');
                    }
                    """,
                    timeout=45000,
                )
                if self.detect_captcha():
                    if not self.wait_manual_continue("Появилась капча при открытии объекта. Введите её вручную."):
                        return None
                text = self.page.inner_text("body", timeout=10000)
                info = parse_card_text(text, flat_no, cad)
                if not re.fullmatch(r"\d{2}:\d{2}:\d{3,}:\d+", info.cadastral_number or ""):
                    info.cadastral_number = cad or "."
                return info
            except Exception as e:
                self.log(f"  Карточка {cad}: попытка {attempt} не удалась: {e}")
                if attempt < int(self.cfg.get("retry_count") or 3):
                    self.controlled_sleep(float(self.cfg.get("retry_delay_seconds") or 10))
                    self.back_to_results()
                else:
                    return None
        return None

    def click_candidate(self, cad: str) -> None:
        assert self.page is not None
        if not cad:
            raise RuntimeError("Пустой кадастровый номер результата")

        # Сначала кликаем через DOM: так надёжнее при таблицах/SPA, чем get_by_text.
        clicked = self.page.evaluate(
            """
            (cad) => {
                function visible(el) {
                    const r = el.getBoundingClientRect();
                    const st = window.getComputedStyle(el);
                    return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
                }
                const nodes = Array.from(document.querySelectorAll('a, button, td, span, div, tr, [role=button]'))
                    .filter(n => visible(n) && (n.innerText || '').includes(cad));
                if (!nodes.length) return false;

                nodes.sort((a, b) => {
                    const at = (a.tagName || '').toLowerCase();
                    const bt = (b.tagName || '').toLowerCase();
                    const ap = at === 'a' ? 0 : at === 'button' ? 1 : at === 'td' ? 2 : 3;
                    const bp = bt === 'a' ? 0 : bt === 'button' ? 1 : bt === 'td' ? 2 : 3;
                    return ap - bp || (a.innerText || '').length - (b.innerText || '').length;
                });

                const el = nodes[0];
                const clickable = el.closest('a, button, [role=button]') || el;
                clickable.scrollIntoView({block: 'center', inline: 'center'});
                clickable.click();
                return true;
            }
            """,
            cad,
        )
        if clicked:
            return

        try:
            self.page.get_by_text(cad, exact=False).first.click(timeout=10000)
            return
        except Exception:
            pass

        raise RuntimeError(f"Не удалось открыть результат {cad}")

    def back_to_results(self) -> None:
        assert self.page is not None
        try:
            self.page.go_back(wait_until="domcontentloaded", timeout=30000)
            try:
                self.page.wait_for_load_state("networkidle", timeout=10000)
            except Exception:
                pass
        except Exception:
            self.open_service_page()

    def write_object_rows(self, info: ObjectInfo) -> None:
        assert self.excel is not None
        lone_dolevaya_placeholder = False
        if not info.rights:
            # Объект есть, прав нет: создаём строку с долей 1/1.
            # По ТЗ: номер регистрации = '.', дата регистрации пустая, признак совместной пустой, ФИО точками.
            rows = [RightInfo(right_type=".", reg_number=".", reg_date="")]
        else:
            rows = list(info.rights)
            # Важный случай старой приватизации:
            # в Росреестре "общая долевая собственность", но найден только один номер права.
            # Это не считаем 1/1: создаём дополнительную строку с точкой под незарегистрированную долю/договор.
            if len(rows) == 1 and "общ" in normalize_text(rows[0].right_type) and "долев" in normalize_text(rows[0].right_type):
                rows.append(RightInfo(right_type="Общая долевая собственность (договор/незарегистрированная доля)", reg_number=".", reg_date=""))
                lone_dolevaya_placeholder = True
                try:
                    self.excel.append_disputed(
                        "",
                        info.flat_no or ".",
                        info.cadastral_number or ".",
                        info.address or ".",
                        "Создана строка с точкой",
                        "Общая долевая собственность, но в Росреестре только один номер права; добавлена строка под старый договор/незарегистрированную долю",
                        info.cadastral_number or "",
                    )
                except Exception:
                    pass

        for right in rows:
            if right.reg_number == "." and not right.reg_date and "долев" not in normalize_text(right.right_type):
                joint, share = "", "1/1"
            else:
                joint, share = share_for_right(right, rows)

            data = {
                "Номер помещения (обяз.)": info.flat_no or ".",
                "Площадь помещения, м2 (обяз.)": info.area or ".",
                "Кадастровый номер": info.cadastral_number or ".",
                "Признак совместной собс. (\"да\" для сов. собс.)": joint,
                "Доля (1/1, 1/2, 1/3 и т.д.) (обяз., сумма в помещении должна быть 1/1)": share,
                "Муницип. собс. (\"да\" для муниц. собс.)": municipal_value(info.ownership_form),
                "Номер регистрации (обяз.)": right.reg_number or ".",
                "Дата регистрации": right.reg_date if right.reg_date else "",
                "Фамилия / Наименование организации (обяз.)": ".",
                "Имя (обяз. для физ лиц)": ".",
                "Отчество": ".",
                "Адрес (местоположение)": info.address or ".",
                "Тип помещения": info.object_type or info.purpose or "Помещение",
                "Вид права": right.right_type or "Собственность",
            }
            self.excel.append_data(data)
        if re.fullmatch(r"\d{2}:\d{2}:\d{3,}:\d+", info.cadastral_number or ""):
            try:
                self.processed_cadastrals.add(info.cadastral_number)
            except Exception:
                pass
        self.log(f"  Добавлено строк: {len(rows)}; КН: {info.cadastral_number}; статус: {info.status or 'не указан'}")


# ----------------------------- GUI -----------------------------



def snt_update_system_headers(sheet_name: str) -> list[str]:
    if normalize_text(sheet_name) == normalize_text("Здания"):
        return list(SNT_BUILDING_HEADERS)
    return list(SNT_LAND_HEADERS)


def snt_update_key_header(sheet_name: str) -> str:
    return "Номер участка/дома" if normalize_text(sheet_name) == normalize_text("Здания") else "Номер участка"


def find_probable_header_row(ws, aliases: list[str] | None = None, max_scan: int = 15) -> int:
    aliases = aliases or ["кадастровый", "номер участка", "номер участка/дома", "адрес", "статус", "площадь"]
    best_row = 1
    best_score = -1
    for r in range(1, min(ws.max_row, max_scan) + 1):
        score = 0
        for c in range(1, ws.max_column + 1):
            n = normalize_text(ws.cell(r, c).value)
            if not n:
                continue
            if any(a in n for a in aliases):
                score += 1
        if score > best_score:
            best_score = score
            best_row = r
    return best_row


def header_map_for_row(ws, header_row: int = 1) -> dict[str, int]:
    result: dict[str, int] = {}
    for c in range(1, ws.max_column + 1):
        value = ws.cell(header_row, c).value
        if value not in (None, ""):
            result[normalize_text(value)] = c
    return result


def find_col_by_alias(headers: dict[str, int], *aliases: str) -> int | None:
    # exact first
    for alias in aliases:
        n = normalize_text(alias)
        if n in headers:
            return headers[n]
    # contains second
    for key, col in headers.items():
        for alias in aliases:
            n = normalize_text(alias)
            if n and n in key:
                return col
    return None


def ensure_header_column(ws, header: str, header_row: int = 1) -> int:
    headers = header_map_for_row(ws, header_row)
    col = find_col_by_alias(headers, header)
    if col:
        return col
    col = ws.max_column + 1
    ws.cell(header_row, col).value = header
    try:
        ws.column_dimensions[ws.cell(header_row, col).column_letter].width = 24
    except Exception:
        pass
    return col


def ensure_snt_update_columns(ws, system_headers: list[str], header_row: int = 1) -> dict[str, int]:
    # Add system columns that are missing, but do not remove user columns.
    cols: dict[str, int] = {}
    for header in system_headers:
        cols[header] = ensure_header_column(ws, header, header_row)
    cols[SNT_UPDATE_STATUS_HEADER] = ensure_header_column(ws, SNT_UPDATE_STATUS_HEADER, header_row)
    cols[SNT_UPDATE_COMMENT_HEADER] = ensure_header_column(ws, SNT_UPDATE_COMMENT_HEADER, header_row)
    return cols


def normalized_cell_value(value: Any) -> str:
    s = str(value or "").strip()
    if s == "None":
        return ""
    return s


def snt_update_row_record(ws, row_idx: int, header_row: int, sheet_name: str) -> dict[str, Any]:
    headers = header_map_for_row(ws, header_row)
    plot_col = find_col_by_alias(headers, "Номер участка/дома", "Номер участка", "№ участка", "участок", "номер")
    cad_col = find_col_by_alias(headers, "Кадастровый номер", "КН")
    reg_col = find_col_by_alias(headers, "Номер регистрации", "Номер права", "Регистрация")
    address_col = find_col_by_alias(headers, "Адрес", "Адрес (местоположение)")
    status_col = find_col_by_alias(headers, "Статус объекта", "Статус")

    def get(col: int | None) -> str:
        return normalized_cell_value(ws.cell(row_idx, col).value) if col else ""

    plot = normalize_plot_number(get(plot_col))
    if not plot:
        address_text = get(address_col)
        plot = extract_snt_plot_number(address_text) or normalize_plot_number(extract_flat_no(address_text))
    cad = extract_cadastral(get(cad_col)) or get(cad_col)
    reg = get(reg_col)
    reg_key = normalize_reg_key(reg)
    values: dict[str, Any] = {}
    all_headers = header_map_for_row(ws, header_row)
    reverse = {v: k for k, v in all_headers.items()}
    for c in range(1, ws.max_column + 1):
        header = ws.cell(header_row, c).value
        if header not in (None, ""):
            values[str(header)] = ws.cell(row_idx, c).value
    return {
        "row": row_idx,
        "plot": plot,
        "cad": str(cad or "").strip(),
        "reg": reg,
        "reg_key": reg_key,
        "status": get(status_col),
        "values": values,
        "sheet": sheet_name,
    }


def read_snt_update_records(ws, sheet_name: str, header_row: int | None = None) -> list[dict[str, Any]]:
    header_row = header_row or find_probable_header_row(ws)
    records: list[dict[str, Any]] = []
    for r in range(header_row + 1, ws.max_row + 1):
        if not any(normalized_cell_value(ws.cell(r, c).value) for c in range(1, ws.max_column + 1)):
            continue
        rec = snt_update_row_record(ws, r, header_row, sheet_name)
        if not rec.get("plot") and not rec.get("cad"):
            continue
        records.append(rec)
    return records


def build_snt_new_indexes(records: list[dict[str, Any]]) -> dict[str, dict[Any, list[dict[str, Any]]]]:
    by_cad_reg: dict[tuple[str, str], list[dict[str, Any]]] = {}
    by_plot_reg: dict[tuple[str, str], list[dict[str, Any]]] = {}
    by_cad: dict[str, list[dict[str, Any]]] = {}
    by_plot: dict[str, list[dict[str, Any]]] = {}
    for rec in records:
        cad = str(rec.get("cad") or "").strip()
        plot = str(rec.get("plot") or "").strip().upper()
        reg_key = str(rec.get("reg_key") or "").strip().upper()
        if cad:
            by_cad.setdefault(cad, []).append(rec)
            if reg_key:
                by_cad_reg.setdefault((cad, reg_key), []).append(rec)
        if plot:
            by_plot.setdefault(plot, []).append(rec)
            if reg_key:
                by_plot_reg.setdefault((plot, reg_key), []).append(rec)
    return {"by_cad_reg": by_cad_reg, "by_plot_reg": by_plot_reg, "by_cad": by_cad, "by_plot": by_plot}


def choose_snt_update_match(old_rec: dict[str, Any], indexes: dict[str, dict[Any, list[dict[str, Any]]]]) -> tuple[dict[str, Any] | None, str, list[dict[str, Any]]]:
    cad = str(old_rec.get("cad") or "").strip()
    plot = str(old_rec.get("plot") or "").strip().upper()
    reg_key = str(old_rec.get("reg_key") or "").strip().upper()
    if cad and reg_key:
        items = indexes["by_cad_reg"].get((cad, reg_key), [])
        if len(items) == 1:
            return items[0], "Совпадение по кадастровому номеру и номеру права", items
        if len(items) > 1:
            return None, "Неоднозначно: несколько строк с тем же КН и номером права", items
    if plot and reg_key:
        items = indexes["by_plot_reg"].get((plot, reg_key), [])
        if len(items) == 1:
            return items[0], "Совпадение по номеру участка и номеру права", items
        if len(items) > 1:
            return None, "Неоднозначно: несколько строк с тем же участком и номером права", items
    if cad:
        items = indexes["by_cad"].get(cad, [])
        if len(items) == 1:
            return items[0], "Совпадение по кадастровому номеру", items
        if len(items) > 1:
            # Same cadastral can have several rights. Use first row but mark possible owner check.
            return items[0], "Совпадение по КН, в новом файле несколько прав", items
    if plot:
        items = indexes["by_plot"].get(plot, [])
        if len(items) == 1:
            return items[0], "Совпадение по номеру участка", items
        if len(items) > 1:
            cads = {str(x.get("cad") or "") for x in items if str(x.get("cad") or "")}
            if len(cads) == 1:
                return items[0], "Совпадение по участку, в новом файле несколько прав", items
            return None, "Неоднозначно: один номер участка найден у разных кадастровых объектов", items
    return None, "Совпадение не найдено", []


def append_update_change(ws, row: list[Any]) -> None:
    idx = ws.max_row + 1
    for c, val in enumerate(row, start=1):
        ws.cell(idx, c).value = val


def ensure_named_sheet(wb, name: str, headers: list[str]):
    if name in wb.sheetnames:
        ws = wb[name]
        ws.delete_rows(1, ws.max_row)
    else:
        ws = wb.create_sheet(name)
    for c, h in enumerate(headers, start=1):
        ws.cell(1, c).value = h
        try:
            ws.column_dimensions[ws.cell(1, c).column_letter].width = 24 if c not in {7, 8, 10} else 50
        except Exception:
            pass
    try:
        ws.freeze_panes = "A2"
    except Exception:
        pass
    return ws


def unique_new_record_id(rec: dict[str, Any]) -> tuple[str, str, str, int]:
    return (str(rec.get("sheet") or ""), str(rec.get("cad") or ""), str(rec.get("reg_key") or ""), int(rec.get("row") or 0))


def apply_snt_update_to_sheet(
    wb_old,
    ws_old,
    ws_new,
    *,
    sheet_name: str,
    options: dict[str, bool],
    changes_ws,
    ambiguous_ws,
) -> dict[str, int]:
    header_row_old = find_probable_header_row(ws_old)
    header_row_new = find_probable_header_row(ws_new)
    system_headers = snt_update_system_headers(sheet_name)
    key_header = snt_update_key_header(sheet_name)
    cols = ensure_snt_update_columns(ws_old, system_headers, header_row_old)
    headers_old = header_map_for_row(ws_old, header_row_old)
    headers_new = header_map_for_row(ws_new, header_row_new)
    old_records = read_snt_update_records(ws_old, sheet_name, header_row_old)
    new_records = read_snt_update_records(ws_new, sheet_name, header_row_new)
    indexes = build_snt_new_indexes(new_records)
    used_new: set[tuple[str, str, str, int]] = set()
    stats = {"matched": 0, "unchanged": 0, "changed": 0, "owner_change": 0, "missing": 0, "new": 0, "ambiguous": 0}

    # helper mapping for new values by exact system header aliases
    def new_value(rec: dict[str, Any], header: str) -> Any:
        vals = rec.get("values") or {}
        if header in vals:
            return vals.get(header)
        # fallback by normalized name
        norm = normalize_text(header)
        for k, v in vals.items():
            if normalize_text(k) == norm:
                return v
        return ""

    def old_cell_value(row: int, header: str) -> Any:
        col = cols.get(header) or find_col_by_alias(headers_old, header)
        return ws_old.cell(row, col).value if col else ""

    for old_rec in old_records:
        row_idx = int(old_rec["row"])
        match, reason, variants = choose_snt_update_match(old_rec, indexes)
        status_col = cols[SNT_UPDATE_STATUS_HEADER]
        comment_col = cols[SNT_UPDATE_COMMENT_HEADER]
        plot = old_rec.get("plot") or normalized_cell_value(ws_old.cell(row_idx, cols.get(key_header, 0)).value if cols.get(key_header) else "")
        cad = old_rec.get("cad") or ""

        if match is None:
            if variants:
                stats["ambiguous"] += 1
                ws_old.cell(row_idx, status_col).value = "Неоднозначное совпадение"
                ws_old.cell(row_idx, comment_col).value = reason
                append_update_change(ambiguous_ws, [
                    datetime.now().strftime("%d.%m.%Y %H:%M:%S"), sheet_name, row_idx, plot, cad, reason,
                    "; ".join(f"{v.get('plot') or ''} / {v.get('cad') or ''} / {v.get('reg') or ''}" for v in variants[:10]),
                    "Строка не обновлена автоматически",
                ])
            else:
                stats["missing"] += 1
                if options.get("keep_missing", True):
                    ws_old.cell(row_idx, status_col).value = "Объект не найден в новом парсинге"
                    ws_old.cell(row_idx, comment_col).value = "Строка оставлена без изменений"
                    append_update_change(changes_ws, [
                        datetime.now().strftime("%d.%m.%Y %H:%M:%S"), sheet_name, row_idx, plot, cad, "Объект", "есть в старом", "нет в новом", "Объект не найден", "Строка не удалена",
                    ])
            continue

        stats["matched"] += 1
        used_new.add(unique_new_record_id(match))
        row_changed = False
        owner_change = False
        changes_count = 0

        if options.get("update_rosreestr_fields", True):
            for header in system_headers:
                col = cols.get(header)
                if not col:
                    continue
                before = ws_old.cell(row_idx, col).value
                after = new_value(match, header)
                # Do not overwrite user-like empty cells with nothing.
                if after is None:
                    after = ""
                if normalized_cell_value(before) != normalized_cell_value(after):
                    ws_old.cell(row_idx, col).value = after
                    row_changed = True
                    changes_count += 1
                    append_update_change(changes_ws, [
                        datetime.now().strftime("%d.%m.%Y %H:%M:%S"), sheet_name, row_idx, plot, cad, header, before, after, "Изменено", reason,
                    ])

        old_reg_key = str(old_rec.get("reg_key") or "").strip().upper()
        new_reg_key = str(match.get("reg_key") or "").strip().upper()
        if options.get("mark_owner_change", True):
            if old_reg_key and new_reg_key and old_reg_key != new_reg_key:
                owner_change = True
            # If old had a right number and the same cad exists but no exact right match, this is a possible ownership change.
            if old_reg_key and reason in {"Совпадение по кадастровому номеру", "Совпадение по КН, в новом файле несколько прав", "Совпадение по номеру участка"}:
                if old_reg_key != new_reg_key:
                    owner_change = True

        if owner_change:
            stats["owner_change"] += 1
            ws_old.cell(row_idx, status_col).value = "Возможная смена собственника"
            ws_old.cell(row_idx, comment_col).value = "Номер права изменился или не найден в новом парсинге. Пользовательские поля сохранены."
        elif row_changed:
            stats["changed"] += 1
            ws_old.cell(row_idx, status_col).value = "Обновлены сведения Росреестра"
            ws_old.cell(row_idx, comment_col).value = f"Изменено полей: {changes_count}. Пользовательские поля сохранены."
        else:
            stats["unchanged"] += 1
            ws_old.cell(row_idx, status_col).value = "Без изменений"
            ws_old.cell(row_idx, comment_col).value = ""

    if options.get("add_new_objects", True):
        existing_new_ids = set(used_new)
        # Anything not matched is appended as a new object/row.
        for rec in new_records:
            rid = unique_new_record_id(rec)
            if rid in existing_new_ids:
                continue
            # Avoid adding another right row when same cad+reg already existed indirectly.
            new_row = ws_old.max_row + 1
            try:
                if new_row > header_row_old + 1:
                    copy_row_style(ws_old, new_row - 1, new_row)
            except Exception:
                pass
            for header in system_headers:
                col = cols.get(header)
                if col:
                    ws_old.cell(new_row, col).value = new_value(rec, header)
            ws_old.cell(new_row, cols[SNT_UPDATE_STATUS_HEADER]).value = "Новый объект"
            ws_old.cell(new_row, cols[SNT_UPDATE_COMMENT_HEADER]).value = "Добавлен из нового парсинга. Пользовательские поля пустые."
            stats["new"] += 1
            append_update_change(changes_ws, [
                datetime.now().strftime("%d.%m.%Y %H:%M:%S"), sheet_name, new_row, rec.get("plot") or "", rec.get("cad") or "", "Объект", "", "добавлен", "Новый объект", "Добавлен из нового парсинга",
            ])

    return stats


def update_snt_registry_files(
    old_registry_path: Path,
    new_result_path: Path,
    output_path: Path,
    *,
    transfer_custom_columns: bool = True,
    update_rosreestr_fields: bool = True,
    mark_owner_change: bool = True,
    keep_missing: bool = True,
    add_new_objects: bool = True,
) -> dict[str, Any]:
    from openpyxl import load_workbook

    if not old_registry_path.exists():
        raise RuntimeError(f"Старый рабочий реестр не найден: {old_registry_path}")
    if not new_result_path.exists():
        raise RuntimeError(f"Новый result-файл не найден: {new_result_path}")

    wb_old = load_workbook(old_registry_path)
    wb_new = load_workbook(new_result_path, data_only=True)
    changes_ws = ensure_named_sheet(wb_old, "Изменения СНТ", SNT_UPDATE_CHANGE_HEADERS)
    ambiguous_ws = ensure_named_sheet(wb_old, "Неоднозначные СНТ", SNT_UPDATE_AMBIGUOUS_HEADERS)

    options = {
        "transfer_custom_columns": transfer_custom_columns,
        "update_rosreestr_fields": update_rosreestr_fields,
        "mark_owner_change": mark_owner_change,
        "keep_missing": keep_missing,
        "add_new_objects": add_new_objects,
    }

    sheet_pairs: list[tuple[str, Any, Any]] = []
    if "Земельные участки" in wb_new.sheetnames:
        if "Земельные участки" in wb_old.sheetnames:
            old_land = wb_old["Земельные участки"]
        else:
            # If old workbook has only a custom sheet, use active sheet as a working land registry and keep its name.
            old_land = wb_old.active
        sheet_pairs.append((old_land.title, old_land, wb_new["Земельные участки"]))
    else:
        raise RuntimeError("В новом result-файле нет листа «Земельные участки»")

    if "Здания" in wb_new.sheetnames:
        if "Здания" not in wb_old.sheetnames:
            ws = wb_old.create_sheet("Здания")
            for c, h in enumerate(SNT_BUILDING_HEADERS, start=1):
                ws.cell(1, c).value = h
        sheet_pairs.append(("Здания", wb_old["Здания"], wb_new["Здания"]))

    total_stats: dict[str, Any] = {
        "sheets": {},
        "old": str(old_registry_path),
        "new": str(new_result_path),
        "output": str(output_path),
    }
    for name, old_ws, new_ws in sheet_pairs:
        stats = apply_snt_update_to_sheet(
            wb_old,
            old_ws,
            new_ws,
            sheet_name=name,
            options=options,
            changes_ws=changes_ws,
            ambiguous_ws=ambiguous_ws,
        )
        total_stats["sheets"][name] = stats

    # If logs are empty except headers, keep them anyway: user needs audit sheets for update mode.
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb_old.save(output_path)
    return total_stats

class App(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title(APP_TITLE)
        self.geometry("1040x760")
        self.minsize(900, 620)

        self.base_dir = app_dir()
        self.config_path = self.base_dir / "config.json"
        self.cfg = load_config(self.config_path)

        self.log_queue: queue.Queue[str] = queue.Queue()
        self.ready_event = threading.Event()
        self.pause_event = threading.Event()
        self.pause_event.set()
        self.stop_event = threading.Event()
        self.worker: ParserWorker | None = None
        self._prompted_close_task_id = 0

        self.vars: dict[str, tk.StringVar] = {}
        self.snt_vars: dict[str, tk.StringVar] = {}
        self.edit_menu: tk.Menu | None = None
        self._build_ui()
        self.protocol("WM_DELETE_WINDOW", self.on_close)
        self.refresh_action_buttons()
        self.after(200, self._drain_log_queue)

    def _make_scrollable_tab(self, notebook: ttk.Notebook, title: str) -> ttk.Frame:
        """Создаёт вкладку с вертикальной прокруткой и прокруткой колесом мыши."""
        outer = ttk.Frame(notebook)
        notebook.add(outer, text=title)

        canvas = tk.Canvas(outer, highlightthickness=0)
        scrollbar = ttk.Scrollbar(outer, orient=tk.VERTICAL, command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        content = ttk.Frame(canvas, padding=12)
        window_id = canvas.create_window((0, 0), window=content, anchor="nw")

        def update_scrollregion(_event=None) -> None:
            canvas.configure(scrollregion=canvas.bbox("all"))

        def resize_content(event) -> None:
            canvas.itemconfigure(window_id, width=event.width)

        def on_mousewheel(event) -> str:
            # Windows / macOS
            if getattr(event, "delta", 0):
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            return "break"

        def on_linux_scroll_up(_event) -> str:
            canvas.yview_scroll(-3, "units")
            return "break"

        def on_linux_scroll_down(_event) -> str:
            canvas.yview_scroll(3, "units")
            return "break"

        def bind_wheel(_event=None) -> None:
            self._active_scroll_canvas = canvas
            self.bind_all("<MouseWheel>", on_mousewheel)
            self.bind_all("<Button-4>", on_linux_scroll_up)
            self.bind_all("<Button-5>", on_linux_scroll_down)

        def unbind_wheel(_event=None) -> None:
            if getattr(self, "_active_scroll_canvas", None) is canvas:
                self._active_scroll_canvas = None
                self.unbind_all("<MouseWheel>")
                self.unbind_all("<Button-4>")
                self.unbind_all("<Button-5>")

        content.bind("<Configure>", update_scrollregion)
        canvas.bind("<Configure>", resize_content)

        # Включаем колесо при наведении на любую область вкладки.
        for widget in (outer, canvas, content):
            widget.bind("<Enter>", bind_wheel, add="+")
            widget.bind("<Leave>", unbind_wheel, add="+")

        # PageUp/PageDown тоже работают, если фокус внутри вкладки.
        canvas.bind("<Prior>", lambda _e: canvas.yview_scroll(-1, "pages"))
        canvas.bind("<Next>", lambda _e: canvas.yview_scroll(1, "pages"))

        return content

    def _build_ui(self) -> None:
        notebook = ttk.Notebook(self)
        notebook.pack(fill=tk.BOTH, expand=True)

        parser_tab = self._make_scrollable_tab(notebook, "Парсинг Росреестра ОСС")
        match_tab = self._make_scrollable_tab(notebook, "Сопоставление реестра ОСС")
        wiki_oss_tab = self._make_scrollable_tab(notebook, "WIKI ОСС")
        snt_tab = self._make_scrollable_tab(notebook, "Парсинг Росреестра СНТ")
        snt_update_tab = self._make_scrollable_tab(notebook, "Актуализация реестра СНТ")
        wiki_snt_tab = self._make_scrollable_tab(notebook, "WIKI СНТ")
        about_tab = self._make_scrollable_tab(notebook, "О программе")
        self.notebook = notebook
        self.bind("<Configure>", self.adjust_log_heights, add="+")

        frm = parser_tab

        title = ttk.Label(frm, text="Локальный парсер справочной информации Росреестра", font=("Segoe UI", 14, "bold"))
        title.pack(anchor="w", pady=(0, 10))

        profile_bar = ttk.LabelFrame(frm, text="Профиль", padding=(8, 6))
        profile_bar.pack(fill=tk.X, pady=(0, 10))

        profile_row1 = ttk.Frame(profile_bar)
        profile_row1.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(profile_row1, text="Профиль:").pack(side=tk.LEFT, padx=(0, 6))
        self.profile_var = tk.StringVar(value="")
        self.profile_combo = ttk.Combobox(profile_row1, textvariable=self.profile_var, width=42, state="readonly")
        self.profile_combo.pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(profile_row1, text="Загрузить", command=self.load_profile).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(profile_row1, text="Сохранить как…", command=self.save_profile).pack(side=tk.LEFT, padx=(0, 6))

        profile_row2 = ttk.Frame(profile_bar)
        profile_row2.pack(fill=tk.X)
        ttk.Button(profile_row2, text="Обновить список", command=self.refresh_profiles).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(profile_row2, text="Удалить", command=self.delete_profile).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(profile_row2, text="Открыть profiles", command=self.open_profiles_folder).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(profile_row2, text="Открыть state", command=self.open_state_folder).pack(side=tk.LEFT)
        self.refresh_profiles()

        grid = ttk.Frame(frm)
        grid.pack(fill=tk.X)

        def add_row(row: int, label: str, key: str, width: int = 60, browse: str | None = None) -> None:
            ttk.Label(grid, text=label).grid(row=row, column=0, sticky="w", padx=(0, 8), pady=4)
            var = tk.StringVar(value=str(self.cfg.get(key, "")))
            self.vars[key] = var
            entry = ttk.Entry(grid, textvariable=var, width=width)
            entry.grid(row=row, column=1, sticky="ew", pady=4)
            self.bind_editable_widget(entry)
            if browse:
                btn = ttk.Button(grid, text="Выбрать…", command=lambda: self._browse(key, browse))
                btn.grid(row=row, column=2, sticky="w", padx=6)

        grid.columnconfigure(1, weight=1)
        add_row(0, "Адрес до номера", "address_prefix")
        add_row(1, "Начальная квартира/объект", "start_flat", width=12)
        add_row(2, "Конечная квартира/объект", "end_flat", width=12)
        add_row(3, "Доп. номера к диапазону: 9А, 15Б, 27/1", "extra_unit_numbers")
        add_row(4, "Ручная очередь помещений: 1, 2, 3А", "manual_unit_numbers")
        add_row(5, "Файл ручной очереди: txt / Excel / Word", "manual_unit_file_path", browse="file")
        add_row(6, "Шаблон Excel", "template_path", browse="file")
        add_row(7, "Итоговый Excel", "output_path", browse="save")
        add_row(8, "Профиль браузера", "browser_profile_dir", browse="dir")
        add_row(9, "Канал браузера: chrome / msedge / пусто", "browser_channel", width=20)
        add_row(10, "URL сервиса", "service_url")

        ttk.Label(grid, text="Режим парсинга").grid(row=11, column=0, sticky="w", padx=(0, 8), pady=4)
        parser_mode_value = str(self.cfg.get("parser_mode") or ("cadastral" if self.cfg.get("collect_cadastral_only") else "oss"))
        self.parser_mode_var = tk.StringVar(value=PARSER_MODE_LABELS.get(parser_mode_value, "Реестр для ОСС"))
        parser_mode_combo = ttk.Combobox(
            grid,
            textvariable=self.parser_mode_var,
            values=["Реестр для ОСС", "Реестр всех сведений", "Реестр кадастровых номеров"],
            state="readonly",
            width=32,
        )
        parser_mode_combo.grid(row=11, column=1, sticky="w", pady=4)

        ttk.Label(grid, text="Формат выгрузки ОСС").grid(row=11, column=2, sticky="w", padx=(14, 8), pady=4)
        oss_format_value = str(self.cfg.get("oss_export_format") or "burmistr")
        self.oss_export_format_var = tk.StringVar(value=OSS_EXPORT_FORMAT_LABELS.get(oss_format_value, "Burmistr"))
        oss_format_combo = ttk.Combobox(
            grid,
            textvariable=self.oss_export_format_var,
            values=["Burmistr", "Roskvartal"],
            state="readonly",
            width=16,
        )
        oss_format_combo.grid(row=11, column=3, sticky="w", pady=4)

        self.search_numbered_var = tk.BooleanVar(value=bool(self.cfg.get("search_numbered_units", True)))
        ttk.Checkbutton(
            grid,
            text="Искать квартиры/помещения по диапазону номеров",
            variable=self.search_numbered_var,
        ).grid(row=12, column=1, sticky="w", pady=2)

        self.try_variants_var = tk.BooleanVar(value=bool(self.cfg.get("try_address_variants", True)))
        ttk.Checkbutton(
            grid,
            text="Пробовать варианты номера: кв / квартира / пом. / помещение / №",
            variable=self.try_variants_var,
        ).grid(row=13, column=1, sticky="w", pady=2)

        self.strict_street_var = tk.BooleanVar(value=bool(self.cfg.get("strict_street_match", False)))
        ttk.Checkbutton(
            grid,
            text="Строго совпадать улице (не смешивать похожие названия)",
            variable=self.strict_street_var,
        ).grid(row=14, column=1, sticky="w", pady=2)

        self.strict_house_var = tk.BooleanVar(value=bool(self.cfg.get("strict_house_match", True)))
        ttk.Checkbutton(
            grid,
            text="Строго совпадать номеру дома",
            variable=self.strict_house_var,
        ).grid(row=15, column=1, sticky="w", pady=2)

        self.include_unnumbered_var = tk.BooleanVar(value=bool(self.cfg.get("include_unnumbered_house_objects", False)))
        ttk.Checkbutton(
            grid,
            text="Поиск помещений без номеров: голый адрес дома после диапазона",
            variable=self.include_unnumbered_var,
        ).grid(row=16, column=1, sticky="w", pady=2)

        self.auto_output_filename_var = tk.BooleanVar(value=bool(self.cfg.get("auto_output_filename", True)))
        ttk.Checkbutton(
            grid,
            text="Автоимя итогового файла по адресу дома",
            variable=self.auto_output_filename_var,
        ).grid(row=17, column=1, sticky="w", pady=2)

        self.resume_enabled_var = tk.BooleanVar(value=bool(self.cfg.get("resume_enabled", True)))
        ttk.Checkbutton(
            grid,
            text="Продолжение с места: использовать файл прогресса",
            variable=self.resume_enabled_var,
        ).grid(row=18, column=1, sticky="w", pady=2)

        self.skip_done_units_var = tk.BooleanVar(value=bool(self.cfg.get("skip_done_units", True)))
        ttk.Checkbutton(
            grid,
            text="Пропускать уже обработанные номера помещений",
            variable=self.skip_done_units_var,
        ).grid(row=19, column=1, sticky="w", pady=2)

        self.skip_done_cadastrals_var = tk.BooleanVar(value=bool(self.cfg.get("skip_done_cadastrals", True)))
        ttk.Checkbutton(
            grid,
            text="Пропускать уже найденные кадастровые номера",
            variable=self.skip_done_cadastrals_var,
        ).grid(row=20, column=1, sticky="w", pady=2)

        ttk.Label(grid, text="Файл прогресса").grid(row=21, column=0, sticky="w", padx=(0, 8), pady=4)
        self.progress_path_var = tk.StringVar(value="")
        progress_entry = ttk.Entry(grid, textvariable=self.progress_path_var, state="readonly")
        progress_entry.grid(row=21, column=1, sticky="ew", pady=4)

        opts = ttk.Frame(frm)
        opts.pack(fill=tk.X, pady=(10, 0))
        ttk.Label(opts, text="Пресет пауз").grid(row=0, column=0, padx=(0, 4), pady=3, sticky="w")
        self.pause_preset_var = tk.StringVar(value=str(self.cfg.get("pause_preset", "Обычно")))
        pause_combo = ttk.Combobox(opts, textvariable=self.pause_preset_var, values=["Быстро", "Обычно", "Медленно", "Очень медленно", "Вручную"], width=16, state="readonly")
        pause_combo.grid(row=0, column=1, padx=(0, 14), pady=3, sticky="w")
        option_fields = [
            ("Пауза мин, сек", "pause_min_seconds"),
            ("Пауза макс, сек", "pause_max_seconds"),
            ("Повторов", "retry_count"),
            ("Пауза повтора, сек", "retry_delay_seconds"),
            ("Замедление браузера, мс", "slow_mo_ms"),
            ("Пауза при лимите, сек", "rate_limit_pause_seconds"),
            ("Повторов лимита", "rate_limit_retry_count"),
            ("Страниц голого адреса", "unnumbered_max_pages"),
        ]
        for idx, (label, key) in enumerate(option_fields):
            row = 1 + idx // 5
            col = (idx % 5) * 2
            ttk.Label(opts, text=label).grid(row=row, column=col, padx=(0, 4), pady=3, sticky="w")
            var = tk.StringVar(value=str(self.cfg.get(key, "")))
            self.vars[key] = var
            opt_entry = ttk.Entry(opts, textvariable=var, width=8)
            opt_entry.grid(row=row, column=col + 1, padx=(0, 14), pady=3, sticky="w")
            self.bind_editable_widget(opt_entry)

        buttons = ttk.Frame(frm)
        buttons.pack(fill=tk.X, pady=12)

        row1 = ttk.Frame(buttons)
        row1.pack(fill=tk.X, pady=(0, 5))
        self.start_btn = ttk.Button(row1, text="Старт", command=self.start_worker)
        self.start_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.continue_btn = ttk.Button(row1, text="Продолжить после входа/капчи", command=self.continue_manual)
        self.continue_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.pause_btn = ttk.Button(row1, text="Пауза", command=self.toggle_pause)
        self.pause_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.stop_btn = ttk.Button(row1, text="Стоп", command=self.stop_worker)
        self.stop_btn.pack(side=tk.LEFT, padx=(0, 8))

        row2 = ttk.Frame(buttons)
        row2.pack(fill=tk.X, pady=(0, 5))
        ttk.Button(row2, text="Сохранить настройки", command=self.save_settings).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(row2, text="Предпросмотр", command=self.show_preview).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(row2, text="Проверить Excel", command=self.validate_excel_from_ui).pack(side=tk.LEFT, padx=(0, 8))

        row3 = ttk.Frame(buttons)
        row3.pack(fill=tk.X)
        ttk.Button(row3, text="Сбросить прогресс", command=self.reset_progress_from_ui).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(row3, text="Обновить путь прогресса", command=self.update_progress_label).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(row3, text="Открыть папку результата", command=self.open_output_folder).pack(side=tk.LEFT, padx=(0, 8))

        note = ttk.Label(
            frm,
            text=(
                "Логин, пароль и SMS-код вводятся только в открытом браузере. "
                "Если появится капча — введите её вручную в браузере и нажмите продолжение."
            ),
            foreground="#555555",
        )
        note.pack(anchor="w", pady=(0, 6))

        log_toolbar = ttk.Frame(frm)
        log_toolbar.pack(fill=tk.X, pady=(0, 4))
        ttk.Button(log_toolbar, text="Копировать выделенное", command=self.copy_selected_log).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(log_toolbar, text="Копировать весь лог", command=self.copy_all_log).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(log_toolbar, text="Выделить весь лог", command=self.select_all_log).pack(side=tk.LEFT, padx=(0, 6))

        log_frame = ttk.Frame(frm)
        log_frame.pack(fill=tk.BOTH, expand=True)
        self.log_text = tk.Text(log_frame, height=10, wrap="word", font=("Consolas", 10), undo=False)
        log_scroll = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=log_scroll.set)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.bind("<Control-a>", lambda e: (self.select_all_log(), "break"))
        self.log_text.bind("<Control-A>", lambda e: (self.select_all_log(), "break"))
        self.log_text.bind("<Control-c>", lambda e: (self.copy_selected_log(), "break"))
        self.log_text.bind("<Control-C>", lambda e: (self.copy_selected_log(), "break"))
        self.log_text.bind("<Control-KeyPress>", self._log_ctrl_keypress)
        self.log_text.bind("<Button-3>", self.show_log_menu)
        self.log_menu = tk.Menu(self, tearoff=0)
        self.log_menu.add_command(label="Копировать", command=self.copy_selected_log)
        self.log_menu.add_command(label="Копировать всё", command=self.copy_all_log)
        self.log_menu.add_command(label="Выделить всё", command=self.select_all_log)

        self._build_snt_ui(snt_tab)
        self._build_snt_update_ui(snt_update_tab)
        self._build_match_ui(match_tab)
        self._build_about_ui(about_tab)
        self._build_wiki_ui(wiki_oss_tab, title_text="WIKI ОСС / МКД", wiki_filename="wiki_oss.md", prefix="wiki_oss")
        self._build_wiki_ui(wiki_snt_tab, title_text="WIKI СНТ / земли", wiki_filename="wiki_snt.md", prefix="wiki_snt")
        self.update_progress_label()
        self.update_snt_progress_label()

        self.log("Готово. Проверьте настройки и нажмите «Старт».")

    def adjust_log_heights(self, _event=None) -> None:
        """Сжимает окна логов по высоте при уменьшении окна программы."""
        try:
            h = max(6, min(16, int((self.winfo_height() - 520) / 24)))
            if hasattr(self, "log_text"):
                self.log_text.configure(height=h)
            if hasattr(self, "match_log"):
                self.match_log.configure(height=max(6, min(14, h)))
            if hasattr(self, "snt_log_text"):
                self.snt_log_text.configure(height=max(6, min(14, h)))
        except Exception:
            pass

    def _build_about_ui(self, frm: ttk.Frame) -> None:
        title = ttk.Label(frm, text="Парсер Росреестра", font=("Segoe UI", 18, "bold"))
        title.pack(anchor="w", pady=(0, 8))

        version = ttk.Label(frm, text=f"Версия: {APP_TITLE}", font=("Segoe UI", 11))
        version.pack(anchor="w", pady=(0, 16))

        info = (
            "Назначение:\n"
            "Локальный сбор бесплатной справочной информации Росреестра "
            "и сопоставление нового реестра со старыми реестрами собственников.\n\n"
            "Режим работы:\n"
            "Авторизация выполняется пользователем вручную в открытом браузере. "
            "Логин, пароль, SMS-код и капча не вводятся в программу.\n\n"
            "Программа не заказывает платные выписки, не отправляет заявления, "
            "не меняет данные аккаунта и не выполняет юридически значимые действия.\n\n"
            "Разработано лентяем для лентяев.\n"
            "Чтобы меньше копировать руками и больше пить чай."
        )

        box = tk.Text(frm, height=14, wrap="word", font=("Segoe UI", 11), padx=12, pady=12)
        box.pack(fill=tk.X, pady=(0, 12))
        box.insert("1.0", info)
        box.configure(state="disabled")
        box.bind("<Button-3>", self.show_about_menu)
        self.about_text = box

        donate_frame = ttk.Frame(frm)
        donate_frame.pack(fill=tk.X, pady=(0, 12))
        ttk.Label(donate_frame, text="На чай разработчику (Карта Сбер): 2202 2088 9799 2260").pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(donate_frame, text="Скопировать номер карты", command=self.copy_donate_card).pack(side=tk.LEFT)

        self.about_menu = tk.Menu(self, tearoff=0)
        self.about_menu.add_command(label="Копировать всё", command=self.copy_about_text)

        update_frame = ttk.LabelFrame(frm, text="Обновления и репозиторий", padding=(10, 8))
        update_frame.pack(fill=tk.BOTH, expand=True)

        repo_url = str(self.cfg.get("repository_url") or "").strip()
        manifest_url = str(self.cfg.get("update_manifest_url") or "").strip()

        ttk.Label(update_frame, text=f"Текущая версия: {APP_TITLE}").grid(row=0, column=0, columnspan=4, sticky="w", pady=(0, 6))
        self.update_status_var = tk.StringVar(value=(
            "Проверка обновлений будет доступна после указания repository_url и update_manifest_url в config.json."
            if not manifest_url else "Нажмите «Проверить обновление»."
        ))
        ttk.Label(update_frame, textvariable=self.update_status_var, wraplength=900).grid(row=1, column=0, columnspan=4, sticky="ew", pady=(0, 8))

        ttk.Button(update_frame, text="Проверить обновление", command=self.check_update).grid(row=2, column=0, sticky="w", padx=(0, 8), pady=3)
        ttk.Button(update_frame, text="Скачать ZIP обновления", command=self.download_update_zip).grid(row=2, column=1, sticky="w", padx=(0, 8), pady=3)
        ttk.Button(update_frame, text="Открыть репозиторий", command=self.open_repository_url).grid(row=2, column=2, sticky="w", padx=(0, 8), pady=3)
        ttk.Button(update_frame, text="Открыть страницу релиза", command=self.open_release_url).grid(row=2, column=3, sticky="w", padx=(0, 8), pady=3)
        ttk.Button(update_frame, text="Открыть config.json", command=self.open_config_file).grid(row=2, column=4, sticky="w", pady=3)

        repo_text = repo_url or "не задано"
        manifest_text = manifest_url or "не задано"
        ttk.Label(update_frame, text=f"Репозиторий: {repo_text}", wraplength=900).grid(row=3, column=0, columnspan=4, sticky="ew", pady=(8, 2))
        ttk.Label(update_frame, text=f"Manifest обновлений: {manifest_text}", wraplength=900).grid(row=4, column=0, columnspan=4, sticky="ew", pady=2)
        update_frame.columnconfigure(3, weight=1)

        self.last_update_info: dict[str, Any] = {}

    def download_update_zip(self) -> None:
        info = getattr(self, "last_update_info", {}) or {}
        download_url = str(info.get("download_url") or "").strip()
        if not download_url:
            self.update_status_var.set("Сначала нажмите «Проверить обновление». В manifest должен быть download_url.")
            return
        file_name = str(info.get("file_name") or "rosreestr_parser.zip").strip() or "rosreestr_parser.zip"
        expected_sha = str(info.get("sha256") or "").strip().lower()
        try:
            updates_dir = self.base_dir / "updates"
            updates_dir.mkdir(parents=True, exist_ok=True)
            dst = updates_dir / file_name
            self.update_status_var.set("Скачиваю обновление...")
            self.update_idletasks()
            with urllib.request.urlopen(download_url, timeout=60) as resp:
                data = resp.read()
            actual_sha = hashlib.sha256(data).hexdigest()
            if expected_sha and actual_sha.lower() != expected_sha:
                self.update_status_var.set(
                    "Обновление скачано, но SHA256 не совпал. Файл не сохранён.\n"
                    f"Ожидалось: {expected_sha}\nПолучено: {actual_sha}"
                )
                return
            dst.write_bytes(data)
            self.update_status_var.set(
                f"ZIP обновления скачан:\n{dst}\n\n"
                "Автоустановка пока не выполняется: распакуйте архив вручную или используйте страницу релиза."
            )
        except Exception as e:
            self.update_status_var.set(f"Не удалось скачать обновление: {e}")

    def open_url_safe(self, url: str, empty_message: str = "Ссылка не задана.") -> None:
        url = str(url or "").strip()
        if not url:
            messagebox.showinfo(APP_TITLE, empty_message)
            return
        try:
            webbrowser.open(url)
        except Exception as e:
            messagebox.showerror(APP_TITLE, f"Не удалось открыть ссылку:\n{url}\n\n{e}")

    def open_repository_url(self) -> None:
        self.open_url_safe(str(self.cfg.get("repository_url") or ""), "Ссылка на репозиторий пока не задана в config.json.")

    def open_release_url(self) -> None:
        info = getattr(self, "last_update_info", {}) or {}
        url = str(info.get("release_url") or info.get("html_url") or info.get("download_url") or "").strip()
        if not url:
            repo = str(self.cfg.get("repository_url") or "").rstrip("/")
            if repo:
                url = repo + "/releases"
        self.open_url_safe(url, "Страница релиза пока не известна. Сначала нажмите «Проверить обновление» или задайте repository_url.")

    def open_config_file(self) -> None:
        try:
            path = self.config_path
            if sys.platform.startswith("win"):
                os.startfile(str(path))
            else:
                webbrowser.open(path.as_uri())
        except Exception as e:
            messagebox.showerror(APP_TITLE, f"Не удалось открыть config.json:\n{e}")

    def _version_code_from_text(self, value: Any) -> int:
        if value is None:
            return 0
        m = re.search(r"(\d+)", str(value))
        return int(m.group(1)) if m else 0

    def check_update(self) -> None:
        manifest_url = str(self.cfg.get("update_manifest_url") or "").strip()
        if not manifest_url:
            self.update_status_var.set(
                "Manifest обновлений не задан. После создания публичного GitHub-репозитория укажите update_manifest_url в config.json."
            )
            return
        self.update_status_var.set("Проверяю обновление...")
        self.update_idletasks()
        try:
            with urllib.request.urlopen(manifest_url, timeout=15) as resp:
                raw = resp.read().decode("utf-8", errors="replace")
            data = json.loads(raw)
            if not isinstance(data, dict):
                raise ValueError("manifest должен быть JSON-объектом")
            self.last_update_info = data
            latest_version = str(data.get("latest_version") or data.get("version") or "").strip()
            latest_code = int(data.get("version_code") or self._version_code_from_text(latest_version))
            current_code = int(globals().get("APP_VERSION_CODE", self._version_code_from_text(APP_TITLE)))
            notes = data.get("notes") or data.get("changes") or []
            if isinstance(notes, list):
                notes_text = "\n".join(f"• {x}" for x in notes[:12])
            else:
                notes_text = str(notes or "").strip()
            if latest_code > current_code:
                msg = f"Доступна новая версия: {latest_version or latest_code}. Текущая версия: {APP_TITLE}."
                if notes_text:
                    msg += "\n\nЧто изменилось:\n" + notes_text
                download_url = str(data.get("download_url") or "").strip()
                if download_url:
                    msg += "\n\nМожно открыть страницу релиза или ссылку скачивания кнопкой «Открыть страницу релиза»."
                self.update_status_var.set(msg)
            else:
                self.update_status_var.set(f"Обновлений нет. Установлена актуальная версия: {APP_TITLE}.")
        except Exception as e:
            self.last_update_info = {}
            self.update_status_var.set(f"Не удалось проверить обновление: {e}")

    def copy_donate_card(self) -> None:
        self.clipboard_clear()
        self.clipboard_append("2202208897992260")

    def copy_about_text(self) -> None:
        parts = []
        widget = getattr(self, "about_text", None)
        if widget is not None:
            try:
                parts.append(widget.get("1.0", "end-1c"))
            except Exception:
                pass
        status = getattr(self, "update_status_var", None)
        if status is not None:
            try:
                parts.append(status.get())
            except Exception:
                pass
        text = "\n\n".join(parts)
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)

    def show_about_menu(self, event: tk.Event) -> str:
        try:
            self.about_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.about_menu.grab_release()
        return "break"

    def wiki_path(self, filename: str = "wiki.md") -> Path:
        # Wiki-файлы лежат рядом с программой, чтобы их можно было править без пересборки.
        p = self.base_dir / "docs" / filename
        if p.exists():
            return p
        return resource_path(f"docs/{filename}")

    def read_wiki_text(self, filename: str = "wiki.md") -> str:
        path = self.wiki_path(filename)
        try:
            return path.read_text(encoding="utf-8")
        except Exception as e:
            return f"Не удалось открыть {filename}:\n{path}\n\nОшибка: {e}"

    def wiki_sections_from_text(self, text: str) -> list[tuple[str, int]]:
        """В меню навигации показываем только заголовки разделов, а не шаги внутри текста."""
        sections: list[tuple[str, int]] = []
        lines = text.splitlines()
        for idx, line in enumerate(lines):
            s = line.strip()
            if not re.match(r"^\d+(?:\.\d+)*\.\s+\S+", s):
                continue
            prev_blank = idx == 0 or not lines[idx - 1].strip()
            next_blank = idx + 1 >= len(lines) or not lines[idx + 1].strip()
            if prev_blank and next_blank:
                sections.append((s, idx + 1))
        return sections

    def _wiki_widget(self, prefix: str, suffix: str):
        return getattr(self, f"{prefix}_{suffix}", None)

    def _build_wiki_ui(self, frm: ttk.Frame, *, title_text: str = "Mini Wiki", wiki_filename: str = "wiki.md", prefix: str = "wiki") -> None:
        title = ttk.Label(frm, text=title_text, font=("Segoe UI", 18, "bold"))
        title.pack(anchor="w", pady=(0, 8))

        nav = ttk.LabelFrame(frm, text="Навигация", padding=(8, 6))
        nav.pack(fill=tk.X, pady=(0, 8))

        ttk.Label(nav, text="Раздел").grid(row=0, column=0, sticky="w", padx=(0, 6), pady=3)
        section_var = tk.StringVar(value="")
        setattr(self, f"{prefix}_section_var", section_var)
        section_combo = ttk.Combobox(nav, textvariable=section_var, state="readonly", width=58)
        setattr(self, f"{prefix}_section_combo", section_combo)
        section_combo.grid(row=0, column=1, sticky="ew", padx=(0, 8), pady=3)
        section_combo.bind("<<ComboboxSelected>>", lambda _e, p=prefix: self.goto_wiki_section(p))
        ttk.Button(nav, text="Перейти", command=lambda p=prefix: self.goto_wiki_section(p)).grid(row=0, column=2, sticky="w", pady=3)

        ttk.Label(nav, text="Поиск").grid(row=1, column=0, sticky="w", padx=(0, 6), pady=3)
        search_var = tk.StringVar(value="")
        setattr(self, f"{prefix}_search_var", search_var)
        search_entry = ttk.Entry(nav, textvariable=search_var)
        search_entry.grid(row=1, column=1, sticky="ew", padx=(0, 8), pady=3)
        search_entry.bind("<Return>", lambda _e, p=prefix: (self.search_wiki_text(p), "break"))
        ttk.Button(nav, text="Найти", command=lambda p=prefix: self.search_wiki_text(p)).grid(row=1, column=2, sticky="w", pady=3)
        ttk.Button(nav, text="Сбросить поиск", command=lambda p=prefix: self.clear_wiki_search(p)).grid(row=1, column=3, sticky="w", padx=(8, 0), pady=3)
        nav.columnconfigure(1, weight=1)

        wiki_frame = ttk.Frame(frm)
        wiki_frame.pack(fill=tk.BOTH, expand=True)
        wiki_text = tk.Text(wiki_frame, height=36, wrap="word", font=("Segoe UI", 10), padx=12, pady=12)
        setattr(self, f"{prefix}_text", wiki_text)
        wiki_scroll = ttk.Scrollbar(wiki_frame, orient=tk.VERTICAL, command=wiki_text.yview)
        wiki_text.configure(yscrollcommand=wiki_scroll.set)
        wiki_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        wiki_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        wiki_text.tag_configure("wiki_match", background="#fff2cc")
        wiki_text.tag_configure("wiki_current", background="#ffd966")
        wiki_text.insert("1.0", self.read_wiki_text(wiki_filename))
        wiki_text.configure(state="disabled")
        wiki_text.bind("<Control-f>", lambda _e: (search_entry.focus_set(), "break"))
        wiki_text.bind("<Control-F>", lambda _e: (search_entry.focus_set(), "break"))

        self.refresh_wiki_sections(prefix)

    def refresh_wiki_sections(self, prefix: str = "wiki") -> None:
        wiki_text = self._wiki_widget(prefix, "text")
        section_combo = self._wiki_widget(prefix, "section_combo")
        section_var = self._wiki_widget(prefix, "section_var")
        if wiki_text is None or section_combo is None or section_var is None:
            return
        text = wiki_text.get("1.0", "end-1c")
        sections = self.wiki_sections_from_text(text)
        setattr(self, f"{prefix}_sections", sections)
        values = [title for title, _line in sections]
        section_combo.configure(values=values)
        if values and not section_var.get():
            section_var.set(values[0])

    def goto_wiki_section(self, prefix: str = "wiki") -> None:
        wiki_text = self._wiki_widget(prefix, "text")
        section_var = self._wiki_widget(prefix, "section_var")
        if wiki_text is None or section_var is None:
            return
        selected = section_var.get()
        for title, line_no in getattr(self, f"{prefix}_sections", []):
            if title == selected:
                wiki_text.see(f"{line_no}.0")
                wiki_text.configure(state="normal")
                wiki_text.tag_remove("wiki_current", "1.0", "end")
                wiki_text.tag_add("wiki_current", f"{line_no}.0", f"{line_no}.end")
                wiki_text.configure(state="disabled")
                break

    def search_wiki_text(self, prefix: str = "wiki") -> None:
        wiki_text = self._wiki_widget(prefix, "text")
        search_var = self._wiki_widget(prefix, "search_var")
        if wiki_text is None or search_var is None:
            return
        query = search_var.get().strip()
        wiki_text.configure(state="normal")
        wiki_text.tag_remove("wiki_match", "1.0", "end")
        wiki_text.tag_remove("wiki_current", "1.0", "end")
        if not query:
            wiki_text.configure(state="disabled")
            return

        start = wiki_text.index("insert +1c")
        pos = wiki_text.search(query, start, nocase=True, stopindex="end")
        if not pos:
            pos = wiki_text.search(query, "1.0", nocase=True, stopindex="end")
        if pos:
            end_pos = f"{pos}+{len(query)}c"
            wiki_text.tag_add("wiki_match", pos, end_pos)
            wiki_text.mark_set("insert", end_pos)
            wiki_text.see(pos)
        wiki_text.configure(state="disabled")

    def clear_wiki_search(self, prefix: str = "wiki") -> None:
        search_var = self._wiki_widget(prefix, "search_var")
        wiki_text = self._wiki_widget(prefix, "text")
        if search_var is not None:
            search_var.set("")
        if wiki_text is not None:
            wiki_text.configure(state="normal")
            wiki_text.tag_remove("wiki_match", "1.0", "end")
            wiki_text.tag_remove("wiki_current", "1.0", "end")
            wiki_text.configure(state="disabled")

    def _build_snt_ui(self, frm: ttk.Frame) -> None:
        title = ttk.Label(frm, text="Реестр земельных участков / СНТ", font=("Segoe UI", 14, "bold"))
        title.pack(anchor="w", pady=(0, 10))

        info = ttk.Label(
            frm,
            text="Этот режим отдельный и не влияет на обычный парсинг квартир/ОСС. Основные признаки: город, название товарищества, номер участка и кадастровый квартал.",
            wraplength=960,
        )
        info.pack(anchor="w", pady=(0, 10))

        grid = ttk.Frame(frm)
        grid.pack(fill=tk.X)

        def add_snt_row(row: int, label: str, key: str, width: int = 60, browse: str | None = None) -> None:
            ttk.Label(grid, text=label).grid(row=row, column=0, sticky="w", padx=(0, 8), pady=4)
            var = tk.StringVar(value=str(self.cfg.get(key, "")))
            self.snt_vars[key] = var
            entry = ttk.Entry(grid, textvariable=var, width=width)
            entry.grid(row=row, column=1, sticky="ew", pady=4)
            self.bind_editable_widget(entry)
            if browse:
                btn = ttk.Button(grid, text="Выбрать…", command=lambda: self._browse_snt(key, browse))
                btn.grid(row=row, column=2, sticky="w", padx=6)

        grid.columnconfigure(1, weight=1)
        add_snt_row(0, "Город / населённый пункт", "snt_city", width=35)
        add_snt_row(1, "Название товарищества", "snt_name", width=35)
        add_snt_row(2, "Кадастровые кварталы: 26:12:032003; 26:12:032001", "snt_cadastral_quarters")
        add_snt_row(3, "Начальный участок", "start_flat", width=12)
        add_snt_row(4, "Конечный участок", "end_flat", width=12)
        add_snt_row(5, "Доп. участки: 860А, 1309/1", "extra_unit_numbers")
        add_snt_row(6, "Ручная очередь участков", "manual_unit_numbers")
        add_snt_row(7, "Файл ручной очереди: txt / Excel / Word", "manual_unit_file_path", browse="file")
        add_snt_row(8, "Шаблон Excel СНТ", "template_path", browse="file")
        add_snt_row(9, "Итоговый Excel", "output_path", browse="save")
        add_snt_row(10, "Профиль браузера", "browser_profile_dir", browse="dir")
        add_snt_row(11, "Канал браузера: chrome / msedge / пусто", "browser_channel", width=20)
        add_snt_row(12, "URL сервиса", "service_url")

        checks = ttk.LabelFrame(frm, text="Настройки СНТ", padding=(8, 6))
        checks.pack(fill=tk.X, pady=(10, 0))

        self.snt_search_numbered_var = tk.BooleanVar(value=bool(self.cfg.get("search_numbered_units", True)))
        ttk.Checkbutton(checks, text="Перебирать участки по диапазону", variable=self.snt_search_numbered_var).grid(row=0, column=0, sticky="w", pady=2)
        self.snt_include_actual_var = tk.BooleanVar(value=bool(self.cfg.get("snt_include_actual", True)))
        ttk.Checkbutton(checks, text="Записывать актуальные участки", variable=self.snt_include_actual_var).grid(row=1, column=0, sticky="w", pady=2)
        self.snt_include_redeemed_var = tk.BooleanVar(value=bool(self.cfg.get("snt_include_redeemed", True)))
        ttk.Checkbutton(checks, text="Записывать погашенные участки", variable=self.snt_include_redeemed_var).grid(row=2, column=0, sticky="w", pady=2)
        self.snt_write_buildings_var = tk.BooleanVar(value=bool(self.cfg.get("snt_write_buildings", True)))
        ttk.Checkbutton(checks, text="СНТ: найденные здания писать в отдельный лист result", variable=self.snt_write_buildings_var).grid(row=3, column=0, sticky="w", pady=2)
        self.snt_write_not_found_var = tk.BooleanVar(value=bool(self.cfg.get("snt_write_not_found", True)))
        ttk.Checkbutton(checks, text="Добавлять не найденные участки в основной файл", variable=self.snt_write_not_found_var).grid(row=4, column=0, sticky="w", pady=2)
        self.snt_try_address_variants_var = tk.BooleanVar(value=bool(self.cfg.get("snt_try_address_variants", True)))
        ttk.Checkbutton(checks, text="Расширенный поиск СНТ — пробовать много вариантов адреса (медленно)", variable=self.snt_try_address_variants_var).grid(row=5, column=0, sticky="w", pady=2)

        self.snt_auto_output_filename_var = tk.BooleanVar(value=bool(self.cfg.get("auto_output_filename", True)))
        ttk.Checkbutton(checks, text="Автоимя result_ по городу и названию СНТ", variable=self.snt_auto_output_filename_var).grid(row=0, column=1, sticky="w", padx=(30, 0), pady=2)
        self.snt_append_existing_var = tk.BooleanVar(value=bool(self.cfg.get("snt_append_existing", True)))
        ttk.Checkbutton(checks, text="Дозаписывать в указанный result-файл, если он уже есть", variable=self.snt_append_existing_var).grid(row=1, column=1, sticky="w", padx=(30, 0), pady=2)
        self.snt_resume_enabled_var = tk.BooleanVar(value=bool(self.cfg.get("resume_enabled", True)))
        ttk.Checkbutton(checks, text="Продолжение с места", variable=self.snt_resume_enabled_var).grid(row=2, column=1, sticky="w", padx=(30, 0), pady=2)
        self.snt_skip_done_units_var = tk.BooleanVar(value=bool(self.cfg.get("skip_done_units", True)))
        ttk.Checkbutton(checks, text="Пропускать уже обработанные участки", variable=self.snt_skip_done_units_var).grid(row=3, column=1, sticky="w", padx=(30, 0), pady=2)
        self.snt_skip_done_cadastrals_var = tk.BooleanVar(value=bool(self.cfg.get("skip_done_cadastrals", True)))
        ttk.Checkbutton(checks, text="Пропускать уже найденные кадастровые номера", variable=self.snt_skip_done_cadastrals_var).grid(row=4, column=1, sticky="w", padx=(30, 0), pady=2)

        ttk.Label(checks, text="Файл прогресса").grid(row=6, column=0, sticky="w", pady=(8, 2))
        self.snt_progress_path_var = tk.StringVar(value="")
        snt_progress_entry = ttk.Entry(checks, textvariable=self.snt_progress_path_var, state="readonly")
        snt_progress_entry.grid(row=6, column=1, sticky="ew", padx=(30, 0), pady=(8, 2))
        checks.columnconfigure(1, weight=1)

        opts = ttk.Frame(frm)
        opts.pack(fill=tk.X, pady=(10, 0))
        ttk.Label(opts, text="Пресет пауз").grid(row=0, column=0, padx=(0, 4), pady=3, sticky="w")
        self.snt_pause_preset_var = tk.StringVar(value=str(self.cfg.get("pause_preset", "Обычно")))
        pause_combo = ttk.Combobox(opts, textvariable=self.snt_pause_preset_var, values=["Быстро", "Обычно", "Медленно", "Очень медленно", "Вручную"], width=16, state="readonly")
        pause_combo.grid(row=0, column=1, padx=(0, 14), pady=3, sticky="w")

        buttons = ttk.Frame(frm)
        buttons.pack(fill=tk.X, pady=12)
        self.snt_start_btn = ttk.Button(buttons, text="Старт СНТ", command=self.start_snt_worker)
        self.snt_start_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.snt_continue_btn = ttk.Button(buttons, text="Продолжить после входа/капчи", command=self.continue_manual)
        self.snt_continue_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.snt_pause_btn = ttk.Button(buttons, text="Пауза", command=self.toggle_pause)
        self.snt_pause_btn.pack(side=tk.LEFT, padx=(0, 8))
        self.snt_stop_btn = ttk.Button(buttons, text="Стоп", command=self.stop_worker)
        self.snt_stop_btn.pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(buttons, text="Сохранить настройки СНТ", command=self.save_snt_settings).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(buttons, text="Сбросить прогресс СНТ", command=self.reset_snt_progress_from_ui).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(buttons, text="Перепроверить не найденные", command=self.start_snt_recheck_not_found).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(buttons, text="Открыть папку результата", command=self.open_output_folder).pack(side=tk.LEFT)

        log_frame = ttk.LabelFrame(frm, text="Лог СНТ", padding=6)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=(8, 0))
        self.snt_log_text = tk.Text(log_frame, height=10, wrap="word", font=("Consolas", 10), undo=False)
        snt_log_scroll = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.snt_log_text.yview)
        self.snt_log_text.configure(yscrollcommand=snt_log_scroll.set)
        self.snt_log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        snt_log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.bind_editable_widget(self.snt_log_text)
        self.update_snt_progress_label()

    def _browse_snt(self, key: str, mode: str) -> None:
        initial = self.snt_vars.get(key).get() if hasattr(self, "snt_vars") and key in self.snt_vars else ""
        initial_dir = self._initial_dir_for_key(key, initial)
        if mode == "file":
            path = filedialog.askopenfilename(
                title="Выберите файл",
                initialdir=initial_dir,
                initialfile=Path(initial).name if initial else "",
                filetypes=[("Файлы", "*.xlsx *.xlsm *.xls *.docx *.txt *.csv"), ("Все файлы", "*.*")]
            )
        elif mode == "dir":
            path = filedialog.askdirectory(title="Выберите папку", initialdir=initial_dir)
        elif mode == "save":
            path = filedialog.asksaveasfilename(
                title="Сохранить Excel",
                initialdir=initial_dir,
                initialfile=Path(initial).name or "result.xlsx",
                defaultextension=".xlsx",
                filetypes=[("Excel", "*.xlsx"), ("Все файлы", "*.*")]
            )
        else:
            path = ""
        if path:
            self.snt_vars[key].set(path)
            self._remember_dir_for_key(key, path)
            self.update_snt_progress_label()

    def collect_snt_cfg(self) -> dict[str, Any]:
        cfg = self.cfg.copy()
        for key, var in getattr(self, "snt_vars", {}).items():
            value: Any = var.get().strip()
            if key in {"start_flat", "end_flat"}:
                try:
                    value = int(value)
                except Exception:
                    value = DEFAULT_CONFIG.get(key)
            cfg[key] = value
        cfg["parser_mode"] = "snt"
        cfg["collect_cadastral_only"] = False
        cfg["search_numbered_units"] = bool(getattr(self, "snt_search_numbered_var", tk.BooleanVar(value=True)).get())
        cfg["auto_output_filename"] = bool(getattr(self, "snt_auto_output_filename_var", tk.BooleanVar(value=True)).get())
        cfg["snt_append_existing"] = bool(getattr(self, "snt_append_existing_var", tk.BooleanVar(value=True)).get())
        cfg["resume_enabled"] = bool(getattr(self, "snt_resume_enabled_var", tk.BooleanVar(value=True)).get())
        cfg["skip_done_units"] = bool(getattr(self, "snt_skip_done_units_var", tk.BooleanVar(value=True)).get())
        cfg["skip_done_cadastrals"] = bool(getattr(self, "snt_skip_done_cadastrals_var", tk.BooleanVar(value=True)).get())
        cfg["snt_include_actual"] = bool(getattr(self, "snt_include_actual_var", tk.BooleanVar(value=True)).get())
        cfg["snt_include_redeemed"] = bool(getattr(self, "snt_include_redeemed_var", tk.BooleanVar(value=True)).get())
        cfg["snt_write_buildings"] = bool(getattr(self, "snt_write_buildings_var", tk.BooleanVar(value=True)).get())
        cfg["snt_only_land"] = False
        cfg["snt_write_not_found"] = bool(getattr(self, "snt_write_not_found_var", tk.BooleanVar(value=True)).get())
        cfg["snt_try_address_variants"] = bool(getattr(self, "snt_try_address_variants_var", tk.BooleanVar(value=True)).get())
        cfg["pause_preset"] = getattr(self, "snt_pause_preset_var", tk.StringVar(value="Обычно")).get()
        city = str(cfg.get("snt_city") or "").strip()
        name = str(cfg.get("snt_name") or "").strip()
        cfg["address_prefix"] = f"{city} {name}".strip()
        preset = str(cfg.get("pause_preset") or "Обычно")
        if preset != "Вручную":
            presets = {
                "Быстро": (3, 7),
                "Обычно": (7, 15),
                "Медленно": (15, 30),
                "Очень медленно": (30, 60),
            }
            if preset in presets:
                cfg["pause_min_seconds"], cfg["pause_max_seconds"] = presets[preset]
        return cfg

    def save_snt_settings(self) -> dict[str, Any]:
        cfg = self.collect_snt_cfg()
        self.cfg.update(cfg)
        save_config(self.config_path, self.cfg)
        self.update_snt_progress_label()
        self.log(f"Настройки СНТ сохранены: {self.config_path}")
        self.apply_live_settings_to_worker(log_message=False)
        return cfg

    def current_snt_progress_path(self) -> Path:
        cfg = self.collect_snt_cfg()
        city = str(cfg.get("snt_city") or "").strip()
        name = str(cfg.get("snt_name") or "").strip()
        quarters = "_".join(parse_cadastral_quarters(cfg.get("snt_cadastral_quarters") or "")) or "all_quarters"
        slug = filename_slug(f"snt_{city}_{name}_{quarters}") or "progress_snt"
        return self.base_dir / "state" / f"progress_{slug}.json"

    def update_snt_progress_label(self) -> None:
        try:
            if hasattr(self, "snt_progress_path_var"):
                self.snt_progress_path_var.set(str(self.current_snt_progress_path()))
        except Exception:
            pass

    def reset_snt_progress_from_ui(self) -> None:
        path = self.current_snt_progress_path()
        if not path.exists():
            messagebox.showinfo(APP_TITLE, f"Файл прогресса СНТ не найден:\n{path}")
            return
        if messagebox.askyesno(APP_TITLE, f"Сбросить прогресс СНТ?\n\n{path}"):
            try:
                path.unlink()
                self.log(f"Прогресс СНТ сброшен: {path}")
            except Exception as e:
                messagebox.showerror(APP_TITLE, f"Не удалось удалить прогресс СНТ: {e}")

    def resolve_snt_result_path_for_recheck(self) -> Path | None:
        """Определяет result-файл СНТ для перепроверки не найденных."""
        # В первую очередь берём файл из progress, потому что при автоимени именно там лежит настоящий путь.
        try:
            progress_path = self.current_snt_progress_path()
            if progress_path.exists():
                data = json.loads(progress_path.read_text(encoding="utf-8"))
                out = Path(str(data.get("output_path") or ""))
                if out.exists():
                    return out
        except Exception:
            pass
        raw = ""
        try:
            raw = self.snt_vars.get("output_path").get().strip()
        except Exception:
            raw = ""
        if raw:
            out = Path(raw)
            if not out.is_absolute():
                out = self.base_dir / out
            if out.exists():
                return out
        return None

    def start_snt_recheck_not_found(self) -> None:
        result_path = self.resolve_snt_result_path_for_recheck()
        if not result_path:
            messagebox.showwarning(APP_TITLE, "Не найден result-файл СНТ для перепроверки.")
            return
        plots = collect_snt_not_found_plots_from_workbook(result_path)
        if not plots:
            messagebox.showinfo(APP_TITLE, "В выбранном result-файле не найдено строк «не найдено».")
            return
        cfg = self.collect_snt_cfg()
        cfg["parser_mode"] = "snt"
        cfg["search_numbered_units"] = False
        cfg["manual_unit_numbers"] = ", ".join(plots)
        cfg["extra_unit_numbers"] = ""
        cfg["manual_unit_file_path"] = ""
        cfg["output_path"] = str(result_path)
        cfg["auto_output_filename"] = False
        cfg["snt_append_existing"] = True
        cfg["skip_done_units"] = False
        cfg["snt_recheck_not_found"] = True
        # Интерфейс тоже обновим, чтобы пользователь видел очередь.
        try:
            self.snt_vars["manual_unit_numbers"].set(cfg["manual_unit_numbers"])
            self.snt_vars["output_path"].set(str(result_path))
        except Exception:
            pass
        self.log(f"СНТ: перепроверка не найденных участков: {len(plots)} шт. ({', '.join(plots[:20])}{'...' if len(plots) > 20 else ''})")
        if self.worker and self.worker.is_alive():
            if getattr(self.worker, "idle", False):
                self.pause_event.set()
                self.ready_event.clear()
                self.cfg = cfg
                self.worker.request_next_task(cfg)
                self.refresh_action_buttons()
                self.log("Перепроверка СНТ запущена в уже открытом браузере.")
                return
            messagebox.showinfo(APP_TITLE, "Парсер уже запущен. Остановите или дождитесь завершения текущей задачи.")
            return
        self.stop_event.clear()
        self.pause_event.set()
        self.ready_event.clear()
        self._prompted_close_task_id = 0
        self.cfg = cfg
        self.worker = ParserWorker(cfg, self.log_queue, self.ready_event, self.pause_event, self.stop_event)
        self.worker.start()
        self.refresh_action_buttons()
        self.log("Старт перепроверки не найденных участков СНТ.")

    def start_snt_worker(self) -> None:
        cfg = self.save_snt_settings()
        if self.worker and self.worker.is_alive():
            if getattr(self.worker, "idle", False):
                self.pause_event.set()
                self.ready_event.clear()
                self.cfg = cfg
                self.worker.request_next_task(cfg)
                self.refresh_action_buttons()
                self.log("Старт СНТ в уже открытом браузере.")
                return
            messagebox.showinfo(APP_TITLE, "Парсер уже запущен.")
            return
        self.stop_event.clear()
        self.pause_event.set()
        self.ready_event.clear()
        self._prompted_close_task_id = 0
        self.cfg = cfg
        self.worker = ParserWorker(cfg, self.log_queue, self.ready_event, self.pause_event, self.stop_event)
        self.worker.start()
        self.refresh_action_buttons()
        self.log(f"Старт СНТ. Браузер откроется автоматически. Версия: {APP_TITLE}")

    def _build_snt_update_ui(self, frm: ttk.Frame) -> None:
        title = ttk.Label(frm, text="Актуализация реестра СНТ", font=("Segoe UI", 14, "bold"))
        title.pack(anchor="w", pady=(0, 10))

        info = ttk.Label(
            frm,
            text="Берёт старый рабочий реестр СНТ и новый result после парсинга. Обновляет поля Росреестра, переносит пользовательские колонки и не удаляет ручные данные.",
            wraplength=1000,
        )
        info.pack(anchor="w", pady=(0, 10))

        self.snt_update_vars: dict[str, tk.StringVar] = {}
        grid = ttk.Frame(frm)
        grid.pack(fill=tk.X)

        def add_row(row: int, label: str, key: str, browse: str | None = None) -> None:
            ttk.Label(grid, text=label).grid(row=row, column=0, sticky="w", padx=(0, 8), pady=5)
            var = tk.StringVar(value=str(self.cfg.get(key, "")))
            self.snt_update_vars[key] = var
            entry = ttk.Entry(grid, textvariable=var, width=82)
            entry.grid(row=row, column=1, sticky="ew", pady=5)
            self.bind_editable_widget(entry)
            if browse:
                ttk.Button(grid, text="Выбрать…", command=lambda: self._browse_snt_update(key, browse)).grid(row=row, column=2, sticky="w", padx=6)

        grid.columnconfigure(1, weight=1)
        add_row(0, "Старый рабочий реестр СНТ", "snt_update_old_registry_path", browse="file")
        add_row(1, "Новый result после парсинга СНТ", "snt_update_new_result_path", browse="file")
        add_row(2, "Итоговый актуализированный файл", "snt_update_output_path", browse="save")

        opts = ttk.LabelFrame(frm, text="Правила актуализации", padding=(8, 6))
        opts.pack(fill=tk.X, pady=(12, 8))
        self.snt_update_transfer_custom_var = tk.BooleanVar(value=bool(self.cfg.get("snt_update_transfer_custom_columns", True)))
        ttk.Checkbutton(opts, text="Переносить все пользовательские колонки", variable=self.snt_update_transfer_custom_var).pack(anchor="w")
        self.snt_update_rosreestr_fields_var = tk.BooleanVar(value=bool(self.cfg.get("snt_update_rosreestr_fields", True)))
        ttk.Checkbutton(opts, text="Обновлять поля Росреестра", variable=self.snt_update_rosreestr_fields_var).pack(anchor="w")
        self.snt_update_mark_owner_change_var = tk.BooleanVar(value=bool(self.cfg.get("snt_update_mark_owner_change", True)))
        ttk.Checkbutton(opts, text="Помечать возможную смену собственника по изменению номера права", variable=self.snt_update_mark_owner_change_var).pack(anchor="w")
        self.snt_update_keep_missing_var = tk.BooleanVar(value=bool(self.cfg.get("snt_update_keep_missing", True)))
        ttk.Checkbutton(opts, text="Не удалять старые строки, если объект не найден в новом парсинге", variable=self.snt_update_keep_missing_var).pack(anchor="w")
        self.snt_update_add_new_objects_var = tk.BooleanVar(value=bool(self.cfg.get("snt_update_add_new_objects", True)))
        ttk.Checkbutton(opts, text="Добавлять новые объекты из нового парсинга", variable=self.snt_update_add_new_objects_var).pack(anchor="w")

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(8, 8))
        ttk.Button(btns, text="Сохранить настройки", command=self.save_snt_update_settings).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(btns, text="Выполнить актуализацию", command=self.run_snt_update).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(btns, text="Открыть папку результата", command=self.open_snt_update_output_folder).pack(side=tk.LEFT)

        log_frame = ttk.LabelFrame(frm, text="Лог актуализации СНТ", padding=6)
        log_frame.pack(fill=tk.BOTH, expand=True, pady=(8, 0))
        self.snt_update_log = tk.Text(log_frame, height=16, wrap="word", font=("Consolas", 10), undo=False)
        scroll = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.snt_update_log.yview)
        self.snt_update_log.configure(yscrollcommand=scroll.set)
        self.snt_update_log.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.bind_editable_widget(self.snt_update_log)

    def _browse_snt_update(self, key: str, mode: str) -> None:
        initial = self.snt_update_vars.get(key).get() if hasattr(self, "snt_update_vars") and key in self.snt_update_vars else ""
        initial_dir = self._initial_dir_for_key(key, initial)
        if mode == "file":
            path = filedialog.askopenfilename(
                title="Выберите Excel-файл",
                initialdir=initial_dir,
                initialfile=Path(initial).name if initial else "",
                filetypes=[("Excel", "*.xlsx *.xlsm"), ("Все файлы", "*.*")],
            )
        elif mode == "save":
            path = filedialog.asksaveasfilename(
                title="Сохранить актуализированный реестр",
                initialdir=initial_dir,
                initialfile=Path(initial).name or "updated_snt_registry.xlsx",
                defaultextension=".xlsx",
                filetypes=[("Excel", "*.xlsx"), ("Все файлы", "*.*")],
            )
        else:
            path = ""
        if path:
            self.snt_update_vars[key].set(path)
            self._remember_dir_for_key(key, path)

    def collect_snt_update_cfg(self) -> dict[str, Any]:
        cfg = self.cfg.copy()
        for key, var in getattr(self, "snt_update_vars", {}).items():
            cfg[key] = var.get().strip()
        cfg["snt_update_transfer_custom_columns"] = bool(getattr(self, "snt_update_transfer_custom_var", tk.BooleanVar(value=True)).get())
        cfg["snt_update_rosreestr_fields"] = bool(getattr(self, "snt_update_rosreestr_fields_var", tk.BooleanVar(value=True)).get())
        cfg["snt_update_mark_owner_change"] = bool(getattr(self, "snt_update_mark_owner_change_var", tk.BooleanVar(value=True)).get())
        cfg["snt_update_keep_missing"] = bool(getattr(self, "snt_update_keep_missing_var", tk.BooleanVar(value=True)).get())
        cfg["snt_update_add_new_objects"] = bool(getattr(self, "snt_update_add_new_objects_var", tk.BooleanVar(value=True)).get())
        return cfg

    def save_snt_update_settings(self) -> dict[str, Any]:
        cfg = self.collect_snt_update_cfg()
        self.cfg.update(cfg)
        save_config(self.config_path, self.cfg)
        self.log_snt_update(f"Настройки актуализации СНТ сохранены: {self.config_path}")
        return cfg

    def log_snt_update(self, msg: str) -> None:
        ts = datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}\n"
        if hasattr(self, "snt_update_log"):
            self.snt_update_log.insert(tk.END, line)
            self.snt_update_log.see(tk.END)
        else:
            self.log(msg)

    def open_snt_update_output_folder(self) -> None:
        cfg = self.collect_snt_update_cfg()
        raw = cfg.get("snt_update_output_path") or cfg.get("snt_update_old_registry_path") or "."
        path = Path(str(raw))
        folder = path.parent if path.suffix else path
        if not folder.is_absolute():
            folder = self.base_dir / folder
        open_folder(folder)

    def run_snt_update(self) -> None:
        cfg = self.save_snt_update_settings()
        old_path = Path(str(cfg.get("snt_update_old_registry_path") or "").strip())
        new_path = Path(str(cfg.get("snt_update_new_result_path") or "").strip())
        out_path = Path(str(cfg.get("snt_update_output_path") or "").strip())
        if not old_path:
            messagebox.showwarning(APP_TITLE, "Выберите старый рабочий реестр СНТ.")
            return
        if not new_path:
            messagebox.showwarning(APP_TITLE, "Выберите новый result-файл после парсинга СНТ.")
            return
        if not out_path:
            out_path = old_path.with_name("updated_" + old_path.name)
        if not old_path.is_absolute():
            old_path = self.base_dir / old_path
        if not new_path.is_absolute():
            new_path = self.base_dir / new_path
        if not out_path.is_absolute():
            out_path = self.base_dir / out_path

        def worker() -> None:
            try:
                self.log_snt_update("Начинаю актуализацию СНТ…")
                result = update_snt_registry_files(
                    old_path,
                    new_path,
                    out_path,
                    transfer_custom_columns=bool(cfg.get("snt_update_transfer_custom_columns", True)),
                    update_rosreestr_fields=bool(cfg.get("snt_update_rosreestr_fields", True)),
                    mark_owner_change=bool(cfg.get("snt_update_mark_owner_change", True)),
                    keep_missing=bool(cfg.get("snt_update_keep_missing", True)),
                    add_new_objects=bool(cfg.get("snt_update_add_new_objects", True)),
                )
                self.log_snt_update(f"Готово: {out_path}")
                for sheet, stats in (result.get("sheets") or {}).items():
                    self.log_snt_update(f"{sheet}: совпало {stats.get('matched', 0)}, изменено {stats.get('changed', 0)}, возможная смена собственника {stats.get('owner_change', 0)}, новых {stats.get('new', 0)}, не найдено {stats.get('missing', 0)}, неоднозначно {stats.get('ambiguous', 0)}")
            except Exception as e:
                self.log_snt_update("Ошибка актуализации СНТ:\n" + traceback.format_exc())
                try:
                    messagebox.showerror(APP_TITLE, f"Ошибка актуализации СНТ: {e}")
                except Exception:
                    pass

        threading.Thread(target=worker, daemon=True).start()

    def _build_match_ui(self, frm: ttk.Frame) -> None:
        title = ttk.Label(frm, text="Сопоставление нового реестра со старым", font=("Segoe UI", 14, "bold"))
        title.pack(anchor="w", pady=(0, 10))

        self.match_vars: dict[str, tk.StringVar] = {}

        grid = ttk.Frame(frm)
        grid.pack(fill=tk.X)

        def add_match_row(row: int, label: str, key: str, browse: str | None = None) -> None:
            ttk.Label(grid, text=label).grid(row=row, column=0, sticky="w", padx=(0, 8), pady=5)
            var = tk.StringVar(value=str(self.cfg.get(key, "")))
            self.match_vars[key] = var
            entry = ttk.Entry(grid, textvariable=var, width=80)
            entry.grid(row=row, column=1, sticky="ew", pady=5)
            self.bind_editable_widget(entry)
            if browse:
                ttk.Button(grid, text="Выбрать…", command=lambda: self._browse_match(key, browse)).grid(row=row, column=2, padx=6, sticky="w")

        grid.columnconfigure(1, weight=1)
        add_match_row(0, "Новый реестр после парсинга", "match_new_registry_path", browse="file")
        add_match_row(1, "Старые реестры собственников (; несколько файлов)", "match_old_registry_path", browse="file")
        add_match_row(2, "Итоговый файл сопоставления", "match_output_path", browse="save")

        opts = ttk.LabelFrame(frm, text="Правила сопоставления", padding=(8, 6))
        opts.pack(fill=tk.X, pady=(12, 8))

        self.match_transfer_shares_var = tk.BooleanVar(value=bool(self.cfg.get("match_transfer_shares", True)))
        ttk.Checkbutton(
            opts,
            text="Переносить доли из старого реестра при строгом совпадении права",
            variable=self.match_transfer_shares_var,
        ).pack(anchor="w")

        self.match_allow_contracts_var = tk.BooleanVar(value=bool(self.cfg.get("match_allow_contracts_for_empty_right", True)))
        ttk.Checkbutton(
            opts,
            text="Если в новом реестре номер права '.', сопоставлять старые договоры приватизации/дарения/мены по номеру помещения",
            variable=self.match_allow_contracts_var,
        ).pack(anchor="w")

        self.match_expand_empty_var = tk.BooleanVar(value=bool(self.cfg.get("match_expand_empty_right_rows", True)))
        ttk.Checkbutton(
            opts,
            text="Разворачивать одну строку с '.' в несколько строк, если в старом реестре несколько старых договоров/собственников",
            variable=self.match_expand_empty_var,
        ).pack(anchor="w")

        self.match_preserve_existing_var = tk.BooleanVar(value=bool(self.cfg.get("match_preserve_existing_fio", True)))
        ttk.Checkbutton(
            opts,
            text="Не перезаписывать уже заполненные ФИО, только дозаполнять пустые строки",
            variable=self.match_preserve_existing_var,
        ).pack(anchor="w")

        self.match_allow_address_surname_var = tk.BooleanVar(value=bool(self.cfg.get("match_allow_address_surname", False)))
        ttk.Checkbutton(
            opts,
            text="Разрешать слабое сопоставление по адресу/номеру помещения и фамилии",
            variable=self.match_allow_address_surname_var,
        ).pack(anchor="w")

        match_format_row = ttk.Frame(opts)
        match_format_row.pack(anchor="w", fill=tk.X, pady=(6, 0))
        ttk.Label(match_format_row, text="Формат итогового файла:").pack(side=tk.LEFT, padx=(0, 8))
        match_format_value = str(self.cfg.get("match_oss_export_format") or self.cfg.get("oss_export_format") or "burmistr")
        self.match_oss_export_format_var = tk.StringVar(value=OSS_EXPORT_FORMAT_LABELS.get(match_format_value, "Burmistr"))
        ttk.Combobox(
            match_format_row,
            textvariable=self.match_oss_export_format_var,
            values=["Burmistr", "Roskvartal"],
            state="readonly",
            width=16,
        ).pack(side=tk.LEFT)

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=10)
        ttk.Button(btns, text="Предпросмотр распознавания", command=self.preview_reconcile).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(btns, text="Сопоставить", command=self.run_reconcile).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(btns, text="Открыть папку результата", command=self.open_match_output_folder).pack(side=tk.LEFT)

        help_text = (
            "Логика: ФИО переносится по строгому совпадению номера регистрации права. "
            "Если в новом реестре номер права '.', можно сопоставлять старые договоры по номеру помещения. "
            "Спорные случаи пишутся в отдельные листы итогового файла."
        )
        ttk.Label(frm, text=help_text, foreground="#555555", wraplength=950).pack(anchor="w", pady=(8, 10))

        self.match_log = tk.Text(frm, height=10, wrap="word", font=("Consolas", 10))
        self.match_log.pack(fill=tk.BOTH, expand=True)
        self.match_log.bind("<Control-a>", lambda e: (self.match_log.tag_add("sel", "1.0", "end-1c"), "break"))
        self.match_log.bind("<Control-c>", lambda e: (self._copy_match_log_selection(), "break"))
        self.match_log.bind("<Button-3>", self.show_match_log_menu)
        self.match_log_menu = tk.Menu(self, tearoff=0)
        self.match_log_menu.add_command(label="Копировать", command=self._copy_match_log_selection)
        self.match_log_menu.add_command(label="Копировать всё", command=self._copy_match_log_all)
        self.match_log_menu.add_command(label="Выделить всё", command=lambda: self.match_log.tag_add("sel", "1.0", "end-1c"))

    def match_log_write(self, msg: str) -> None:
        ts = datetime.now().strftime("%H:%M:%S")
        if hasattr(self, "match_log"):
            self.match_log.insert(tk.END, f"[{ts}] {msg}\n")
            self.match_log.see(tk.END)
        self.log(msg)

    def _copy_match_log_selection(self) -> None:
        try:
            text = self.match_log.get("sel.first", "sel.last")
        except Exception:
            text = ""
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)

    def _copy_match_log_all(self) -> None:
        text = self.match_log.get("1.0", "end-1c")
        self.clipboard_clear()
        self.clipboard_append(text)

    def show_match_log_menu(self, event: tk.Event) -> str:
        try:
            self.match_log_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.match_log_menu.grab_release()
        return "break"

    def _browse_match(self, key: str, mode: str) -> None:
        initial = self.match_vars.get(key).get() if hasattr(self, "match_vars") and key in self.match_vars else ""
        initial_dir = self._initial_dir_for_key(key, initial)
        if mode == "file":
            if key == "match_old_registry_path":
                paths = filedialog.askopenfilenames(
                    title="Выберите один или несколько старых реестров",
                    initialdir=initial_dir,
                    filetypes=[("Реестры", "*.xlsx *.xlsm *.xls *.docx"), ("Excel", "*.xlsx *.xlsm *.xls"), ("Word", "*.docx"), ("Все файлы", "*.*")]
                )
                path = "; ".join(paths)
            else:
                path = filedialog.askopenfilename(
                    title="Выберите новый реестр после парсинга",
                    initialdir=initial_dir,
                    initialfile=Path(initial).name if initial else "",
                    filetypes=[("Реестры Excel", "*.xlsx *.xlsm"), ("Excel", "*.xlsx *.xlsm"), ("Все файлы", "*.*")]
                )
        elif mode == "save":
            path = filedialog.asksaveasfilename(
                title="Сохранить результат сопоставления",
                initialdir=initial_dir,
                initialfile=Path(initial).name or "matched.xlsx",
                defaultextension=".xlsx",
                filetypes=[("Excel", "*.xlsx"), ("Все файлы", "*.*")]
            )
        else:
            path = ""
        if path:
            self.match_vars[key].set(path)
            self._remember_dir_for_key(key, path)
            if key == "match_new_registry_path" and hasattr(self, "match_vars") and "match_output_path" in self.match_vars:
                current_out = self.match_vars["match_output_path"].get().strip()
                if not current_out or Path(current_out).name.lower() == "matched_registry.xlsx":
                    self.match_vars["match_output_path"].set(str(build_matched_output_path(Path(path), self.base_dir / "output")))

    def collect_match_cfg(self) -> dict[str, Any]:
        cfg = self.cfg.copy()
        if hasattr(self, "match_vars"):
            for key, var in self.match_vars.items():
                cfg[key] = var.get().strip()
        cfg["match_transfer_shares"] = bool(getattr(self, "match_transfer_shares_var", tk.BooleanVar(value=True)).get())
        cfg["match_allow_contracts_for_empty_right"] = bool(getattr(self, "match_allow_contracts_var", tk.BooleanVar(value=True)).get())
        cfg["match_expand_empty_right_rows"] = bool(getattr(self, "match_expand_empty_var", tk.BooleanVar(value=True)).get())
        cfg["match_preserve_existing_fio"] = bool(getattr(self, "match_preserve_existing_var", tk.BooleanVar(value=True)).get())
        cfg["match_allow_address_surname"] = bool(getattr(self, "match_allow_address_surname_var", tk.BooleanVar(value=False)).get())
        format_label = getattr(self, "match_oss_export_format_var", tk.StringVar(value="Burmistr")).get()
        cfg["match_oss_export_format"] = OSS_EXPORT_FORMATS.get(format_label, "burmistr")
        return cfg

    def save_match_settings(self) -> dict[str, Any]:
        cfg = self.collect_match_cfg()
        self.cfg.update(cfg)
        save_config(self.config_path, self.cfg)
        return cfg

    def preview_reconcile(self) -> None:
        cfg = self.save_match_settings()
        old_paths = resolve_old_registry_paths(cfg.get("match_old_registry_path") or "")
        new_path = Path(cfg.get("match_new_registry_path") or "")
        if not old_paths or any(not p.exists() for p in old_paths):
            messagebox.showerror(APP_TITLE, "Укажите существующий старый реестр или несколько реестров через ';'.")
            return
        if not new_path.exists():
            messagebox.showerror(APP_TITLE, "Укажите существующий новый реестр.")
            return
        try:
            records = []
            for old_path in old_paths:
                records.extend(extract_old_registry_records(old_path))
            records = dedupe_old_records(records)
            reg_count = sum(1 for r in records if r.right_key and not r.is_contract)
            contract_count = sum(1 for r in records if r.right_key and r.is_contract)
            fio_count = sum(1 for r in records if r.fio_original and r.fio_original != ".")
            self.match_log_write(f"Старый реестр распознан: записей {len(records)}, с номерами регистрации {reg_count}, договоров {contract_count}, с ФИО {fio_count}.")
            for rec in records[:20]:
                self.match_log_write(f"  {rec.flat_no or '-'} | {rec.reg_number or rec.contract_text or rec.raw_right[:60]} | {rec.fio_original} | доля {rec.share or '-'}")
            if len(records) > 20:
                self.match_log_write(f"  ... показаны первые 20 записей из {len(records)}")
        except Exception as e:
            self.match_log_write("Ошибка предпросмотра:\n" + traceback.format_exc())
            messagebox.showerror(APP_TITLE, f"Ошибка предпросмотра: {e}")

    def run_reconcile(self) -> None:
        cfg = self.save_match_settings()
        new_path = Path(cfg.get("match_new_registry_path") or "")
        old_paths = resolve_old_registry_paths(cfg.get("match_old_registry_path") or "")
        output_raw = str(cfg.get("match_output_path") or "").strip()
        default_names = {"matched_registry.xlsx", "output/matched_registry.xlsx"}
        if not output_raw or output_raw.replace("\\", "/") in default_names:
            output_path = build_matched_output_path(new_path, self.base_dir / "output")
        else:
            output_path = Path(output_raw)
            if not output_path.is_absolute():
                output_path = self.base_dir / output_path
        if not new_path.exists():
            messagebox.showerror(APP_TITLE, "Укажите существующий новый реестр.")
            return
        if not old_paths or any(not p.exists() for p in old_paths):
            messagebox.showerror(APP_TITLE, "Укажите существующий старый реестр или несколько реестров через ';'.")
            return
        try:
            self.match_log_write("Начинаю сопоставление…")
            self.match_log_write("Старые реестры: " + "; ".join(str(p) for p in old_paths))
            result = reconcile_registries(
                new_path,
                old_paths,
                output_path,
                transfer_shares=bool(cfg.get("match_transfer_shares", True)),
                allow_contracts_for_empty_right=bool(cfg.get("match_allow_contracts_for_empty_right", True)),
                expand_empty_right_rows=bool(cfg.get("match_expand_empty_right_rows", True)),
                preserve_existing_fio=bool(cfg.get("match_preserve_existing_fio", True)),
                allow_address_surname_match=bool(cfg.get("match_allow_address_surname", False)),
                export_format=str(cfg.get("match_oss_export_format") or "burmistr"),
            )
            self.match_vars["match_output_path"].set(str(output_path))
            self.match_log_write(
                "Сопоставление завершено: "
                f"старых записей {result.get('old_records')}, "
                f"по номеру права {result.get('matched_by_reg')}, "
                f"по договорам {result.get('matched_by_contract')}, "
                f"адрес+фамилия {result.get('matched_by_address_surname')}, "
                f"добавлено строк {result.get('expanded_rows')}, "
                f"неоднозначных {result.get('ambiguous')}, "
                f"не найдено {result.get('not_found')}."
            )
            self.match_log_write(f"Итоговый файл: {output_path}")
            if result.get("log_output"):
                self.match_log_write(f"Лог-файл: {result.get('log_output')}")
            messagebox.showinfo(APP_TITLE, f"Сопоставление завершено.\n\nФайл:\n{output_path}\n\nЛог:\n{result.get('log_output', '')}")
        except Exception as e:
            self.match_log_write("Ошибка сопоставления:\n" + traceback.format_exc())
            messagebox.showerror(APP_TITLE, f"Ошибка сопоставления: {e}")

    def open_match_output_folder(self) -> None:
        cfg = self.collect_match_cfg()
        out = Path(str(cfg.get("match_output_path") or "output/matched_registry.xlsx"))
        if not out.is_absolute():
            out = self.base_dir / out
        folder = out.parent
        folder.mkdir(parents=True, exist_ok=True)
        try:
            os.startfile(folder)  # type: ignore[attr-defined]
        except Exception as e:
            messagebox.showerror(APP_TITLE, f"Не удалось открыть папку: {e}")

    def bind_editable_widget(self, widget: tk.Widget) -> None:
        """Добавляет нормальную вставку/копирование для полей ввода: Ctrl+V и ПКМ.

        На Windows при русской раскладке Tkinter иногда не обрабатывает Ctrl+V
        стандартно, поэтому дополнительно ловим keycode клавиш A/C/V/X.
        """
        widget.bind("<Control-v>", self._edit_paste)
        widget.bind("<Control-V>", self._edit_paste)
        widget.bind("<Control-c>", self._edit_copy)
        widget.bind("<Control-C>", self._edit_copy)
        widget.bind("<Control-x>", self._edit_cut)
        widget.bind("<Control-X>", self._edit_cut)
        widget.bind("<Control-a>", self._edit_select_all)
        widget.bind("<Control-A>", self._edit_select_all)
        widget.bind("<Control-KeyPress>", self._edit_ctrl_keypress)
        widget.bind("<Shift-Insert>", self._edit_paste)
        widget.bind("<Button-3>", self.show_edit_menu)
        widget.bind("<Button-2>", self.show_edit_menu)

    def _edit_ctrl_keypress(self, event: tk.Event) -> str | None:
        # Windows virtual-key codes: A=65, C=67, V=86, X=88.
        keycode = getattr(event, "keycode", None)
        if keycode == 86:
            return self._edit_paste(event)
        if keycode == 67:
            return self._edit_copy(event)
        if keycode == 88:
            return self._edit_cut(event)
        if keycode == 65:
            return self._edit_select_all(event)
        return None

    def _edit_select_all(self, event: tk.Event) -> str:
        widget = event.widget
        try:
            if isinstance(widget, tk.Text):
                widget.tag_add("sel", "1.0", "end-1c")
            else:
                widget.selection_range(0, "end")
            widget.focus_set()
        except Exception:
            pass
        return "break"

    def _edit_copy(self, event: tk.Event) -> str:
        widget = event.widget
        try:
            text = widget.selection_get()
            self.clipboard_clear()
            self.clipboard_append(text)
        except Exception:
            try:
                widget.event_generate("<<Copy>>")
            except Exception:
                pass
        return "break"

    def _edit_cut(self, event: tk.Event) -> str:
        widget = event.widget
        try:
            widget.event_generate("<<Cut>>")
        except Exception:
            pass
        return "break"

    def _edit_paste(self, event: tk.Event) -> str:
        widget = event.widget
        try:
            text = self.clipboard_get()
        except Exception:
            text = ""
        if text:
            try:
                if isinstance(widget, tk.Text):
                    try:
                        widget.delete("sel.first", "sel.last")
                    except Exception:
                        pass
                    widget.insert("insert", text)
                else:
                    try:
                        widget.delete("sel.first", "sel.last")
                    except Exception:
                        pass
                    widget.insert("insert", text)
            except Exception:
                try:
                    widget.event_generate("<<Paste>>")
                except Exception:
                    pass
        else:
            try:
                widget.event_generate("<<Paste>>")
            except Exception:
                pass
        return "break"

    def show_edit_menu(self, event: tk.Event) -> str:
        widget = event.widget
        try:
            widget.focus_set()
        except Exception:
            pass
        if self.edit_menu is None:
            self.edit_menu = tk.Menu(self, tearoff=0)
            self.edit_menu.add_command(label="Вырезать", command=lambda: self._edit_cut_from_focus())
            self.edit_menu.add_command(label="Копировать", command=lambda: self._edit_copy_from_focus())
            self.edit_menu.add_command(label="Вставить", command=lambda: self._edit_paste_from_focus())
            self.edit_menu.add_separator()
            self.edit_menu.add_command(label="Выделить всё", command=lambda: self._edit_select_all_from_focus())
        try:
            self.edit_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.edit_menu.grab_release()
        return "break"

    def _focus_event(self) -> tk.Event:
        ev = tk.Event()
        ev.widget = self.focus_get()
        return ev

    def _edit_cut_from_focus(self) -> None:
        self._edit_cut(self._focus_event())

    def _edit_copy_from_focus(self) -> None:
        self._edit_copy(self._focus_event())

    def _edit_paste_from_focus(self) -> None:
        self._edit_paste(self._focus_event())

    def _edit_select_all_from_focus(self) -> None:
        self._edit_select_all(self._focus_event())

    def _initial_dir_for_key(self, key: str, current_value: str = "") -> str:
        dir_key_map = {
            "template_path": "last_dir_template",
            "output_path": "last_dir_output",
            "manual_unit_file_path": "last_dir_manual_queue",
            "browser_profile_dir": "last_dir_browser_profile",
            "match_new_registry_path": "last_dir_new_registry",
            "match_old_registry_path": "last_dir_old_registry",
            "match_output_path": "last_dir_match_output",
        }
        cfg_key = dir_key_map.get(key)
        saved = str(self.cfg.get(cfg_key, "") or "").strip() if cfg_key else ""
        if saved and Path(saved).exists():
            return saved

        if current_value:
            first = str(current_value).split(";")[0].strip().strip('"')
            p = Path(first)
            if p.exists():
                return str(p if p.is_dir() else p.parent)
            if p.parent and str(p.parent) != "." and p.parent.exists():
                return str(p.parent)

        return str(self.base_dir)

    def _remember_dir_for_key(self, key: str, selected_path: str) -> None:
        if not selected_path:
            return
        dir_key_map = {
            "template_path": "last_dir_template",
            "output_path": "last_dir_output",
            "manual_unit_file_path": "last_dir_manual_queue",
            "browser_profile_dir": "last_dir_browser_profile",
            "match_new_registry_path": "last_dir_new_registry",
            "match_old_registry_path": "last_dir_old_registry",
            "match_output_path": "last_dir_match_output",
        }
        cfg_key = dir_key_map.get(key)
        if not cfg_key:
            return
        first = str(selected_path).split(";")[0].strip().strip('"')
        p = Path(first)
        folder = p if p.is_dir() else p.parent
        if folder and str(folder) != ".":
            self.cfg[cfg_key] = str(folder)
            try:
                save_config(self.config_path, self.cfg)
            except Exception:
                pass

    def _browse(self, key: str, mode: str) -> None:
        initial = self.vars.get(key).get() if key in self.vars else ""
        initial_dir = self._initial_dir_for_key(key, initial)
        if mode == "file":
            if key == "manual_unit_file_path":
                path = filedialog.askopenfilename(
                    title="Выберите файл ручной очереди",
                    initialdir=initial_dir,
                    initialfile=Path(initial).name if initial else "",
                    filetypes=[
                        ("Списки номеров", "*.txt *.csv *.xlsx *.xlsm *.docx"),
                        ("Текст", "*.txt *.csv"),
                        ("Excel", "*.xlsx *.xlsm"),
                        ("Word", "*.docx"),
                        ("Все файлы", "*.*"),
                    ],
                )
            else:
                path = filedialog.askopenfilename(
                    title="Выберите файл",
                    initialdir=initial_dir,
                    initialfile=Path(initial).name if initial else "",
                    filetypes=[("Excel", "*.xlsx *.xlsm"), ("Все файлы", "*.*")]
                )
        elif mode == "save":
            path = filedialog.asksaveasfilename(
                title="Сохранить как",
                initialdir=initial_dir,
                initialfile=Path(initial).name or "result.xlsx",
                defaultextension=".xlsx",
                filetypes=[("Excel", "*.xlsx"), ("Все файлы", "*.*")]
            )
        else:
            path = filedialog.askdirectory(title="Выберите папку", initialdir=initial_dir)
        if path:
            self.vars[key].set(path)
            self._remember_dir_for_key(key, path)

    def select_all_log(self) -> None:
        self.log_text.tag_add("sel", "1.0", "end-1c")
        self.log_text.focus_set()

    def copy_selected_log(self) -> None:
        try:
            text = self.log_text.get("sel.first", "sel.last")
        except tk.TclError:
            text = ""
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)

    def copy_all_log(self) -> None:
        text = self.log_text.get("1.0", "end-1c")
        self.clipboard_clear()
        self.clipboard_append(text)

    def show_log_menu(self, event: tk.Event) -> str:
        try:
            self.log_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.log_menu.grab_release()
        return "break"

    def _log_ctrl_keypress(self, event: tk.Event) -> str | None:
        # Для русской раскладки на Windows: A=65, C=67.
        keycode = getattr(event, "keycode", None)
        if keycode == 65:
            self.select_all_log()
            return "break"
        if keycode == 67:
            self.copy_selected_log()
            return "break"
        return None

    def profiles_dir(self) -> Path:
        d = self.base_dir / "profiles"
        d.mkdir(parents=True, exist_ok=True)
        return d

    def safe_profile_name(self, name: str) -> str:
        name = re.sub(r"[<>:\"/\\|?*.,;№]+", " ", name or "").strip()
        name = re.sub(r"\s+", "_", name)
        return name or "profile"

    def refresh_profiles(self) -> None:
        try:
            names = sorted(p.stem for p in self.profiles_dir().glob("*.json"))
            if hasattr(self, "profile_combo"):
                self.profile_combo.configure(values=names)
                current = self.profile_var.get().strip() if hasattr(self, "profile_var") else ""
                if current and current not in names:
                    self.profile_var.set("")
        except Exception as e:
            self.log(f"Не удалось обновить список профилей: {e}")

    def profile_path_by_name(self, name: str) -> Path:
        return self.profiles_dir() / f"{self.safe_profile_name(name)}.json"

    def ask_profile_name(self, initial: str = "") -> str | None:
        win = tk.Toplevel(self)
        win.title("Сохранить профиль")
        win.transient(self)
        win.grab_set()
        win.geometry("560x150")
        win.minsize(480, 140)

        frm = ttk.Frame(win, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frm, text="Название профиля:").pack(anchor="w")
        var = tk.StringVar(value=initial)
        entry = ttk.Entry(frm, textvariable=var, width=70)
        entry.pack(fill=tk.X, pady=(6, 12))
        self.bind_editable_widget(entry)

        result: dict[str, str | None] = {"value": None}

        def ok() -> None:
            value = var.get().strip()
            if not value:
                messagebox.showwarning(APP_TITLE, "Введите название профиля.", parent=win)
                return
            result["value"] = value
            win.destroy()

        def cancel() -> None:
            result["value"] = None
            win.destroy()

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X)
        ttk.Button(btns, text="OK", command=ok).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(btns, text="Отмена", command=cancel).pack(side=tk.LEFT)

        entry.focus_set()
        entry.selection_range(0, "end")
        win.bind("<Return>", lambda _e: ok())
        win.bind("<Escape>", lambda _e: cancel())
        self.wait_window(win)
        return result["value"]

    def save_profile(self) -> None:
        cfg = self.collect_cfg()
        selected = self.profile_var.get().strip() if hasattr(self, "profile_var") else ""
        default_name = selected or filename_slug(house_address_from_prefix(str(cfg.get("address_prefix", ""))) or "profile")
        name = self.ask_profile_name(default_name)
        if not name:
            return
        path = self.profile_path_by_name(name)
        if path.exists() and not messagebox.askyesno(APP_TITLE, f"Профиль уже существует. Перезаписать?\n\n{path}"):
            return
        with path.open("w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
        self.refresh_profiles()
        self.profile_var.set(path.stem)
        self.log(f"Профиль сохранён: {path}")

    def load_profile(self) -> None:
        selected = self.profile_var.get().strip() if hasattr(self, "profile_var") else ""
        path: Path | None = None
        if selected:
            path = self.profile_path_by_name(selected)
            if not path.exists():
                path = None
        if path is None:
            chosen = filedialog.askopenfilename(title="Выберите профиль", initialdir=str(self.profiles_dir()), filetypes=[("JSON", "*.json"), ("Все файлы", "*.*")])
            if not chosen:
                return
            path = Path(chosen)
        try:
            with path.open("r", encoding="utf-8") as f:
                cfg = json.load(f)
            if not isinstance(cfg, dict):
                raise ValueError("Профиль не является JSON-объектом")
            self.cfg.update(cfg)
            for key, var in self.vars.items():
                if key in self.cfg:
                    var.set(str(self.cfg.get(key, "")))
            for key, var in getattr(self, "snt_vars", {}).items():
                if key in self.cfg:
                    var.set(str(self.cfg.get(key, "")))
            for key, var in getattr(self, "snt_update_vars", {}).items():
                if key in self.cfg:
                    var.set(str(self.cfg.get(key, "")))
            self.search_numbered_var.set(bool(self.cfg.get("search_numbered_units", True)))
            self.try_variants_var.set(bool(self.cfg.get("try_address_variants", True)))
            self.strict_street_var.set(bool(self.cfg.get("strict_street_match", False)))
            self.strict_house_var.set(bool(self.cfg.get("strict_house_match", True)))
            self.include_unnumbered_var.set(bool(self.cfg.get("include_unnumbered_house_objects", False)))
            self.auto_output_filename_var.set(bool(self.cfg.get("auto_output_filename", True)))
            if hasattr(self, "snt_auto_output_filename_var"):
                self.snt_auto_output_filename_var.set(bool(self.cfg.get("auto_output_filename", True)))
            if hasattr(self, "snt_append_existing_var"):
                self.snt_append_existing_var.set(bool(self.cfg.get("snt_append_existing", True)))
            if hasattr(self, "snt_include_actual_var"):
                self.snt_include_actual_var.set(bool(self.cfg.get("snt_include_actual", True)))
            if hasattr(self, "snt_include_redeemed_var"):
                self.snt_include_redeemed_var.set(bool(self.cfg.get("snt_include_redeemed", True)))
            if hasattr(self, "snt_write_buildings_var"):
                self.snt_write_buildings_var.set(bool(self.cfg.get("snt_write_buildings", True)))
            if hasattr(self, "snt_write_not_found_var"):
                self.snt_write_not_found_var.set(bool(self.cfg.get("snt_write_not_found", True)))
            if hasattr(self, "snt_try_address_variants_var"):
                self.snt_try_address_variants_var.set(bool(self.cfg.get("snt_try_address_variants", True)))
            mode_value = str(self.cfg.get("parser_mode") or ("cadastral" if self.cfg.get("collect_cadastral_only") else "oss"))
            if hasattr(self, "parser_mode_var"):
                self.parser_mode_var.set(PARSER_MODE_LABELS.get(mode_value, "Реестр для ОСС") if mode_value in {"oss", "full", "cadastral"} else "Реестр для ОСС")
            if hasattr(self, "oss_export_format_var"):
                self.oss_export_format_var.set(OSS_EXPORT_FORMAT_LABELS.get(str(self.cfg.get("oss_export_format") or "burmistr"), "Burmistr"))
            self.pause_preset_var.set(str(self.cfg.get("pause_preset", "Обычно")))
            self.refresh_profiles()
            self.profile_var.set(path.stem)
            self.log(f"Профиль загружен: {path}")
        except Exception as e:
            messagebox.showerror(APP_TITLE, f"Не удалось загрузить профиль: {e}")

    def delete_profile(self) -> None:
        selected = self.profile_var.get().strip() if hasattr(self, "profile_var") else ""
        path: Path | None = None
        if selected:
            path = self.profile_path_by_name(selected)
        if path is None or not path.exists():
            chosen = filedialog.askopenfilename(title="Выберите профиль для удаления", initialdir=str(self.profiles_dir()), filetypes=[("JSON", "*.json"), ("Все файлы", "*.*")])
            if not chosen:
                return
            path = Path(chosen)
        if messagebox.askyesno(APP_TITLE, f"Удалить профиль?\n\n{path}"):
            try:
                path.unlink()
                self.log(f"Профиль удалён: {path}")
                self.refresh_profiles()
                if hasattr(self, "profile_var"):
                    self.profile_var.set("")
            except Exception as e:
                messagebox.showerror(APP_TITLE, f"Не удалось удалить профиль: {e}")

    def show_preview(self) -> None:
        cfg = self.collect_cfg()
        lines: list[str] = []
        address_prefix = str(cfg.get("address_prefix") or "")
        start_flat = int(cfg.get("start_flat") or 1)
        end_flat = int(cfg.get("end_flat") or start_flat)
        try_variants = bool(cfg.get("try_address_variants", True))

        lines.append("ПРЕДПРОСМОТР ПОИСКОВЫХ ЗАПРОСОВ")
        lines.append("")
        try:
            unit_queue = build_full_unit_queue(
                start_flat,
                end_flat,
                cfg.get("extra_unit_numbers", ""),
                cfg.get("manual_unit_numbers", ""),
                cfg.get("manual_unit_file_path", ""),
                include_range=bool(cfg.get("search_numbered_units", True)),
            )
        except Exception as e:
            unit_queue = build_full_unit_queue(
                start_flat,
                end_flat,
                cfg.get("extra_unit_numbers", ""),
                cfg.get("manual_unit_numbers", ""),
                "",
                include_range=bool(cfg.get("search_numbered_units", True)),
            )
            lines.append(f"Файл ручной очереди не прочитан: {e}")
            lines.append("")
        lines.append("Очередь номеров:")
        lines.append(", ".join(unit_queue) if unit_queue else "(пусто)")
        lines.append("")
        for flat in unit_queue[:31]:
            variants = build_search_addresses(address_prefix, str(flat), enabled=try_variants)
            lines.append(f"№ {flat}:")
            for v in variants:
                lines.append(f"  - {v}")
            lines.append("")
        if len(unit_queue) > 31:
            lines.append(f"... показан 31 номер из очереди {len(unit_queue)}")
        if not cfg.get("search_numbered_units", True):
            lines.append("")
            lines.append("Диапазон квартир/помещений выключен; используется только ручная очередь, если она заполнена.")

        if cfg.get("include_unnumbered_house_objects", False):
            lines.append("")
            lines.append("Поиск помещений без номеров / голый адрес дома:")
            lines.append(f"  - {house_address_from_prefix(address_prefix)}")
        if cfg.get("collect_cadastral_only", False):
            lines.append("")
            lines.append("Режим: только собрать кадастровые номера, карточки не открывать.")

        win = tk.Toplevel(self)
        win.title("Предпросмотр поиска")
        win.geometry("900x650")
        txt = tk.Text(win, wrap="word", font=("Consolas", 10))
        txt.pack(fill=tk.BOTH, expand=True)
        txt.insert("1.0", "\n".join(lines))
        txt.configure(state="normal")
        self.bind_editable_widget(txt)

    def collect_cfg(self) -> dict[str, Any]:
        cfg = self.cfg.copy()
        cfg["search_numbered_units"] = bool(getattr(self, "search_numbered_var", tk.BooleanVar(value=True)).get())
        cfg["try_address_variants"] = bool(getattr(self, "try_variants_var", tk.BooleanVar(value=True)).get())
        cfg["strict_street_match"] = bool(getattr(self, "strict_street_var", tk.BooleanVar(value=False)).get())
        cfg["strict_house_match"] = bool(getattr(self, "strict_house_var", tk.BooleanVar(value=True)).get())
        cfg["include_unnumbered_house_objects"] = bool(getattr(self, "include_unnumbered_var", tk.BooleanVar(value=False)).get())
        cfg["auto_output_filename"] = bool(getattr(self, "auto_output_filename_var", tk.BooleanVar(value=True)).get())
        mode_label = getattr(self, "parser_mode_var", tk.StringVar(value="Реестр для ОСС")).get()
        cfg["parser_mode"] = PARSER_MODES.get(mode_label, "oss")
        format_label = getattr(self, "oss_export_format_var", tk.StringVar(value="Burmistr")).get()
        cfg["oss_export_format"] = OSS_EXPORT_FORMATS.get(format_label, "burmistr")
        cfg["collect_cadastral_only"] = cfg["parser_mode"] == "cadastral"
        cfg["resume_enabled"] = bool(getattr(self, "resume_enabled_var", tk.BooleanVar(value=True)).get())
        cfg["skip_done_units"] = bool(getattr(self, "skip_done_units_var", tk.BooleanVar(value=True)).get())
        cfg["skip_done_cadastrals"] = bool(getattr(self, "skip_done_cadastrals_var", tk.BooleanVar(value=True)).get())
        cfg["pause_preset"] = getattr(self, "pause_preset_var", tk.StringVar(value="Обычно")).get()
        for key, var in self.vars.items():
            value: Any = var.get().strip()
            if key in {"start_flat", "end_flat", "retry_count", "slow_mo_ms", "rate_limit_retry_count", "unnumbered_max_pages"}:
                try:
                    value = int(value)
                except Exception:
                    value = DEFAULT_CONFIG.get(key)
            if key in {"pause_min_seconds", "pause_max_seconds", "retry_delay_seconds", "rate_limit_pause_seconds"}:
                try:
                    value = float(value)
                except Exception:
                    value = DEFAULT_CONFIG.get(key)
            cfg[key] = value
        preset = str(cfg.get("pause_preset") or "Обычно")
        if preset != "Вручную":
            presets = {
                "Быстро": (3, 7),
                "Обычно": (7, 15),
                "Медленно": (15, 30),
                "Очень медленно": (30, 60),
            }
            if preset in presets:
                cfg["pause_min_seconds"], cfg["pause_max_seconds"] = presets[preset]
                if "pause_min_seconds" in self.vars:
                    self.vars["pause_min_seconds"].set(str(cfg["pause_min_seconds"]))
                if "pause_max_seconds" in self.vars:
                    self.vars["pause_max_seconds"].set(str(cfg["pause_max_seconds"]))
        return cfg

    def current_progress_path(self) -> Path:
        cfg = self.collect_cfg()
        start_flat = int(cfg.get("start_flat") or 1)
        end_flat = int(cfg.get("end_flat") or start_flat)
        address_prefix = str(cfg.get("address_prefix") or "").strip()
        house = house_address_from_prefix(address_prefix) or address_prefix or "house"
        slug = filename_slug(f"{house}_kv_{start_flat}-{end_flat}") or "progress"
        return self.base_dir / "state" / f"progress_{slug}.json"

    def reset_progress_from_ui(self) -> None:
        path = self.current_progress_path()
        if not path.exists():
            messagebox.showinfo(APP_TITLE, f"Файл прогресса не найден:\n{path}")
            return
        if messagebox.askyesno(APP_TITLE, f"Сбросить прогресс для текущего дома?\n\n{path}"):
            try:
                path.unlink()
                self.log(f"Прогресс сброшен: {path}")
            except Exception as e:
                messagebox.showerror(APP_TITLE, f"Не удалось удалить прогресс: {e}")

    def update_progress_label(self) -> None:
        try:
            if hasattr(self, "progress_path_var"):
                self.progress_path_var.set(str(self.current_progress_path()))
        except Exception:
            pass

    def open_profiles_folder(self) -> None:
        folder = self.profiles_dir()
        folder.mkdir(parents=True, exist_ok=True)
        try:
            os.startfile(folder)  # type: ignore[attr-defined]
        except Exception as e:
            messagebox.showerror(APP_TITLE, f"Не удалось открыть папку profiles: {e}")

    def open_state_folder(self) -> None:
        folder = self.base_dir / "state"
        folder.mkdir(parents=True, exist_ok=True)
        try:
            os.startfile(folder)  # type: ignore[attr-defined]
        except Exception as e:
            messagebox.showerror(APP_TITLE, f"Не удалось открыть папку state: {e}")

    def validate_excel_from_ui(self) -> None:
        input_path = filedialog.askopenfilename(
            title="Выберите Excel для проверки",
            filetypes=[("Excel", "*.xlsx *.xlsm"), ("Все файлы", "*.*")]
        )
        if not input_path:
            return
        src = Path(input_path)
        default_name = f"{src.stem}_checked.xlsx"
        output_path = filedialog.asksaveasfilename(
            title="Сохранить проверенный Excel",
            initialdir=str(src.parent),
            initialfile=default_name,
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx"), ("Все файлы", "*.*")]
        )
        if not output_path:
            return
        try:
            result = validate_existing_registry_excel(src, Path(output_path))
            self.log(f"Проверка Excel завершена: {output_path}; замечаний: {result.get('checks')}, дублей: {result.get('duplicates')}")
            messagebox.showinfo(APP_TITLE, f"Проверка завершена.\n\nФайл:\n{output_path}\n\nЗамечаний: {result.get('checks')}\nДублей: {result.get('duplicates')}")
        except Exception as e:
            self.log("Ошибка проверки Excel:\n" + traceback.format_exc())
            messagebox.showerror(APP_TITLE, f"Ошибка проверки Excel: {e}")

    def save_settings(self) -> None:
        self.cfg = self.collect_cfg()
        save_config(self.config_path, self.cfg)
        self.update_progress_label()
        self.log(f"Настройки сохранены: {self.config_path}")

    def collect_live_cfg_for_worker(self) -> dict[str, Any]:
        """Считывает безопасные настройки, которые можно менять во время паузы."""
        worker_mode = ""
        try:
            worker_mode = str((self.worker.cfg if self.worker else {}).get("parser_mode") or "")
        except Exception:
            worker_mode = ""
        if worker_mode == "snt":
            return self.collect_snt_cfg()
        return self.collect_cfg()

    def apply_live_settings_to_worker(self, *, log_message: bool = True) -> None:
        """Применяет к работающему парсеру только безопасные настройки без смены задачи."""
        worker = self.worker
        if not worker or not worker.is_alive():
            return
        try:
            new_cfg = self.collect_live_cfg_for_worker()
            safe_keys = {
                "pause_preset", "pause_min_seconds", "pause_max_seconds",
                "retry_count", "retry_delay_seconds", "rate_limit_pause_seconds", "rate_limit_retry_count",
                "try_address_variants", "strict_street_match", "strict_house_match",
                "snt_try_address_variants", "snt_include_actual", "snt_include_redeemed", "snt_write_buildings", "snt_write_not_found",
                "skip_done_cadastrals",
            }
            for key in safe_keys:
                if key in new_cfg:
                    worker.cfg[key] = new_cfg[key]
                    self.cfg[key] = new_cfg[key]
            if log_message:
                self.log("Настройки скорости/поиска применены к текущей задаче.")
        except Exception as e:
            self.log(f"Не удалось применить настройки к текущей задаче: {e}")

    def refresh_action_buttons(self) -> None:
        """Обновляет подписи и доступность кнопок Старт/Пауза/Стоп."""
        worker_active = bool(self.worker and self.worker.is_alive())
        paused = worker_active and not self.pause_event.is_set()
        idle = worker_active and bool(getattr(self.worker, "idle", False))

        pause_text = "Продолжить" if paused else "Пауза"
        for name in ("pause_btn", "snt_pause_btn"):
            btn = getattr(self, name, None)
            if btn:
                try:
                    btn.configure(text=pause_text, state=("normal" if worker_active and not idle else "disabled"))
                except Exception:
                    pass

        for name in ("stop_btn", "snt_stop_btn"):
            btn = getattr(self, name, None)
            if btn:
                try:
                    btn.configure(state=("normal" if worker_active else "disabled"))
                except Exception:
                    pass

        # Старт активен, если парсер не работает, или если браузер оставлен открытым и поток idle.
        start_state = "normal" if (not worker_active or idle) else "disabled"
        for name in ("start_btn", "snt_start_btn"):
            btn = getattr(self, name, None)
            if btn:
                try:
                    btn.configure(state=start_state)
                except Exception:
                    pass

        # Кнопка входа/капчи не нужна во время обычной паузы.
        continue_state = "disabled" if paused else ("normal" if worker_active else "disabled")
        for name in ("continue_btn", "snt_continue_btn"):
            btn = getattr(self, name, None)
            if btn:
                try:
                    btn.configure(state=continue_state)
                except Exception:
                    pass

    def start_worker(self) -> None:
        if self.worker and self.worker.is_alive():
            if getattr(self.worker, "idle", False):
                self.save_settings()
                self.pause_event.set()
                self.ready_event.clear()
                self.worker.request_next_task(self.cfg)
                self.refresh_action_buttons()
                self.log("Старт новой задачи в уже открытом браузере.")
                return
            messagebox.showinfo(APP_TITLE, "Парсер уже запущен.")
            return
        self.save_settings()
        self.stop_event.clear()
        self.pause_event.set()
        self.ready_event.clear()
        self._prompted_close_task_id = 0
        self.worker = ParserWorker(self.cfg, self.log_queue, self.ready_event, self.pause_event, self.stop_event)
        self.worker.start()
        self.refresh_action_buttons()
        self.log(f"Старт. Браузер откроется автоматически. Версия: {APP_TITLE}")

    def continue_manual(self) -> None:
        self.ready_event.set()
        self.log("Продолжение подтверждено пользователем.")

    def toggle_pause(self) -> None:
        if self.pause_event.is_set():
            self.pause_event.clear()
            self.log("Пауза включена. Текущий шаг завершится, затем парсер остановится на паузе.")
        else:
            self.apply_live_settings_to_worker(log_message=True)
            self.pause_event.set()
            self.log("Парсер продолжает работу.")
        self.refresh_action_buttons()

    def stop_worker(self) -> None:
        self.stop_event.set()
        self.pause_event.set()
        self.ready_event.set()
        self.refresh_action_buttons()
        self.log("Получена команда Стоп. Ожидаю остановки текущего шага.")

    def open_output_folder(self) -> None:
        cfg = self.collect_cfg()
        out = Path(str(cfg.get("output_path") or "output/result.xlsx"))
        if not out.is_absolute():
            out = self.base_dir / out
        folder = out.parent
        folder.mkdir(parents=True, exist_ok=True)
        try:
            os.startfile(folder)  # type: ignore[attr-defined]
        except Exception as e:
            messagebox.showerror(APP_TITLE, f"Не удалось открыть папку: {e}")

    def log(self, msg: str) -> None:
        ts = datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}\n"
        if hasattr(self, "log_text"):
            self.log_text.insert(tk.END, line)
            self.log_text.see(tk.END)
        if hasattr(self, "snt_log_text"):
            self.snt_log_text.insert(tk.END, line)
            self.snt_log_text.see(tk.END)

    def on_close(self) -> None:
        worker = self.worker
        if worker and worker.is_alive():
            close_it = messagebox.askyesno(
                APP_TITLE,
                "Закрыть программу и временный браузер?\n\n"
                "Да — закрыть браузер и выйти.\n"
                "Нет — не закрывать программу.",
            )
            if not close_it:
                return
            self.stop_event.set()
            self.pause_event.set()
            self.ready_event.set()
            try:
                worker.next_task_event.set()
            except Exception:
                pass
            try:
                worker.join(timeout=5)
            except Exception:
                pass
        self.destroy()

    def _drain_log_queue(self) -> None:
        try:
            while True:
                msg = self.log_queue.get_nowait()
                self.log(msg)
        except queue.Empty:
            pass

        # После завершения задачи спрашиваем, закрывать ли временный браузер.
        # Если оставить браузер, следующая задача сможет стартовать в этой же сессии.
        try:
            worker = self.worker
            if (
                worker
                and worker.is_alive()
                and getattr(worker, "idle", False)
                and getattr(worker, "completed_task_id", 0) > self._prompted_close_task_id
            ):
                task_id = getattr(worker, "completed_task_id", 0)
                self._prompted_close_task_id = task_id
                close_it = messagebox.askyesno(
                    APP_TITLE,
                    "Задача завершена.\n\nЗакрыть временный браузер?\n\n"
                    "Да — закрыть браузер.\n"
                    "Нет — оставить открытым, чтобы запустить следующую задачу без новой авторизации.",
                )
                if close_it:
                    self.stop_worker()
                else:
                    self.log("Браузер оставлен открытым. Измените настройки и нажмите «Старт» для следующей задачи.")
        except Exception as e:
            self.log(f"Не удалось показать вопрос о закрытии браузера: {e}")

        self.refresh_action_buttons()
        self.after(200, self._drain_log_queue)


if __name__ == "__main__":
    App().mainloop()
